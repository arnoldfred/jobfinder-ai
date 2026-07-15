"""
JobFinder AI — Views IA
Chaque endpoint utilise le profil COMPLET du candidat pour des réponses
100% personnalisées, pertinentes et cohérentes.
"""
import json
import re
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.views.decorators.http import require_POST
from .groq_client import ask_groq, ask_groq_chat


# ═══════════════════════════════════════════════════════════════
#  CONSTRUCTION DU PROFIL COMPLET
# ═══════════════════════════════════════════════════════════════

def _get_full_profile(user):
    """
    Retourne un dict structuré avec TOUTES les données réelles du candidat.
    C'est la source unique de vérité pour tous les prompts IA.
    """
    try:
        p = user.profile
    except Exception:
        return {
            'nom': user.email, 'email': user.email,
            'titre_souhaite': '', 'localisation': '', 'phone': '',
            'resume': '', 'linkedin': '', 'github': '',
            'comp_tech': [], 'comp_tool': [], 'comp_soft': [], 'comp_lang': [],
            'all_skills': [],
            'experiences': [], 'formations': [],
            'annees_exp': 0, 'nb_exp': 0, 'nb_formations': 0,
            'cv_disponible': False, 'score_profil': 0,
            'domaines_exp': [],
        }

    from django.utils import timezone

    # ── Compétences par catégorie ──────────────────────────────
    comp_tech = []
    comp_tool = []
    comp_soft = []
    comp_lang = []

    for s in p.skills.all():
        level_str = f' ({s.get_level_display()})' if s.level else ''
        entry = s.name + level_str
        if s.category == 'technical':
            comp_tech.append(s.name)
        elif s.category == 'tool':
            comp_tool.append(s.name)
        elif s.category == 'soft':
            comp_soft.append(s.name)
        elif s.category == 'language':
            comp_lang.append(entry)

    all_skills = comp_tech + comp_tool + comp_soft + [s.split(' (')[0] for s in comp_lang]

    # ── Expériences DÉTAILLÉES ─────────────────────────────────
    total_months = 0
    experiences  = []
    domaines_exp = []

    for e in p.experiences.all().order_by('-start_date'):
        if e.start_date:
            end = e.end_date or timezone.now().date()
            months = (end.year - e.start_date.year) * 12 + end.month - e.start_date.month
            total_months += max(0, months)

        techs = [t.strip() for t in (e.technologies or '').split(',') if t.strip()]
        period = ''
        if e.start_date:
            period = e.start_date.strftime('%m/%Y')
            period += ' – ' + ('Présent' if e.is_current else (e.end_date.strftime('%m/%Y') if e.end_date else '?'))
            period += f'  [{e.duration}]'

        experiences.append({
            'titre':       e.title,
            'entreprise':  e.company,
            'lieu':        e.location or '',
            'periode':     period,
            'actuel':      e.is_current,
            'description': (e.description or '')[:400],
            'techs':       techs,
        })
        if e.company:
            domaines_exp.append(e.company)

    # ── Formations DÉTAILLÉES ──────────────────────────────────
    formations = []
    for ed in p.educations.all().order_by('-start_year'):
        formations.append({
            'diplome':    ed.degree,
            'institution': ed.institution,
            'lieu':        ed.location or '',
            'periode':     f'{ed.start_year} – {ed.end_year or "En cours"}',
            'mention':     ed.gpa or '',
        })

    annees_exp = round(total_months / 12, 1)

    return {
        'nom':            user.get_full_name() or user.email,
        'email':          user.email,
        'titre_souhaite': p.desired_title or '',
        'localisation':   p.location or '',
        'phone':          p.phone or '',
        'linkedin':       p.linkedin_url or '',
        'github':         p.github_url or '',
        'resume':         p.summary or '',
        'comp_tech':      comp_tech,
        'comp_tool':      comp_tool,
        'comp_soft':      comp_soft,
        'comp_lang':      comp_lang,
        'all_skills':     all_skills,
        'experiences':    experiences,
        'formations':     formations,
        'nb_exp':         len(experiences),
        'nb_formations':  len(formations),
        'annees_exp':     annees_exp,
        'domaines_exp':   domaines_exp,
        'cv_disponible':  bool(p.cv_file),
        'score_profil':   p.completion_score,
    }


def _profile_text(profile: dict) -> str:
    """
    Transforme le profil en bloc texte structuré, lisible par l'IA.
    Format délibérément dense et complet — l'IA doit tout voir.
    """
    nom = profile.get('nom', '')
    titre = profile.get('titre_souhaite', '') or 'Non précisé'
    loc   = profile.get('localisation', '') or 'Non précisée'
    annees = profile.get('annees_exp', 0)
    nb_exp = profile.get('nb_exp', 0)
    score  = profile.get('score_profil', 0)

    lines = [
        '┌─────────────────────────────────────────────────────┐',
        '│              PROFIL CANDIDAT — DONNÉES RÉELLES       │',
        '└─────────────────────────────────────────────────────┘',
        '',
        f'NOM            : {nom}',
        f'TITRE SOUHAITÉ : {titre}',
        f'LOCALISATION   : {loc}',
        f'PROFIL COMPLÉTÉ: {score}%',
        f'EXPÉRIENCE     : {annees} an(s) — {nb_exp} poste(s)',
        '',
    ]

    # Résumé
    resume = profile.get('resume', '')
    if resume:
        lines += ['RÉSUMÉ PROFESSIONNEL:', resume, '']

    # Compétences techniques
    comp_tech = profile.get('comp_tech', [])
    lines.append('COMPÉTENCES TECHNIQUES:')
    lines.append('  ' + (', '.join(comp_tech) if comp_tech else 'Aucune renseignée'))

    comp_tool = profile.get('comp_tool', [])
    if comp_tool:
        lines += ['', 'OUTILS & LOGICIELS:', '  ' + ', '.join(comp_tool)]

    comp_soft = profile.get('comp_soft', [])
    if comp_soft:
        lines += ['', 'SOFT SKILLS:', '  ' + ', '.join(comp_soft)]

    comp_lang = profile.get('comp_lang', [])
    if comp_lang:
        lines += ['', 'LANGUES:', '  ' + ', '.join(comp_lang)]

    # Expériences
    lines += ['', 'EXPÉRIENCES PROFESSIONNELLES:']
    exps = profile.get('experiences', [])
    if exps:
        for i, e in enumerate(exps, 1):
            lines.append(f'  [{i}] {e["titre"]} — {e["entreprise"]}')
            if e.get('lieu'):
                lines.append(f'       Lieu: {e["lieu"]}')
            if e.get('periode'):
                lines.append(f'       Période: {e["periode"]}')
            if e.get('techs'):
                lines.append(f'       Technologies: {", ".join(e["techs"])}')
            if e.get('description'):
                lines.append(f'       Missions: {e["description"][:300]}')
    else:
        lines.append('  ⚠ Aucune expérience renseignée')

    # Formations
    lines += ['', 'FORMATIONS:']
    fmts = profile.get('formations', [])
    if fmts:
        for f in fmts:
            lines.append(f'  • {f["diplome"]} — {f["institution"]} ({f["periode"]})')
            if f.get('mention'):
                lines.append(f'    Mention: {f["mention"]}')
    else:
        lines.append('  ⚠ Aucune formation renseignée')

    linkedin = profile.get('linkedin', '')
    github   = profile.get('github', '')
    if linkedin:
        lines += ['', f'LINKEDIN: {linkedin}']
    if github:
        lines += [f'GITHUB:   {github}']

    lines += ['', '──────────────────────────────────────────────────────']
    return '\n'.join(lines)


def _job_text(job_title='', company='', job_desc='', job_skills='', salary='') -> str:
    """Formate les données d'une offre pour les prompts."""
    parts = ['OFFRE D\'EMPLOI:']
    if job_title: parts.append(f'  Poste    : {job_title}')
    if company:   parts.append(f'  Société  : {company}')
    if salary:    parts.append(f'  Salaire  : {salary}')
    if job_skills:
        parts.append(f'  Compétences requises: {job_skills}')
    if job_desc:
        parts.append(f'  Description:\n{job_desc[:600]}')
    return '\n'.join(parts)


# ═══════════════════════════════════════════════════════════════
#  PAGE OUTILS
# ═══════════════════════════════════════════════════════════════

@login_required
def tools_page(request):
    has_cv = False
    try:
        has_cv = bool(request.user.profile.cv_file)
    except Exception:
        pass
    return render(request, 'ai_tools/tools.html', {'has_cv': has_cv})


# ═══════════════════════════════════════════════════════════════
#  CHAT ASSISTANT PERSONNALISÉ
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def chat(request):
    data = json.loads(request.body)
    msg  = data.get('message', '').strip()
    hist = data.get('history', [])
    if not msg:
        return JsonResponse({'error': 'Message vide'}, status=400)

    profile = _get_full_profile(request.user)
    nom     = profile['nom'].split()[0] if profile['nom'] else 'le candidat'

    system = (
        f'Tu es l\'assistant carrière JobFinder AI, expert emploi en Côte d\'Ivoire et Afrique de l\'Ouest.\n'
        f'Tu parles à {nom} et tu connais son profil complet ci-dessous.\n'
        f'Tes réponses sont TOUJOURS personnalisées : cite ses vraies compétences, '
        f'ses vraies expériences, son vrai niveau.\n'
        f'Sois direct, concret, actionnable. Maximum 150 mots. Réponds en français.\n\n'
        + _profile_text(profile)
    )

    messages = [{'role': 'system', 'content': system}] + hist[-6:] + [{'role': 'user', 'content': msg}]
    reply = ask_groq_chat(messages, max_tokens=400)
    return JsonResponse({'reply': reply})


# ═══════════════════════════════════════════════════════════════
#  ANALYSE CV — FORMAT CLAIR ET SECTIONS BIEN SÉPARÉES
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def analyze_cv(request):
    import io

    # Lecture du contenu CV
    cv_text    = ''
    target_job = ''
    job_desc   = ''

    if request.content_type and 'multipart' in request.content_type:
        cv_text    = request.POST.get('cv_text', '').strip()
        target_job = request.POST.get('target_job', '').strip()
        job_desc   = request.POST.get('job_description', '').strip()
        if 'cv_file' in request.FILES:
            f   = request.FILES['cv_file']
            raw = f.read()
            if f.name.lower().endswith('.pdf'):
                try:
                    from pypdf import PdfReader
                    cv_text = '\n'.join(p.extract_text() or '' for p in PdfReader(io.BytesIO(raw)).pages)
                except Exception:
                    pass
            elif f.name.lower().endswith('.docx'):
                try:
                    from docx import Document
                    cv_text = '\n'.join(p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text.strip())
                except Exception:
                    pass
    else:
        body       = json.loads(request.body)
        cv_text    = body.get('cv_text', '').strip()
        target_job = body.get('target_job', '').strip()
        job_desc   = body.get('job_description', '').strip()

        # Option: utiliser le CV enregistré sur le profil
        if not cv_text and body.get('use_profile_cv'):
            try:
                cv_file = request.user.profile.cv_file
                if cv_file:
                    raw = cv_file.read()
                    if cv_file.name.lower().endswith('.pdf'):
                        from pypdf import PdfReader
                        cv_text = '\n'.join(p.extract_text() or '' for p in PdfReader(io.BytesIO(raw)).pages)
                    elif cv_file.name.lower().endswith('.docx'):
                        from docx import Document
                        cv_text = '\n'.join(p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text.strip())
            except Exception:
                pass

    if not cv_text:
        return JsonResponse({'error': 'Importez votre CV (PDF/DOCX) ou collez son texte.'}, status=400)

    profile = _get_full_profile(request.user)

    # Contexte offre ciblée
    offre_ctx = ''
    job_context_hint = ''
    if target_job or job_desc:
        offre_ctx = (
            f'\n\nOFFRE CIBLÉE : {target_job}\n'
            + (f'Description : {job_desc[:500]}' if job_desc else '')
        )
        job_context_hint = (
            f'\n⚠ CONTEXTE : Cet analyse est SPÉCIFIQUEMENT pour le poste "{target_job}".'
            f'\n  Mets l\'accent sur les éléments du CV qui correspondent à ce rôle.'
        )

    system = (
        'Tu es un expert RH senior basé en Côte d\'Ivoire, '
        'spécialisé dans l\'optimisation de CVs pour le marché africain et les systèmes ATS.\n'
        'Tu analyses avec précision et bienveillance. '
        'Tu mentionnes les ÉLÉMENTS RÉELS du CV fourni, tu ne génères rien de générique.\n'
        'Tes scores reflètent objectivement la qualité réelle du CV.\n'
        'RÈGLE ABSOLUE SUR LA FORMATION : '
        'La hiérarchie est : BEP/CAP < BAC < BTS/DUT < Licence < Master/Master 2 < Doctorat. '
        'Si le candidat a un diplôme SUPÉRIEUR OU ÉGAL au requis, le score Formation est 8/10 minimum. '
        'Un Master 2 satisfait toute exigence BTS, Licence ou inférieure. '
        'Ne jamais pénaliser un candidat pour ne pas avoir un diplôme INFÉRIEUR à celui qu\'il possède.'
    )

    prompt = (
        _profile_text(profile) + offre_ctx +
        '\n\n═══ CV À ANALYSER ═══\n' + cv_text[:4500] +
        '\n\n═══ INSTRUCTIONS D\'ANALYSE ═══\n'
        + ('ANALYSE CONTEXTUALISÉE : Cette analyse cible spécifiquement le poste de ' + target_job + '\n'
           'Concentre-toi sur la pertinence pour CE rôle précis.\n\n' if target_job else '') +
        'Analyse ce CV en 5 sections séparées par des lignes vides. '
        'Chaque section commence par son titre en majuscules suivi de deux-points.\n\n'
        'SCORES (chaque score sur 100, justifié par une phrase):\n'
        '  Score Global: XX/100 — [justification basée sur CE CV' + (', par rapport au poste ' + target_job if target_job else '') + ']\n'
        '  Score ATS: XX/100 — [compatibilité systèmes de tri automatique]\n'
        '  Score Impact: XX/100 — [force des réalisations et chiffres]\n\n'
        'POINTS FORTS (3 éléments CONCRETS tirés du CV' + (', directement pertinents pour ' + target_job if target_job else ', liés au poste ciblé si précisé') + '):\n'
        '  1. [élément réel du CV]\n'
        '  2. [élément réel du CV]\n'
        '  3. [élément réel du CV]\n\n'
        'CORRECTIONS PRIORITAIRES (cite exactement ce qui est écrit dans le CV):\n'
        '  1. Section [X] — Problème: [ce qui ne va pas] → Correction: remplacez "[texte actuel]" par "[texte amélioré]"\n'
        '  2. ...\n\n'
        'MOTS-CLÉS MANQUANTS (pour l\'ATS' + (f' et spécifiques au poste {target_job}' if target_job else ' et le poste ciblé') + '):\n'
        '  - [mot-clé spécifique au secteur/poste]\n\n'
        'ACTION PRIORITAIRE:\n'
        '  [Une seule action concrète, la plus impactante' + (f', pour le poste {target_job}' if target_job else '') + ', à faire dans les 24h]'
    )

    result = ask_groq(system, prompt, max_tokens=1000, use_smart_model=True)

    # Extraction des scores pour les widgets visuels
    scores = []
    for pat, lbl in [
        (r'Score Global\s*:\s*(\d+)', 'Score Global'),
        (r'Score ATS\s*:\s*(\d+)',    'Score ATS'),
        (r'Score Impact\s*:\s*(\d+)', 'Score Impact'),
    ]:
        m = re.search(pat, result, re.IGNORECASE)
        if m:
            scores.append({'label': lbl, 'value': int(m.group(1))})

    return JsonResponse({'result': result, 'scores': scores})


# ═══════════════════════════════════════════════════════════════
#  ESTIMATION DES CHANCES — CALCUL RÉALISTE
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def estimate_chances(request):
    data            = json.loads(request.body)
    job_id          = data.get('job_id')
    job_title       = data.get('job_title', '')
    job_desc        = data.get('job_description', '')
    job_skills      = data.get('job_skills', '')
    job_implicit    = data.get('implicit_skills', '')
    salary          = data.get('salary', '')
    nlp_score       = data.get('nlp_score')  # Score NLP déjà affiché — ne pas le modifier

    profile         = _get_full_profile(request.user)

    system = (
        'Tu es un expert RH senior en Côte d\'Ivoire avec une connaissance approfondie des métiers.\n'
        'Tu comprends chaque poste AU-DELÀ des mots-clés : tu connais les vraies missions, les outils typiques, '
        'les livrables attendus et le niveau de séniorité requis pour chaque métier '
        '(Data Analyst, Développeur, Comptable, Marketing, RH, Commercial, Ingénieur, Chef de projet, etc.).\n'
        'Tu utilises cette connaissance métier pour évaluer si le profil correspond RÉELLEMENT au poste, '
        'même si l\'offre est mal rédigée ou incomplète.\n'
        'RÈGLE ABSOLUE SUR LA FORMATION : '
        'Hiérarchie : BEP/CAP < BAC < BTS/DUT < Licence < Master/Master 2/MIAGE < Doctorat. '
        'Diplôme supérieur ou égal au requis → score Formation 8/10 minimum. '
        'Ne jamais pénaliser un diplôme plus élevé pour absence d\'un diplôme inférieur.\n'
        + (
            f'Le score de compatibilité calculé est {nlp_score}%. '
            'Ton analyse doit être cohérente avec ce score : '
            f'{"encourage à postuler" if nlp_score >= 65 else "mets en avant les points à améliorer" if nlp_score >= 45 else "sois honnête sur les lacunes"}.\n'
            if nlp_score is not None else ''
        ) +
        'Réponds en français, sois direct et bienveillant.'
    )

    prompt = (
        _profile_text(profile)
        + '\n\n' + _job_text(job_title, '', job_desc, job_skills, salary)
        + (f'\n\nCOMPÉTENCES IMPLICITES DÉTECTÉES DANS LE PROFIL : {job_implicit}\n' if job_implicit else '')
        + '\nANALYSE DEMANDÉE :\n'
        + 'Évalue les chances de CE candidat précis pour CETTE offre précise.\n'
        + 'Base-toi sur ses vraies compétences, son vrai niveau d\'expérience, sa vraie localisation.\n'
        + (f'Pour chaque compétence implicite, explique clairement ce qui, dans le profil, le résumé ou les expériences, te pousse à la considérer comme implicite. '
           'Si elle n\'est pas mentionnée textuellement dans les compétences du profil, indique si elle est déduite d\'une réalisation, d\'un projet, d\'une fonction ou d\'un mot-clé présent dans le résumé.\n' if job_implicit else '')
        + (f'Les compétences implicites listées ci-dessus doivent être justifiées et expliquées par rapport à ce poste.\n' if job_implicit else '')
        + '\n'
        + 'FORMAT OBLIGATOIRE :\n\n'
        + 'VOS ATOUTS POUR CE POSTE:\n'
        + '- [compétence/expérience réelle du profil qui correspond à l\'offre]\n'
        + '- ...\n\n'
        + 'CE QUI MANQUE:\n'
        + '- [compétence/expérience requise absente du profil]\n'
        + '- ...\n\n'
        + 'ANALYSE PAR CRITÈRE:\n'
        + '- Compétences techniques : X/10 [justification courte]\n'
        + '- Années d\'expérience    : X/10 [justification courte]\n'
        + '- Formation              : X/10 [justification courte]\n'
        + '- Localisation           : X/10 [justification courte]\n\n'
        + 'VERDICT HONNÊTE:\n'
        + '[2 phrases directes sur les chances réelles de décrocher ce poste]\n\n'
        + '3 ACTIONS CONCRÈTES pour maximiser vos chances:\n'
        + '1. [action spécifique et réalisable]\n'
        + '2. [action spécifique et réalisable]\n'
        + '3. [action spécifique et réalisable]'
    )

    result = ask_groq(system, prompt, max_tokens=900, use_smart_model=True)

    return JsonResponse({'result': result})


# ═══════════════════════════════════════════════════════════════
#  QUESTIONS D'ENTRETIEN — ULTRA PERSONNALISÉES
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def interview_questions(request):
    data       = json.loads(request.body)
    job_title  = data.get('job_title', '')
    job_desc   = data.get('job_description', '')
    level      = data.get('level', 'Mid-Senior')
    qtype      = data.get('type', 'Toutes')

    profile    = _get_full_profile(request.user)
    nom        = profile['nom']

    system = (
        'Tu es un recruteur senior expert en Côte d\'Ivoire.\n'
        'Tu génères des questions d\'entretien ULTRA PERSONNALISÉES.\n'
        'Chaque question doit citer explicitement des éléments réels du profil : '
        'une technologie spécifique qu\'il connaît, une entreprise où il a travaillé, '
        'une durée d\'expérience précise, un diplôme qu\'il a obtenu.\n'
        'Les questions vagues et génériques sont INTERDITES.\n'
        'Réponds en français.'
    )

    gaps_note = ''
    req_skills = [s.strip().lower() for s in job_desc[:200].split() if len(s) > 4]
    missing = [s for s in (profile['comp_tech'] or []) if s.lower() not in ' '.join(req_skills)]
    if missing:
        gaps_note = f'\nGaps identifiés à creuser: {", ".join(missing[:3])}'

    prompt = (
        _profile_text(profile) + gaps_note +
        f'\n\n{_job_text(job_title, "", job_desc, "", "")}'
        f'\nNiveau requis: {level} | Type de questions: {qtype}\n\n'
        'Génère exactement 8 questions d\'entretien pour CE candidat pour CE poste.\n'
        'Chaque question doit être spécifique — elle doit contenir le nom d\'une de ses '
        'compétences, ou faire référence à une de ses expériences, ou cibler un gap précis.\n\n'
        'Format strict, séparé par --- :\n'
        'QUESTION: [question personnalisée citant des éléments réels du profil]\n'
        'TYPE: [Comportementale | Technique | Motivation | Situationnelle]\n'
        'POURQUOI: [pourquoi ce recruteur posera cette question à CE candidat spécifiquement]\n'
        'REPONSE: [réponse modèle STAR utilisant les vraies expériences du profil]\n'
        '---'
    )

    result    = ask_groq(system, prompt, max_tokens=2000, use_smart_model=True, temperature=0.4)
    questions = []
    for block in result.split('---'):
        lines = block.strip().split('\n')
        q = t = w = a = ''
        for line in lines:
            l = line.strip()
            if l.startswith('QUESTION:'):  q = l[9:].strip()
            elif l.startswith('TYPE:'):    t = l[5:].strip()
            elif l.startswith('POURQUOI:'): w = l[9:].strip()
            elif l.startswith('REPONSE:'): a = l[8:].strip()
        if q:
            questions.append({
                'question': q,
                'type_q':   t or 'Question d\'entretien',
                'why':      w,
                'answer':   a,
            })

    return JsonResponse({'questions': questions[:8], 'raw': result})


# ═══════════════════════════════════════════════════════════════
#  ÉVALUATION RÉPONSE — SCORING OBJECTIF MÉTHODE STAR
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def evaluate_answer(request):
    data      = json.loads(request.body)
    answer    = data.get('answer', '').strip()
    question  = data.get('question', '').strip()
    job_title = data.get('job_title', '').strip()
    q_num     = data.get('q_num', 1)
    q_total   = data.get('q_total', 8)

    if not answer:
        return JsonResponse({'error': 'Réponse vide'}, status=400)

    profile = _get_full_profile(request.user)
    nb_mots = len(answer.split())

    system = (
        'Tu es coach carrière expert, spécialisé en entretiens pour le marché africain.\n'
        'Tu évalues STRICTEMENT la réponse fournie — pas une réponse idéale théorique.\n'
        'Critères de scoring (total /10) :\n'
        '  Structure STAR présente   : 0–3 points\n'
        '  Précision (chiffres/faits): 0–2 points\n'
        '  Pertinence à la question  : 0–2 points\n'
        '  Usage des vraies expériences: 0–2 points\n'
        '  Impact et conviction      : 0–1 point\n\n'
        'CALIBRAGE SCORES :\n'
        '  1–3 : Réponse hors sujet ou totalement vague\n'
        '  4–5 : Éléments présents mais incomplets ou sans structure\n'
        '  6–7 : Bonne réponse avec quelques manques\n'
        '  8–9 : Très bonne réponse, structurée et précise\n'
        '  10  : Réponse parfaite (rare)\n\n'
        'Réponds UNIQUEMENT en JSON valide, sans texte avant ni après.'
    )

    prompt = (
        _profile_text(profile) +
        f'\n\nPOSTE VISÉ: {job_title or "Non précisé"}\n'
        f'Question {q_num}/{q_total}: {question}\n'
        f'Nombre de mots dans la réponse: {nb_mots}\n\n'
        f'RÉPONSE EXACTE DU CANDIDAT:\n"{answer}"\n\n'
        'Évalue cette réponse telle quelle. '
        'Si la réponse est courte (<20 mots) ou vague, le score doit être bas (1–4). '
        'Pour la version améliorée, utilise les vraies expériences du profil.\n\n'
        'JSON attendu:\n'
        '{\n'
        '  "score": 0-10,\n'
        '  "score_label": "label",\n'
        '  "strengths": ["point fort réel de cette réponse"],\n'
        '  "improvements": ["ce qui manque dans cette réponse"],\n'
        '  "star_feedback": "analyse de la structure STAR de cette réponse",\n'
        '  "improved_answer": "version améliorée utilisant les vraies expériences du profil",\n'
        '  "tip": "un seul conseil actionnable"\n'
        '}'
    )

    raw = ask_groq(system, prompt, max_tokens=900, use_smart_model=True, temperature=0.2)

    # Parse JSON robuste
    try:
        clean = re.sub(r'```json|```', '', raw).strip()
        match = re.search(r'\{.*\}', clean, re.DOTALL)
        if match:
            result_obj = json.loads(match.group())
            # Sanity check du score
            score = int(result_obj.get('score', 5))
            if nb_mots < 10 and score > 4:
                result_obj['score'] = min(score, 3)
            elif nb_mots < 30 and score > 6:
                result_obj['score'] = min(score, 5)
            return JsonResponse({'ok': True, 'result': result_obj, 'raw': raw})
    except Exception:
        pass

    return JsonResponse({'ok': False, 'result': None, 'raw': raw})


# ═══════════════════════════════════════════════════════════════
#  CONSEIL SALAIRE — BASÉ SUR LE PROFIL RÉEL
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def salary_advice(request):
    data    = json.loads(request.body)
    job     = data.get('job', '').strip()
    country = data.get('country', "Côte d'Ivoire")
    offer   = data.get('offer', '').strip()

    profile    = _get_full_profile(request.user)
    annees_exp = profile['annees_exp']
    nb_skills  = len(profile['comp_tech']) + len(profile['comp_tool'])

    system = (
        'Tu es expert en rémunération sur le marché du travail africain, '
        'particulièrement en Côte d\'Ivoire, Sénégal et Maroc.\n'
        'Tes fourchettes salariales sont RÉALISTES et basées sur le marché actuel 2024-2025.\n'
        'Tu tiens compte du profil réel du candidat pour personnaliser les montants.\n'
        'Montants en FCFA pour CI, en XOF pour Sénégal. Réponds en français.'
    )

    offer_ctx = f'\nOFFRE REÇUE: {offer} FCFA/mois' if offer else ''

    prompt = (
        _profile_text(profile) +
        f'\n\nPOSTE VISÉ: {job}\n'
        f'MARCHÉ: {country}\n'
        f'PROFIL: {annees_exp} ans d\'expérience, {nb_skills} compétences techniques/outils'
        + offer_ctx +
        '\n\nAnalyse salariale PERSONNALISÉE pour CE profil:\n\n'
        'FOURCHETTE DU MARCHÉ:\n'
        '  Junior (0-2 ans)    : [min] – [max] FCFA/mois\n'
        '  Confirmé (3-5 ans)  : [min] – [max] FCFA/mois\n'
        '  Senior (6+ ans)     : [min] – [max] FCFA/mois\n\n'
        'VOTRE POSITION:\n'
        f'  Avec {annees_exp} ans d\'exp. et votre profil, vous vous situez dans la fourchette [X].\n'
        '  Prétention recommandée: [montant] FCFA/mois\n'
        '  Justification: [pourquoi ce montant, basé sur vos vraies compétences]\n\n'
        + (
            'ANALYSE DE L\'OFFRE REÇUE:\n'
            '  [Est-ce dans la fourchette marché ? En dessous ? Au-dessus ?]\n'
            '  [Marge de négociation estimée]\n\n'
            if offer else ''
        ) +
        'SCRIPT DE NÉGOCIATION:\n'
        '  Version prudente  : "[formule de négociation adaptée au profil]"\n'
        '  Version assertive : "[formule plus directe utilisant vos atouts réels]"\n\n'
        'AVANTAGES À NÉGOCIER SI LE SALAIRE EST BLOQUÉ:\n'
        '  - [avantage pertinent pour ce type de poste]\n'
        '  - ...'
    )

    return JsonResponse({'result': ask_groq(system, prompt, max_tokens=800, use_smart_model=True)})


# ═══════════════════════════════════════════════════════════════
#  ANALYSE D'OFFRE — MATCH PRÉCIS PROFIL/OFFRE
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def analyze_job(request):
    data = json.loads(request.body)
    jd   = data.get('job_description', '').strip()
    if not jd:
        return JsonResponse({'error': 'Collez une description d\'offre.'}, status=400)

    profile = _get_full_profile(request.user)

    # Pré-calcul local pour calibrer le score IA
    words_offre  = set(re.findall(r'\b\w{4,}\b', jd.lower()))
    skills_lower = {s.lower() for s in profile['all_skills']}
    overlap      = skills_lower & words_offre
    rough_score  = min(95, max(10, int(len(overlap) / max(1, len(words_offre) * 0.15) * 100)))

    system = (
        'Tu es expert en matching candidat/offre en Côte d\'Ivoire.\n'
        'Tu compares PRÉCISÉMENT les exigences réelles de l\'offre avec le profil réel du candidat.\n'
        'Ton score est réaliste et calibré : un fort overlap compétences = score élevé.\n'
        f'Calibrage initial estimé (overlap keywords) : ~{rough_score}%.\n'
        'Cite toujours les éléments RÉELS du profil dans ton analyse. Réponds en français.'
    )

    prompt = (
        _profile_text(profile) +
        '\n\n═══ OFFRE À ANALYSER ═══\n' + jd +
        '\n\nANALYSE DE COMPATIBILITÉ:\n\n'
        'SCORE DE COMPATIBILITÉ: XX%\n'
        '(une phrase expliquant ce score en citant des éléments réels)\n\n'
        'VOS ATOUTS POUR CETTE OFFRE:\n'
        '- [compétence/expérience du profil qui correspond à une exigence précise de l\'offre]\n'
        '- ...\n\n'
        'GAPS IDENTIFIÉS:\n'
        '- [ce que l\'offre demande explicitement et qui est absent du profil]\n'
        '- ...\n\n'
        'VERDICT: [Postuler maintenant | Postuler en préparant ces points | Ne pas postuler]\n'
        'Justification: [2 phrases basées sur des éléments réels]\n\n'
        'MOTS-CLÉS ATS À INCLURE dans votre candidature:\n'
        '- [terme exact de l\'offre à reprendre]\n\n'
        'STRATÉGIE POUR SE DÉMARQUER:\n'
        '1. [action précise basée sur un atout réel du profil]\n'
        '2. [comment compenser un gap identifié]\n'
        '3. [angle d\'approche spécifique à ce candidat pour cette offre]'
    )

    result = ask_groq(system, prompt, max_tokens=900, use_smart_model=True)
    m = re.search(r'(\d{1,3})\s*%', result)
    return JsonResponse({'result': result, 'score': int(m.group(1)) if m else rough_score})


# ═══════════════════════════════════════════════════════════════
#  RÉDACTION EMAIL PRO
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def draft_email(request):
    data      = json.loads(request.body)
    etype     = data.get('type', 'followup')
    recipient = data.get('recipient', 'le/la Responsable RH')
    position  = data.get('position', '')
    tone      = data.get('tone', 'Professionnel')
    context   = data.get('context', '')
    
    # Détecte si on a un contexte
    has_job_context = bool(position or context)

    labels = {
        'followup':    'relance après candidature',
        'spontaneous': 'candidature spontanée',
        'thanks':      'remerciement après entretien',
        'decline':     'refus poli d\'une offre',
    }

    profile = _get_full_profile(request.user)
    nom     = profile['nom']
    titre   = profile['titre_souhaite'] or 'professionnel'

    system = (
        'Tu es expert en rédaction de courriers professionnels (marché ivoirien / Afrique de l\'Ouest).\n'
        'Tu rédiges des emails UNIQUES, adaptés au contexte exact fourni — jamais génériques ni répétitifs.\n'
        'Tu exploites TOUJOURS les vraies compétences, expériences et réalisations du profil.\n'
        + ('STRATÉGIE CONTEXTUELLE :\n'
           '  - Si poste et/ou description fournis : email TRÈS SPÉCIFIQUE à cette offre/entreprise\n'
           '  - Si aucun contexte : email générique mais professionnel et personnalisé au profil\n' if has_job_context else
           'Reste à un niveau générique mais professionnel.\n')
        + 'INTERDICTIONS ABSOLUES:\n'
        '- Ne jamais commencer par "Je me permets de vous contacter"\n'
        '- Ne jamais utiliser "Je reste à votre disposition", "N\'hésitez pas à me contacter", "Cordialement"\n'
        '- Zéro phrase générique copiable pour n\'importe quel candidat\n'
        '- Aucun placeholder entre crochets [...]\n'
        'Chaque email doit avoir une accroche différente, ancrée dans le contexte spécifique (poste, entreprise, secteur).\n'
        'Style : soutenu, professionnel, humain. Réponds en français.'
    )

    # Formules de clôture variées selon le type
    closings = {
        'followup':    'Dans l\'attente de votre retour, je vous adresse mes cordiales salutations.',
        'spontaneous': 'Je serais ravi(e) d\'échanger avec vous sur la contribution que je pourrais apporter à votre équipe.',
        'thanks':      'Cet entretien a renforcé mon intérêt pour ce poste et je reste pleinement disponible pour la suite.',
        'decline':     'Je vous remercie de votre compréhension et espère qu\'une collaboration pourra se présenter à l\'avenir.',
    }
    closing = closings.get(etype, 'Dans l\'attente de votre retour, veuillez recevoir mes cordiales salutations.')

    ctx_note = f'\nContexte supplémentaire: {context}' if context else ''

    # Instructions d'accroche spécifiques au type ET au contexte
    accroche_rules = {
        'followup': (
            'Accroche : rappelle ta candidature avec une date précise et un enthousiasme renouvelé — sans répéter mot pour mot ta lettre initiale.' if has_job_context else
            'Accroche : rappelle ta candidature avec une date précise et un enthousiasme renouvelé.'
        ),
        'spontaneous': (
            'Accroche : commence par ce qui t\'attire SPÉCIFIQUEMENT dans CETTE ENTREPRISE ou SECTEUR (innovation concrète, projet mentionné, valeur reconnue) — PAS "je me permets".' if has_job_context else
            'Accroche : commence par ce qui t\'attire dans le secteur/domaine en général — PAS "je me permets".'
        ),
        'thanks': (
            'Accroche : cite UN élément SPÉCIFIQUE de l\'entretien (sujet exact abordé, valeur partagée, projet discuté).' if has_job_context else
            'Accroche : cite un élément marquant de l\'entretien (sujet abordé, échange positif).'
        ),
        'decline': (
            'Accroche : remercie chaleureusement pour CETTE OPPORTUNITÉ SPÉCIFIQUE tout en déclinant avec élégance.' if has_job_context else
            'Accroche : remercie chaleureusement tout en déclinant avec élégance.'
        ),
    }
    accroche_rule = accroche_rules.get(etype, 'Accroche : phrase d\'ouverture percutante ancrée dans le contexte.')

    prompt = (
        _profile_text(profile) +
        f'\n\nTYPE D\'EMAIL: {labels.get(etype, "email professionnel")}\n'
        f'Entreprise / destinataire: {recipient or "le recruteur"}\n'
        f'Poste concerné: {position or titre}\n'
        f'Ton souhaité: {tone}'
        + ctx_note +
        ('\n\n⚡ CONTEXTE POSTE/OFFRE FOURNI : Rédige un email TRÈS ADAPTÉ À CETTE OFFRE ET CETTE ENTREPRISE.\n'
         '→ Cite des éléments spécifiques de l\'offre ou de l\'entreprise si fournis\n'
         '→ Montre que tu connais CE QU\'ON CHERCHE\n' if has_job_context else
         '\n\n⚡ AUCUN CONTEXTE POSTE : Rédige un email générique mais personnalisé.\n'
         '→ Basé sur le profil général\n'
         '→ Ne mentionne pas une offre inexistante\n') +
        '\n\nRédige UNIQUEMENT le corps de la lettre (sans "De:", "OBJET:", ni en-tête).\n'
        'STRUCTURE:\n'
        '- Salutation: "Madame, Monsieur,"\n'
        f'- {accroche_rule}\n'
        '- Paragraphe valeur : cite 2–3 compétences/réalisations RÉELLES du profil avec des verbes d\'action forts et des chiffres si disponibles\n'
        '- Paragraphe closing : appel à l\'action affirmé, projection concrète dans le rôle\n'
        f'- Formule de clôture : "{closing}"\n'
        '- Ligne vide puis nom complet\n'
        '- 150–180 mots maximum\n'
        '- Vocabulaire riche et varié — chaque phrase doit apporter quelque chose de nouveau'
    )

    return JsonResponse({'result': ask_groq(system, prompt, max_tokens=600, temperature=0.85)})


# ═══════════════════════════════════════════════════════════════
#  PLAN CARRIÈRE 5 ANS — BASÉ SUR LE PROFIL RÉEL
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def career_path(request):
    data        = json.loads(request.body)
    goal        = data.get('goal', '')
    sector      = data.get('sector', '')
    constraints = data.get('constraints', '')

    profile    = _get_full_profile(request.user)
    annees_exp = profile['annees_exp']
    titre      = profile['titre_souhaite'] or 'professionnel'

    system = (
        'Tu es coach carrière senior spécialisé marché africain (CI, Sénégal, Maroc).\n'
        'Tu construis des plans carrière RÉALISTES basés sur le profil réel du candidat.\n'
        'Cite ses vraies compétences, ses vraies lacunes, son vrai niveau actuel.\n'
        'Les entreprises citées sont de vraies entreprises africaines dans ce secteur.\n'
        'Réponds en français.'
    )

    constraints_note = f'\nContraintes: {constraints}' if constraints else ''

    prompt = (
        _profile_text(profile) +
        f'\n\nOBJECTIF: {goal or "Évoluer dans mon domaine"}\n'
        f'SECTEUR CIBLE: {sector or "Non précisé"}'
        + constraints_note +
        f'\n\nPLAN CARRIÈRE 5 ANS pour {titre} avec {annees_exp} an(s) d\'expérience:\n\n'
        'SITUATION ACTUELLE:\n'
        '  [Évaluation honnête du profil : forces, lacunes, positionnement marché]\n\n'
        'ÉTAPE 1 (6–12 mois) — Consolidation:\n'
        '  Objectif: [basé sur le profil actuel]\n'
        '  Actions: [compétences à renforcer, certifications à obtenir]\n'
        '  Entreprises cibles réalistes en Côte d\'Ivoire: [noms réels]\n\n'
        'ÉTAPE 2 (12–24 mois) — Montée en compétences:\n'
        '  Objectif: [poste suivant réaliste]\n'
        '  Formations spécifiques: [certifications, formations concrètes]\n'
        '  Salaire cible: [fourchette FCFA]\n\n'
        'ÉTAPE 3 (2–3 ans) — Positionnement sénior:\n'
        '  Objectif: []\n'
        '  Entreprises cibles: []\n\n'
        'ÉTAPE 4 (3–5 ans) — Leadership:\n'
        '  Objectif: []\n\n'
        'PLAN B (si objectif principal ne se réalise pas):\n'
        '  [Alternative réaliste basée sur les compétences actuelles]\n\n'
        '3 ACTIONS IMMÉDIATES À FAIRE CETTE SEMAINE:\n'
        '1. [action précise et réalisable]\n'
        '2. [action précise et réalisable]\n'
        '3. [action précise et réalisable]'
    )

    return JsonResponse({'result': ask_groq(system, prompt, max_tokens=1000, use_smart_model=True)})


# ═══════════════════════════════════════════════════════════════
#  GÉNÉRATION DOCUMENT — CV ADAPTÉ OU LETTRE DE MOTIVATION
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def generate_doc(request, doc_type):
    data       = json.loads(request.body)
    job_title  = data.get('job_title', '').strip()
    company    = data.get('company', '').strip()
    job_desc   = data.get('job_description', '').strip()
    job_skills = data.get('job_skills', '').strip()
    lm_style   = data.get('lm_style', 'classique').strip().lower()

    profile    = _get_full_profile(request.user)
    nom        = profile['nom']
    loc        = profile['localisation']
    titre_dyn  = profile['titre_souhaite']
    
    # Détecte si on a un contexte d'offre (poste/entreprise/description)
    has_context = bool(job_title or company or job_desc)

    if doc_type == 'cv':
        system = (
            'Tu es expert CV pour le marché africain. '
            'RÈGLE ABSOLUE : utilise UNIQUEMENT les informations du profil fourni. '
            'ADAPTATION CONTEXTUELLE : '
            '  - Si une offre est fournie (poste/entreprise/description), adapte LE CV À CETTE OFFRE SPÉCIFIQUE. '
            '  - Si AUCUNE offre n\'est fournie, génère un CV GÉNÉRIQUE équilibré sans inventer de contexte. '
            '  - JAMAIS d\'invention de compétences, diplômes, entreprises ou réalisations manquantes. '
            'Le CV doit être optimisé ATS et tenir en 1 page. Réponds en français.'
        )

        # Calcule les compétences qui matchent réellement
        req_list = [s.strip() for s in job_skills.split(',') if s.strip()]
        cand_tech = profile['comp_tech']
        cand_tool = profile['comp_tool']
        matching  = [s for s in req_list if any(s.lower() in c.lower() or c.lower() in s.lower()
                                                 for c in cand_tech + cand_tool)]
        match_note = (f'\nCompétences requises matchées dans le profil: {", ".join(matching)}'
                      if matching else '')

        # Titre du poste ciblé pour l'en-tête du CV
        poste_cible = job_title if job_title else titre_dyn
        
        # Séparer soft skills et hard skills depuis le profil
        soft_keywords = {'communication', 'travail en équipe', 'leadership', 'autonomie',
                         'rigueur', 'organisation', 'adaptabilité', 'esprit d\'analyse',
                         'gestion du stress', 'créativité', 'ponctualité', 'écoute',
                         'prise de décision', 'esprit d\'initiative', 'sens des responsabilités'}
        all_soft = profile.get('comp_soft') or []
        all_tech = (profile.get('comp_tech') or []) + (profile.get('comp_tool') or [])
        # Si pas de comp_soft dédié, extraire depuis comp_tech les mots-clés doux
        if not all_soft:
            all_soft = [c for c in all_tech if any(k in c.lower() for k in soft_keywords)]
            all_tech = [c for c in all_tech if c not in all_soft]

        soft_str = ', '.join(all_soft) if all_soft else '[Rigueur, autonomie, esprit d\'analyse, communication]'
        hard_str = ', '.join(all_tech) if all_tech else '[Compétences techniques du profil]'

        prompt = (
            _profile_text(profile) + match_note +
            (f'\n\n{_job_text(job_title, company, job_desc, job_skills)}' if has_context else '') +
            '\n\nRÈGLES DE GÉNÉRATION:\n' +
            '⛔ INTERDIT: inventer des compétences, diplômes, entreprises, réalisations\n' +
            '⛔ INTERDIT: inclure des compétences requises si elles sont absentes du profil\n' +
            '⛔ INTERDIT: ajouter du contenu générique sans pertinence\n' +
            ('✅ OFFRE FOURNIE : Adapte le CV À CE POSTE SPÉCIFIQUE, reformule les missions avec ses mots-clés\n'
             '✅ Réorganise l\'ordre des sections pour montrer la pertinence maximale\n'
             '✅ Met en avant UNIQUEMENT les compétences/expériences qui correspondent À CETTE OFFRE\n' if has_context else 
             '✅ AUCUNE OFFRE : Génère un CV générique équilibré basé sur le profil\n'
             '✅ Utilise des termes larges et transférables\n'
             '✅ Ordonne par pertinence générale, pas par offre spécifique\n') +
            '\n\nGÉNÈRE LE CV ADAPTÉ en respectant EXACTEMENT cette structure:\n\n' +
            f'**{nom}**\n' +
            (f'Candidature : {poste_cible}\n' if has_context else '') +
            f'{loc}\n\n' +
            '**RÉSUMÉ PROFESSIONNEL**\n' +
            ('2-3 phrases précises ciblant CE POSTE SPÉCIFIQUE, basées sur le vrai profil\n' if has_context else 
             '2-3 phrases générales sur le profil, ses forces principales et adaptabilité\n') +
            '\n**COMPÉTENCES TECHNIQUES (Hard Skills)**\n' +
            ('Sélectionne UNIQUEMENT parmi : ' + hard_str + ' — celles qui matchent CETTE OFFRE\n' if has_context else
             'Sélectionne parmi : ' + hard_str + ' — les plus transférables et reconnues\n') +
            '\n**COMPÉTENCES COMPORTEMENTALES (Soft Skills)**\n' +
            ('Sélectionne parmi : ' + soft_str + ' — les plus PERTINENTES POUR CE RÔLE\n' if has_context else
             'Sélectionne parmi : ' + soft_str + ' — les plus universelles et applicables\n') +
            '\n**EXPÉRIENCES PROFESSIONNELLES**\n' +
            ('Vraies expériences, reformulées avec LES MOTS-CLÉS DE CETTE OFFRE\n' if has_context else
             'Vraies expériences, ordonnées par pertinence générale et impact\n') +
            '\n**FORMATION**\n' +
            'Vraies formations du profil\n\n' +
            '**LANGUES**\n' +
            'Vraies langues avec niveaux'
        )

    else:  # cover_letter — styles multiples
        best_exps = profile.get('experiences', [])[:2]
        exp_note = ''
        if best_exps:
            exp_note = '\nExpériences clés à valoriser:\n'
            for exp in best_exps:
                exp_note += f"  - {exp.get('titre', '')} chez {exp.get('entreprise', '')} ({exp.get('periode', '')})\n"

        dest_role = 'Responsable du recrutement' if job_title else 'Responsable RH'
        dest_co = company or 'le service RH'

        if has_context:
            context_header = f'OFFRE SPÉCIFIÉE: "{job_title}" chez {company or "[entreprise inconnue]"}\n'
            context_note = (
                '→ La lettre doit être TRÈS SPÉCIFIQUE À CETTE OFFRE\n'
                '→ Cite les challenges/projets de cette offre si mentionnés\n'
                '→ Montre comment le profil résout LES PROBLÈMES SPÉCIFIQUES de ce rôle\n\n'
            )
        else:
            context_header = 'AUCUNE OFFRE SPÉCIFIÉE — Lettre générale\n'
            context_note = (
                '→ Reste à un niveau générique mais pertinent\n'
                '→ Ne mentionne pas une offre spécifique qui n\'existe pas\n'
                '→ Focalise sur le profil et ses forces intrinsèques\n\n'
            )

        header_lines = (
            f'DESTINATAIRE: {dest_role} chez {dest_co}\n'
            f'OBJET: Candidature au poste de {job_title or "votre offre"}\n\n'
        )

        rules_communes = (
            f"{context_header}{context_note}"
            "RÈGLES DE FORMAT OBLIGATOIRES:\n"
            "- Les deux premières lignes DOIVENT commencer par DESTINATAIRE: et OBJET:\n"
            "- Aucun titre ni libellé dans les paragraphes\n"
            "- NE PAS inclure la formule de politesse finale, la date, l'adresse ni la signature "
            "(elles sont ajoutées automatiquement)\n"
            "- Chaque paragraphe commence par 4 espaces d'indentation\n"
            "- Paragraphes séparés par une ligne vide\n"
            "- Utilise UNIQUEMENT les vraies informations du profil — n'invente rien\n"
            "- INTERDIT : tirets (-), puces (•), listes — uniquement des paragraphes en prose\n\n"
            "STRUCTURE INSPIRÉE DU MODÈLE DE LETTRE FRANÇAISE PROFESSIONNELLE :\n"
            "Chaque paragraphe doit être complet, développé, fluide — comme dans une vraie lettre "
            "rédigée par un humain qualifié. Pas de phrases isolées. Pas d'accroche en \"Je m'appelle\"."
        )

        if lm_style == 'impact':
            system = 'Tu rédiges des lettres de motivation axées sur les résultats et les chiffres. Réponds en français.'
            style_instructions = 'STYLE : Impact & Résultats.\n\nParagraphe 1 : accroche directe et spécifique.\n\nParagraphe 2 : développe 2 à 3 accomplissements concrets.\n\nParagraphe 3 : projection claire sur le poste ciblé.'
        elif lm_style == 'storytelling':
            system = 'Tu rédiges des lettres de motivation sous forme de récit personnel captivant. Réponds en français.'
            style_instructions = 'STYLE : Storytelling.\n\nParagraphe 1 : raconte une expérience qui te prépare à ce poste.\n\nParagraphe 2 : développe le récit avec des éléments concrets.\n\nParagraphe 3 : conclue avec une motivation naturelle.'
        elif lm_style == 'enthousiaste':
            system = 'Tu rédiges des lettres de motivation avec un ton chaleureux, enthousiaste et humain. Réponds en français.'
            style_instructions = 'STYLE : Enthousiaste & Humain.\n\nParagraphe 1 : exprime sincèrement ce qui t attire dans l offre.\n\nParagraphe 2 : parle de tes expériences avec énergie.\n\nParagraphe 3 : termine avec une envie claire de rejoindre l équipe.'
        else:
            system = 'Tu rédiges des lettres de motivation professionnelles au format français classique. Réponds en français.'
            style_instructions = 'STYLE : Classique professionnel.\n\nParagraphe 1 : accroche personnelle et pertinente.\n\nParagraphe 2 : développe des réalisations concrètes.\n\nParagraphe 3 : motive sincèrement pour le poste.'

        prompt_parts = [
            _profile_text(profile),
            exp_note,
            (f'\n\n{_job_text(job_title, company, job_desc, job_skills)}' if has_context else ''),
            '\n\nGÉNÈRE LA LETTRE DE MOTIVATION en respectant EXACTEMENT ce format:\n\n',
            header_lines,
            style_instructions,
            '\n\n',
            rules_communes,
        ]
        prompt = ''.join(part for part in prompt_parts if part is not None)

    max_tok = 520 if doc_type == 'cover_letter' else 800
    result = ask_groq(system, prompt, max_tokens=max_tok, use_smart_model=True, temperature=0.3)
    return JsonResponse({'result': result, 'doc_type': doc_type})


# ═══════════════════════════════════════════════════════════════
#  IMPORT CV — EXTRACTION IA + PRÉ-REMPLISSAGE PROFIL
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def import_cv(request):
    """Lit un CV (PDF/DOCX/texte) et extrait les données structurées via l'IA."""
    import io
    cv_text = ''

    # Cas 1: JSON avec use_profile_cv=True — retourner le texte du CV enregistré
    if request.content_type and 'application/json' in request.content_type:
        try:
            req_json = json.loads(request.body)
            if req_json.get('use_profile_cv'):
                cv_file = request.user.profile.cv_file
                if not cv_file:
                    return JsonResponse({'ok': False, 'error': 'Aucun CV enregistré sur votre profil.'})
                raw = cv_file.read()
                if cv_file.name.lower().endswith('.pdf'):
                    from pypdf import PdfReader
                    cv_text = '\n'.join(p.extract_text() or '' for p in PdfReader(io.BytesIO(raw)).pages)
                elif cv_file.name.lower().endswith('.docx'):
                    from docx import Document
                    cv_text = '\n'.join(p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text.strip())
                if not cv_text:
                    return JsonResponse({'ok': False, 'error': 'Impossible de lire le fichier CV.'})
                return JsonResponse({'ok': True, 'cv_text': cv_text[:5000]})
        except Exception as e:
            return JsonResponse({'ok': False, 'error': str(e)})

    # Cas 2: upload fichier
    if 'cv_file' in request.FILES:
        f   = request.FILES['cv_file']
        raw = f.read()
        try:
            if f.name.lower().endswith('.pdf'):
                from pypdf import PdfReader
                cv_text = '\n'.join(p.extract_text() or '' for p in PdfReader(io.BytesIO(raw)).pages)
            elif f.name.lower().endswith('.docx'):
                from docx import Document
                cv_text = '\n'.join(p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text.strip())
            else:
                cv_text = raw.decode('utf-8', errors='ignore')
        except ImportError as ie:
            return JsonResponse({'ok': False, 'error': f'Module manquant: {ie}. pip install pypdf python-docx'})
        except Exception as e:
            return JsonResponse({'ok': False, 'error': f'Erreur lecture fichier: {e}'})
    else:
        try:
            cv_text = json.loads(request.body).get('cv_text', '')
        except Exception:
            pass

    if not cv_text.strip():
        return JsonResponse({'ok': False, 'error': 'CV vide ou illisible.'})

    # Extraction IA — JSON structuré
    system = (
        'Tu es expert RH africain spécialisé en extraction de données de CV.\n'
        'Analyse le CV et extrais TOUTES les informations en JSON strict.\n'
        'Réponds UNIQUEMENT avec le JSON — aucun texte avant ou après, pas de backticks.\n'
        'Sois exhaustif : langues, outils, certifications, soft skills.\n'
        'Dates au format YYYY-MM-DD (ou YYYY si seule l\'année est disponible).\n'
        'Champs absents du CV = chaîne vide ou tableau vide.'
    )

    prompt = (
        'Analyse ce CV et réponds UNIQUEMENT avec ce JSON valide:\n'
        '{\n'
        '  "name": "Prénom Nom",\n'
        '  "email": "email@exemple.com",\n'
        '  "phone": "+225...",\n'
        '  "location": "Abidjan, Côte d\'Ivoire",\n'
        '  "title": "Titre professionnel principal",\n'
        '  "summary": "Résumé professionnel 2-3 phrases",\n'
        '  "skills_technical": ["Python", "SQL", "Django"],\n'
        '  "skills_tools": ["Excel", "Power BI", "Figma"],\n'
        '  "skills_soft": ["Leadership", "Communication"],\n'
        '  "skills_certifications": ["AWS Certified", "PMP"],\n'
        '  "languages": [{"name": "Français", "level": "bilingue"}, {"name": "Anglais", "level": "avance"}],\n'
        '  "experiences": [\n'
        '    {\n'
        '      "title": "Titre du poste",\n'
        '      "company": "Entreprise",\n'
        '      "location": "Ville",\n'
        '      "start_date": "2022-01-01",\n'
        '      "end_date": "2024-06-01",\n'
        '      "is_current": false,\n'
        '      "description": "Missions principales",\n'
        '      "technologies": "Python, SQL, React"\n'
        '    }\n'
        '  ],\n'
        '  "educations": [\n'
        '    {\n'
        '      "degree": "Licence Informatique",\n'
        '      "institution": "Université",\n'
        '      "location": "Ville",\n'
        '      "start_year": 2019,\n'
        '      "end_year": 2022,\n'
        '      "gpa": "Mention Bien"\n'
        '    }\n'
        '  ]\n'
        '}\n\n'
        'CV à analyser:\n\n' + cv_text[:4500]
    )

    raw_result = ask_groq(system, prompt, max_tokens=1800, use_smart_model=True, temperature=0.1)

    try:
        import json as json_lib
        clean = re.sub(r'```json|```', '', raw_result).strip()
        match = re.search(r'\{.*\}', clean, re.DOTALL)
        if match:
            data_extracted = json_lib.loads(match.group())
            return JsonResponse({'ok': True, 'data': data_extracted, 'cv_text': cv_text[:500]})
        return JsonResponse({'ok': False, 'error': 'Impossible de parser les données du CV.', 'raw': raw_result[:300]})
    except Exception as e:
        return JsonResponse({'ok': False, 'error': f'Erreur parsing: {e}', 'raw': raw_result[:300]})


# ═══════════════════════════════════════════════════════════════
#  GÉNÉRATION PDF
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
@login_required
@require_POST
def generate_pdf(request):
    """
    Génère un PDF professionnel : Times New Roman 12pt, interligne 1.5.
    Photo bien placée en en-tête pour CV et LM.
    """
    import io, re as re_mod
    from django.http import HttpResponse

    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({'error': 'JSON invalide'}, status=400)

    content_text  = data.get('content', '')
    title         = data.get('title', 'Document JobFinder AI')
    include_photo = data.get('include_photo', False)

    if not content_text:
        return JsonResponse({'error': 'Contenu vide'}, status=400)

    user        = request.user
    user_name   = user.get_full_name() or user.email
    user_email  = user.email
    user_phone  = ''
    user_loc    = ''
    user_linkedin = ''
    photo_path  = None
    try:
        p = user.profile
        user_loc      = p.location or ''
        user_phone    = p.phone or ''
        user_linkedin = p.linkedin_url or ''
        if include_photo and p.avatar:
            photo_path = p.avatar.path
    except Exception:
        pass

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                         HRFlowable, Table, TableStyle, KeepTogether)
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # ── Couleurs discrètes ───────────────────────────────────────────────
        DARK   = colors.HexColor('#1A1A2E')
        GRAY   = colors.HexColor('#555555')
        LIGHT  = colors.HexColor('#888888')
        ACCENT = colors.HexColor('#C9A84C')  # or premium
        LBORDER= colors.HexColor('#C8C8C8')

        is_cv    = 'cv' in title.lower() or 'curriculum' in title.lower()
        is_lm    = 'lettre' in title.lower() or 'motivation' in title.lower()
        is_email = title.lower().startswith('email_') or title.lower().startswith('email ')

        # ── Page setup ──────────────────────────────────────────────────────
        buf = io.BytesIO()
        # LM/Email : marges légèrement réduites pour tenir sur 1 page
        _top = 1.8*cm if (is_lm or is_email) else 2.0*cm
        _bot = 1.8*cm if (is_lm or is_email) else 2.0*cm
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=2.5*cm, rightMargin=2.0*cm,
            topMargin=_top, bottomMargin=_bot
        )

        # ── Styles Times New Roman ───────────────────────────────────────────
        FS   = 12
        LD   = 17 if (is_lm or is_email) else 18   # interligne légèrement réduit pour LM

        body_style = ParagraphStyle(
            'Body',
            fontName='Times-Roman', fontSize=FS, leading=LD,
            spaceAfter=4, spaceBefore=0, textColor=DARK,
            alignment=TA_JUSTIFY
        )
        body_left = ParagraphStyle(
            'BodyLeft',
            parent=body_style, alignment=TA_LEFT
        )
        h1_style = ParagraphStyle(
            'H1',
            fontName='Times-Bold', fontSize=18, leading=22,
            spaceAfter=2, spaceBefore=0, textColor=DARK
        )
        h2_style = ParagraphStyle(
            'H2',
            fontName='Times-Bold', fontSize=13, leading=18,
            spaceAfter=4, spaceBefore=10, textColor=ACCENT
        )
        h3_style = ParagraphStyle(
            'H3',
            fontName='Times-Bold', fontSize=FS, leading=LD,
            spaceAfter=2, spaceBefore=6, textColor=DARK
        )
        subtitle_style = ParagraphStyle(
            'Sub',
            fontName='Times-Roman', fontSize=11, leading=16,
            spaceAfter=2, textColor=GRAY
        )
        bullet_style = ParagraphStyle(
            'Bullet', parent=body_style,
            leftIndent=22, firstLineIndent=-10,
            spaceAfter=3, spaceBefore=1,
            alignment=TA_LEFT
        )
        sub_bullet_style = ParagraphStyle(
            'SubBullet', parent=body_style,
            leftIndent=38, firstLineIndent=-10,
            spaceAfter=2, spaceBefore=0,
            alignment=TA_LEFT
        )
        footer_style = ParagraphStyle(
            'Footer',
            fontName='Times-Roman', fontSize=9, leading=12,
            textColor=LIGHT, alignment=TA_CENTER
        )
        center_style = ParagraphStyle(
            'Center', parent=body_style, alignment=TA_CENTER
        )

        story = []
        skip_normal_body = False

        # ── EN-TÊTE ──────────────────────────────────────────────────────────
        if is_cv:
            # Extraire le titre de poste ciblé depuis le contenu généré
            _poste_cv = ''
            for _ln in content_text.split('\n'):
                _m = re_mod.match(r'^Candidature\s*:\s*(.+)', _ln.strip(), re_mod.IGNORECASE)
                if _m:
                    _poste_cv = _m.group(1).strip()
                    break

            if photo_path and include_photo:
                try:
                    from reportlab.platypus import Image as RLImage
                    img = RLImage(photo_path, width=3.0*cm, height=3.5*cm)
                    img.hAlign = 'CENTER'

                    # Bloc texte identité
                    id_block = []
                    id_block.append(Paragraph(user_name, h1_style))

                    # Sous-titre (titre souhaité si dispo)
                    try:
                        desired = request.user.profile.desired_title or ''
                        if desired:
                            id_block.append(Paragraph(desired, subtitle_style))
                    except Exception:
                        pass

                    # Coordonnées
                    coords = []
                    if user_phone:   coords.append(user_phone)
                    if user_email:   coords.append(user_email)
                    if user_loc:     coords.append(user_loc)
                    if user_linkedin: coords.append(user_linkedin)
                    if coords:
                        coord_text = ' | '.join(coords)
                        id_block.append(Spacer(1, 3))
                        id_block.append(Paragraph(coord_text, subtitle_style))

                    # Tableau identité (gauche) + photo (droite) — ATS-friendly
                    header_table = Table(
                        [[id_block, img]],
                        colWidths=[13.0*cm, 3.5*cm]
                    )
                    header_table.setStyle(TableStyle([
                        ('VALIGN',       (0, 0), (-1, -1), 'MIDDLE'),
                        ('LEFTPADDING',  (0, 0), (0, 0),  0),
                        ('RIGHTPADDING', (0, 0), (0, 0),  12),
                        ('LEFTPADDING',  (1, 0), (1, 0),  0),
                        ('RIGHTPADDING', (1, 0), (1, 0),  0),
                    ]))
                    story.append(header_table)
                except Exception as photo_err:
                    # Fallback sans photo
                    story.append(Paragraph(user_name, h1_style))
            else:
                # En-tête sans photo — centré
                story.append(Paragraph(user_name, ParagraphStyle(
                    'N', fontName='Times-Bold', fontSize=18, leading=22,
                    alignment=TA_CENTER, textColor=DARK
                )))
                # Titre souhaité
                try:
                    desired = request.user.profile.desired_title or ''
                    if desired:
                        story.append(Paragraph(desired, ParagraphStyle(
                            'DS', fontName='Times-Roman', fontSize=12, leading=16,
                            alignment=TA_CENTER, textColor=ACCENT, spaceAfter=4
                        )))
                except Exception:
                    pass
                # Coordonnées
                coords = [x for x in [user_phone, user_email, user_loc, user_linkedin] if x]
                if coords:
                    story.append(Paragraph(' | '.join(coords), ParagraphStyle(
                        'CO', fontName='Times-Roman', fontSize=10, leading=14,
                        alignment=TA_CENTER, textColor=GRAY
                    )))

            # Titre du poste ciblé — au-dessus du trait, plus grand que le nom
            if _poste_cv:
                story.append(Spacer(1, 0.15*cm))
                story.append(Paragraph(_poste_cv, ParagraphStyle(
                    'CV_Poste', fontName='Times-Bold', fontSize=20, leading=24,
                    alignment=TA_CENTER, textColor=DARK, spaceAfter=3
                )))
            story.append(Spacer(1, 0.2*cm))
            story.append(HRFlowable(width='100%', thickness=1.5, color=ACCENT, spaceAfter=8))

        elif is_lm:
            # ── LETTRE DE MOTIVATION — format français officiel ──────────────
            skip_normal_body = True
            from datetime import date as _date_cls

            # Parsing du contenu structuré généré par l'IA
            dest_line   = ''
            objet_line  = ''
            body_paras  = []
            header_done = False
            cur_para    = []

            for raw_ln in content_text.split('\n'):
                sl = raw_ln.strip()
                if not header_done:
                    if sl.startswith('DESTINATAIRE:'):
                        dest_line = sl[13:].strip()
                    elif sl.startswith('OBJET:'):
                        objet_line = sl[6:].strip()
                        header_done = True
                else:
                    if sl:
                        cur_para.append(sl)
                    elif cur_para:
                        body_paras.append(' '.join(cur_para))
                        cur_para = []
            if cur_para:
                body_paras.append(' '.join(cur_para))
            # Fallback si l'IA n'a pas suivi le format
            if not body_paras:
                body_paras = [p.strip() for p in content_text.split('\n\n') if p.strip()
                              and not p.strip().startswith(('DESTINATAIRE:', 'OBJET:'))]

            city_name = user_loc.split(',')[0].strip() if user_loc else 'Abidjan'
            today_str = _date_cls.today().strftime('%d/%m/%Y')

            # ─── Ligne 1 : Expéditeur (gauche, décalé) | Date (droite, en haut) ──
            snd_name_p = Paragraph(user_name, ParagraphStyle(
                'LM_SN', fontName='Times-Bold', fontSize=FS, leading=LD,
                textColor=DARK, spaceBefore=6))
            snd_items = [snd_name_p]
            for info_line in [user_loc, user_email, user_phone]:
                if info_line:
                    snd_items.append(Paragraph(info_line, ParagraphStyle(
                        'LM_SI', fontName='Times-Roman', fontSize=FS,
                        leading=LD, textColor=DARK, spaceAfter=0)))

            date_p = Paragraph(f'{city_name}, le {today_str}', ParagraphStyle(
                'LM_DT', fontName='Times-Roman', fontSize=FS, leading=LD,
                textColor=DARK, alignment=TA_RIGHT))

            lm_hdr_tbl = Table([[snd_items, [date_p]]], colWidths=[9.5*cm, 7.0*cm])
            lm_hdr_tbl.setStyle(TableStyle([
                ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
                ('TOPPADDING',    (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            story.append(lm_hdr_tbl)
            story.append(Spacer(1, 0.75*cm))

            # ─── Destinataire (aligné à droite) ──────────────────────────────
            LM_LD = 18   # interligne 1.5
            dest_right = ParagraphStyle(
                'LM_Dest', fontName='Times-Roman', fontSize=FS, leading=LM_LD,
                textColor=DARK, alignment=TA_RIGHT, spaceAfter=0)
            story.append(Paragraph('A', dest_right))
            if dest_line:
                if ' chez ' in dest_line:
                    role_p, co_p = dest_line.split(' chez ', 1)
                    story.append(Paragraph('Madame/Monsieur,', dest_right))
                    story.append(Paragraph(role_p, dest_right))
                    story.append(Paragraph(co_p, dest_right))
                else:
                    story.append(Paragraph('Madame/Monsieur,', dest_right))
                    story.append(Paragraph(dest_line, dest_right))
            else:
                story.append(Paragraph('Madame/Monsieur,', dest_right))
            story.append(Spacer(1, 0.55*cm))

            # ─── Objet ────────────────────────────────────────────────────────
            objet_txt = objet_line or 'Candidature'
            # Supprimer "chez …" en fin d'objet
            objet_txt = re_mod.sub(r'\s+chez\s+.+$', '', objet_txt, flags=re_mod.IGNORECASE).strip()
            # "Objet :" souligné gras + texte objet gras non souligné
            story.append(Paragraph(
                f'<u><b>Objet :</b></u><b> {objet_txt}</b>',
                ParagraphStyle('LM_Obj', fontName='Times-Bold', fontSize=FS,
                               leading=LM_LD, textColor=DARK, alignment=TA_LEFT,
                               spaceAfter=0)
            ))
            story.append(Spacer(1, 0.35*cm))

            # ─── Appel "Monsieur/Madame," centré ─────────────────────────────
            story.append(Paragraph('Monsieur/Madame,', ParagraphStyle(
                'LM_Call', fontName='Times-Roman', fontSize=FS, leading=LM_LD,
                textColor=DARK, alignment=TA_CENTER, spaceAfter=0)))
            story.append(Spacer(1, 0.3*cm))

            # ─── Corps — alinéa rentrant ──────────────────────────────────────
            indent_s = ParagraphStyle(
                'LM_Body', fontName='Times-Roman', fontSize=FS, leading=LM_LD,
                firstLineIndent=1.0*cm, alignment=TA_JUSTIFY,
                textColor=DARK, spaceAfter=8)
            for para_txt in body_paras:
                # Supprimer les tirets/puces résiduels générés par l'IA
                clean = re_mod.sub(r'^\s*[-•*]\s+', '', para_txt)
                clean = re_mod.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean)
                clean = clean.strip()
                if clean:
                    story.append(Paragraph(clean, indent_s))

            story.append(Spacer(1, 0.5*cm))

            # ─── Formule de politesse ─────────────────────────────────────────
            story.append(Paragraph(
                "Dans l'attente d'une réponse favorable, je vous prie d'agréer "
                "Madame/Monsieur, l'expression de ma considération distinguée.",
                ParagraphStyle('LM_Close', fontName='Times-Roman', fontSize=FS,
                               leading=LM_LD, firstLineIndent=1.0*cm,
                               alignment=TA_JUSTIFY, textColor=DARK)
            ))
            story.append(Spacer(1, 1.3*cm))

            # ─── Signature ────────────────────────────────────────────────────
            story.append(Paragraph(user_name, ParagraphStyle(
                'LM_Sig', fontName='Times-Roman', fontSize=FS, leading=LM_LD,
                textColor=DARK, alignment=TA_RIGHT)))

        elif is_email:
            # ── EMAIL PROFESSIONNEL — même mise en forme que LM ──────────────
            skip_normal_body = True
            from datetime import date as _date_cls

            # Parsing du contenu structuré : DESTINATAIRE: / OBJET: / corps
            dest_line   = ''
            objet_line  = ''
            body_paras  = []
            header_done = False
            cur_para    = []

            for raw_ln in content_text.split('\n'):
                sl = raw_ln.strip()
                if not header_done:
                    if sl.startswith('DESTINATAIRE:'):
                        dest_line = sl[13:].strip()
                    elif sl.startswith('OBJET:'):
                        objet_line = sl[6:].strip()
                        header_done = True
                else:
                    if sl:
                        cur_para.append(sl)
                    elif cur_para:
                        body_paras.append(' '.join(cur_para))
                        cur_para = []
            if cur_para:
                body_paras.append(' '.join(cur_para))
            if not body_paras:
                body_paras = [p.strip() for p in content_text.split('\n\n') if p.strip()
                              and not p.strip().startswith(('DESTINATAIRE:', 'OBJET:'))]

            city_name = user_loc.split(',')[0].strip() if user_loc else 'Abidjan'
            today_str = _date_cls.today().strftime('%d/%m/%Y')
            EM_LD = 18   # interligne 1.5

            # ─── Expéditeur (gauche) | Date (droite) ─────────────────────────
            snd_name_p = Paragraph(user_name, ParagraphStyle(
                'EM_SN', fontName='Times-Bold', fontSize=FS, leading=EM_LD,
                textColor=DARK, spaceBefore=6))
            snd_items = [snd_name_p]
            for info_line in [user_loc, user_email, user_phone]:
                if info_line:
                    snd_items.append(Paragraph(info_line, ParagraphStyle(
                        'EM_SI', fontName='Times-Roman', fontSize=FS,
                        leading=EM_LD, textColor=DARK, spaceAfter=0)))

            date_p = Paragraph(f'{city_name}, le {today_str}', ParagraphStyle(
                'EM_DT', fontName='Times-Roman', fontSize=FS, leading=EM_LD,
                textColor=DARK, alignment=TA_RIGHT))

            em_hdr_tbl = Table([[snd_items, [date_p]]], colWidths=[9.5*cm, 7.0*cm])
            em_hdr_tbl.setStyle(TableStyle([
                ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
                ('TOPPADDING',    (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            story.append(em_hdr_tbl)
            story.append(Spacer(1, 0.9*cm))

            # ─── Destinataire (aligné à droite) ──────────────────────────────
            dest_right = ParagraphStyle(
                'EM_Dest', fontName='Times-Roman', fontSize=FS, leading=EM_LD,
                textColor=DARK, alignment=TA_RIGHT, spaceAfter=0)
            if dest_line:
                if ' chez ' in dest_line:
                    role_p, co_p = dest_line.split(' chez ', 1)
                    story.append(Paragraph('A', dest_right))
                    story.append(Paragraph('Madame/Monsieur,', dest_right))
                    story.append(Paragraph(role_p, dest_right))
                    story.append(Paragraph(co_p, dest_right))
                else:
                    story.append(Paragraph('A', dest_right))
                    story.append(Paragraph(dest_line, dest_right))
            else:
                story.append(Paragraph('A', dest_right))
                story.append(Paragraph('Madame/Monsieur,', dest_right))
            story.append(Spacer(1, 0.7*cm))

            # ─── Objet ────────────────────────────────────────────────────────
            objet_txt = objet_line or 'Sujet'
            story.append(Paragraph(
                f'<u><b>Objet : {objet_txt}</b></u>',
                ParagraphStyle('EM_Obj', fontName='Times-Bold', fontSize=FS,
                               leading=EM_LD, textColor=DARK, alignment=TA_LEFT,
                               spaceAfter=0)
            ))
            story.append(Spacer(1, 0.5*cm))

            # ─── Corps — alinéa rentrant ──────────────────────────────────────
            salut_re = re_mod.compile(r'^(madame|monsieur|cher|chère|chère?s?)', re_mod.I)
            indent_s = ParagraphStyle(
                'EM_Body', fontName='Times-Roman', fontSize=FS, leading=EM_LD,
                firstLineIndent=1.0*cm, alignment=TA_JUSTIFY,
                textColor=DARK, spaceAfter=8)
            center_s = ParagraphStyle(
                'EM_Sal', fontName='Times-Roman', fontSize=FS, leading=EM_LD,
                alignment=TA_CENTER, textColor=DARK, spaceAfter=8)
            for para_txt in body_paras:
                clean = re_mod.sub(r'^\s*[-•*]\s+', '', para_txt).strip()
                clean = re_mod.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean)
                if clean:
                    style = center_s if salut_re.match(para_txt.strip()) else indent_s
                    story.append(Paragraph(clean, style))

            story.append(Spacer(1, 1.3*cm))

            # ─── Signature ────────────────────────────────────────────────────
            story.append(Paragraph(user_name, ParagraphStyle(
                'EM_Sig', fontName='Times-Roman', fontSize=FS, leading=EM_LD,
                textColor=DARK, alignment=TA_RIGHT)))

        else:
            # Document générique
            story.append(Paragraph(title, ParagraphStyle(
                'T', fontName='Times-Bold', fontSize=16, leading=20,
                textColor=DARK, spaceAfter=8, alignment=TA_CENTER
            )))
            story.append(HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8))

        # ── CONTENU (CV et documents génériques — pas lettre) ────────────────
        if not skip_normal_body:
            # Pour le CV : supprimer les premières lignes qui dupliquent l'en-tête
            # (nom, titre, localisation déjà affichés dans le bloc photo/identité)
            cv_header_skip = set()
            if is_cv:
                _name_norm = re_mod.sub(r'\s+', ' ', user_name).strip().lower()
                cv_header_skip.add(_name_norm)
                try:
                    _desired = (request.user.profile.desired_title or '').strip().lower()
                    if _desired:
                        cv_header_skip.add(_desired)
                except Exception:
                    pass
                if user_loc:
                    cv_header_skip.add(user_loc.strip().lower())
                # Sauter la ligne "Candidature : POSTE" (déjà rendue dans l'en-tête)
                if _poste_cv:
                    cv_header_skip.add(_poste_cv.strip().lower())
                    cv_header_skip.add(f'candidature : {_poste_cv.strip().lower()}')

            _cur_section  = ''   # section courante (pour numérotation expérience)
            _exp_counter  = 0   # compteur d'entrées expérience pro

            _EXP_KEYWORDS = ('expérience', 'experience', 'professionnel', 'emploi', 'poste')

            for raw_line in content_text.split('\n'):
                is_sub = bool(re_mod.match(r'^  +[-•*]\s', raw_line))
                line = raw_line.strip()
                if not line:
                    story.append(Spacer(1, 0.15*cm))
                    continue

                # Sauter les lignes d'en-tête dupliquées dans le CV
                if is_cv and cv_header_skip:
                    _clean = re_mod.sub(r'[*#_]', '', line).strip().lower()
                    if _clean in cv_header_skip:
                        continue
                    # Sauter toute ligne "Candidature : ..." (déjà dans l'en-tête)
                    if re_mod.match(r'^candidature\s*:', _clean):
                        continue

                # Section titre (** ... **)
                if re_mod.match(r'^\*\*[^*]+\*\*$', line):
                    text = line.strip('*').strip()
                    _cur_section = text.lower()
                    _exp_counter = 0
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph(text.upper(), h2_style))
                    story.append(HRFlowable(width='100%', thickness=0.5, color=LBORDER, spaceAfter=4))
                    continue

                # Titre Markdown ###
                if re_mod.match(r'^#{1,3}\s', line):
                    level = len(re_mod.match(r'^(#+)', line).group(1))
                    text  = re_mod.sub(r'^#+\s*', '', line)
                    _cur_section = text.lower()
                    _exp_counter = 0
                    s = h2_style if level <= 2 else h3_style
                    story.append(Paragraph(text, s))
                    continue

                # Puces
                if re_mod.match(r'^[-•*]\s', line):
                    text = re_mod.sub(r'^[-•*]\s+', '', line)
                    text = re_mod.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                    in_exp = any(k in _cur_section for k in _EXP_KEYWORDS)
                    if in_exp and not is_sub:
                        # Numérotation pour l'expérience professionnelle
                        _exp_counter += 1
                        num = f'<font color="#C9A84C"><b>{_exp_counter}.</b></font>'
                        story.append(Paragraph(num + '  ' + text, bullet_style))
                    elif is_sub:
                        marker = '<font color="#888888">\u2013</font>'
                        story.append(Paragraph(marker + '  ' + text, sub_bullet_style))
                    else:
                        marker = '<font color="#C9A84C" size="10"><b>\u25a0</b></font>'
                        story.append(Paragraph(marker + '  ' + text, bullet_style))
                    continue

                # Ligne normale
                text = re_mod.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
                story.append(Paragraph(text, body_style))


        # ── CONSTRUCTION ─────────────────────────────────────────────────────
        doc.build(story)
        buf.seek(0)

        resp = HttpResponse(buf.read(), content_type='application/pdf')
        safe   = re_mod.sub(r'[^a-zA-Z0-9_-]', '_', title)[:40]
        prefix = 'EMAIL' if is_email else ('LM' if is_lm else ('CV' if is_cv else 'JobFinder'))
        resp['Content-Disposition'] = f'attachment; filename="{prefix}_{safe}.pdf"'
        return resp

    except ImportError:
        resp = HttpResponse(content_text, content_type='text/plain; charset=utf-8')
        resp['Content-Disposition'] = 'attachment; filename="document.txt"'
        return resp
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# ═══════════════════════════════════════════════════════════════
#  GÉNÉRATION DOCX — Times New Roman 12pt, interligne 1.5
# ═══════════════════════════════════════════════════════════════

@login_required
@require_POST
def generate_docx(request):
    """
    Génère un DOCX professionnel : Times New Roman 12pt, interligne 1.5.
    Photo en haut à gauche pour CV et LM, mise en page soignée.
    """
    import io, re as re_mod
    from django.http import HttpResponse

    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({'error': 'JSON invalide'}, status=400)

    content_text  = data.get('content', '')
    title         = data.get('title', 'Document JobFinder AI')
    include_photo = data.get('include_photo', False)

    if not content_text:
        return JsonResponse({'error': 'Contenu vide'}, status=400)

    user        = request.user
    user_name   = user.get_full_name() or user.email
    user_email  = user.email
    user_phone  = ''
    user_loc    = ''
    user_linkedin = ''
    user_title  = ''
    photo_path  = None

    try:
        p = user.profile
        user_loc      = p.location or ''
        user_phone    = p.phone or ''
        user_linkedin = p.linkedin_url or ''
        user_title    = p.desired_title or ''
        if include_photo and p.avatar:
            photo_path = p.avatar.path
    except Exception:
        pass

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import lxml.etree as etree

        # ── Palette ─────────────────────────────────────────────────────────
        DARK   = RGBColor(0x1A, 0x1A, 0x2E)
        ACCENT = RGBColor(0xC9, 0xA8, 0x4C)  # or premium
        GRAY   = RGBColor(0x55, 0x55, 0x55)
        LIGHT  = RGBColor(0x88, 0x88, 0x88)
        WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

        is_cv    = 'cv' in title.lower() or 'curriculum' in title.lower()
        is_lm    = 'lettre' in title.lower() or 'motivation' in title.lower()
        is_email = title.lower().startswith('email_') or title.lower().startswith('email ')

        doc = Document()

        # ── Marges ──────────────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.0)

        # ── Fonction: appliquer TNR 12pt interligne 1.5 ─────────────────────
        def set_run_tnr(run, size=12, bold=False, italic=False, color=None):
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold   = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color
            # Force font East-Asia aussi
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'),    'Times New Roman')
            rFonts.set(qn('w:hAnsi'),    'Times New Roman')
            rFonts.set(qn('w:cs'),       'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            existing = rPr.find(qn('w:rFonts'))
            if existing is not None:
                rPr.remove(existing)
            rPr.insert(0, rFonts)

        def set_para_spacing(para, size=12, bold=False, italic=False,
                             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=0, space_after=6):
            para.alignment = align
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing      = Pt(size * 1.5)
            pf.space_before      = Pt(space_before)
            pf.space_after       = Pt(space_after)

        def add_tnr_para(text, bold=False, italic=False, size=12, color=None,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=6):
            para = doc.add_paragraph()
            set_para_spacing(para, size=size, space_before=space_before, space_after=space_after, align=align)
            run = para.add_run(text)
            set_run_tnr(run, size=size, bold=bold, italic=italic, color=color)
            return para

        def add_horizontal_line(color_hex='2C5F8A', thickness=12):
            """Ajoute une ligne horizontale via border de paragraphe."""
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),   'single')
            bot.set(qn('w:sz'),    str(thickness))
            bot.set(qn('w:space'), '1')
            bot.set(qn('w:color'), color_hex)
            pBdr.append(bot)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(4)
            return p

        # ── EN-TÊTE ──────────────────────────────────────────────────────────
        skip_normal_body = False

        if is_cv:
            if photo_path and include_photo:
                try:
                    # Tableau 2 colonnes: photo | infos identité
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                    tbl.style = 'Table Grid'
                    # Remove all borders
                    tbl_el2 = tbl._tbl
                    tbl_pr = tbl_el2.find(qn('w:tblPr'))
                    if tbl_pr is None:
                        tbl_pr = OxmlElement('w:tblPr')
                        tbl_el2.insert(0, tbl_pr)
                    tbl_bdr = OxmlElement('w:tblBorders')
                    for side in ('top','left','bottom','right','insideH','insideV'):
                        el = OxmlElement(f'w:{side}')
                        el.set(qn('w:val'), 'none')
                        tbl_bdr.append(el)
                    tbl_pr.append(tbl_bdr)

                    # Colonne identité (gauche) — ATS-friendly : texte en premier
                    cell_id = tbl.rows[0].cells[0]
                    cell_id.width = Cm(13.0)
                    cell_id.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # Nom
                    p_name = cell_id.paragraphs[0]
                    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_name = p_name.add_run(user_name)
                    set_run_tnr(run_name, size=18, bold=True, color=DARK)

                    # Titre souhaité
                    if user_title:
                        p_t = cell_id.add_paragraph()
                        p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        r_t = p_t.add_run(user_title)
                        set_run_tnr(r_t, size=12, italic=True, color=ACCENT)

                    # Coordonnées
                    coords = [x for x in [user_phone, user_email, user_loc, user_linkedin] if x]
                    if coords:
                        p_c = cell_id.add_paragraph()
                        p_c.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        r_c = p_c.add_run(' | '.join(coords))
                        set_run_tnr(r_c, size=10, color=GRAY)

                    # Colonne photo (droite)
                    cell_photo = tbl.rows[0].cells[1]
                    cell_photo.width = Cm(4)
                    cell_photo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    ph_para = cell_photo.paragraphs[0]
                    ph_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = ph_para.add_run()
                    run_img.add_picture(photo_path, width=Cm(3.0), height=Cm(3.5))

                except Exception:
                    # Fallback sans photo
                    add_tnr_para(user_name, bold=True, size=18, color=DARK,
                                 align=WD_ALIGN_PARAGRAPH.LEFT)
                    if user_title:
                        add_tnr_para(user_title, italic=True, size=12, color=ACCENT,
                                     align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                # Sans photo — centré
                add_tnr_para(user_name, bold=True, size=18, color=DARK,
                             align=WD_ALIGN_PARAGRAPH.CENTER)
                if user_title:
                    add_tnr_para(user_title, italic=True, size=12, color=ACCENT,
                                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
                coords = [x for x in [user_phone, user_email, user_loc, user_linkedin] if x]
                if coords:
                    add_tnr_para(' | '.join(coords), size=10, color=GRAY,
                                 align=WD_ALIGN_PARAGRAPH.CENTER)

            doc.add_paragraph()
            add_horizontal_line('C9A84C', 12)

        elif is_lm:
            # ── LETTRE DE MOTIVATION — format français officiel ──────────────
            skip_normal_body = True
            from datetime import date as _date_cls

            # Parsing du contenu structuré
            dest_line   = ''
            objet_line  = ''
            body_paras  = []
            header_done = False
            cur_para    = []

            for raw_ln in content_text.split('\n'):
                sl = raw_ln.strip()
                if not header_done:
                    if sl.startswith('DESTINATAIRE:'):
                        dest_line = sl[13:].strip()
                    elif sl.startswith('OBJET:'):
                        objet_line = sl[6:].strip()
                        header_done = True
                else:
                    if sl:
                        cur_para.append(sl)
                    elif cur_para:
                        body_paras.append(' '.join(cur_para))
                        cur_para = []
            if cur_para:
                body_paras.append(' '.join(cur_para))
            if not body_paras:
                body_paras = [p.strip() for p in content_text.split('\n\n') if p.strip()
                              and not p.strip().startswith(('DESTINATAIRE:', 'OBJET:'))]

            city_name = user_loc.split(',')[0].strip() if user_loc else 'Abidjan'
            today_str = _date_cls.today().strftime('%d/%m/%Y')

            # ─── Expéditeur (gauche) | Date (droite) ─────────────────────────
            lm_tbl = doc.add_table(rows=1, cols=2)
            lm_tbl.style = 'Table Grid'
            # Suppression des bordures (compatible toutes versions python-docx)
            lm_tbl_el = lm_tbl._tbl
            lm_tblPr = lm_tbl_el.find(qn('w:tblPr'))
            if lm_tblPr is None:
                lm_tblPr = OxmlElement('w:tblPr')
                lm_tbl_el.insert(0, lm_tblPr)
            lm_bdr = OxmlElement('w:tblBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                lm_bdr.append(el)
            lm_tblPr.append(lm_bdr)

            # Cellule gauche : expéditeur
            cell_snd = lm_tbl.rows[0].cells[0]
            cell_snd.width = Cm(9.5)
            cell_snd.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p_snd = cell_snd.paragraphs[0]
            p_snd.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_snd.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_snd.paragraph_format.line_spacing = Pt(18)
            p_snd.paragraph_format.space_before = Pt(14)  # décale le nom sous la date
            p_snd.paragraph_format.space_after  = Pt(0)
            r_snd = p_snd.add_run(user_name)
            set_run_tnr(r_snd, size=12, bold=True, color=DARK)
            for info in [user_loc, user_email, user_phone]:
                if info:
                    p_i = cell_snd.add_paragraph()
                    p_i.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_i.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p_i.paragraph_format.line_spacing = Pt(18)
                    p_i.paragraph_format.space_before = Pt(0)
                    p_i.paragraph_format.space_after  = Pt(0)
                    r_i = p_i.add_run(info)
                    set_run_tnr(r_i, size=12, color=DARK)

            # Cellule droite : date
            cell_dt = lm_tbl.rows[0].cells[1]
            cell_dt.width = Cm(7.0)
            cell_dt.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p_dt = cell_dt.paragraphs[0]
            p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_dt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_dt.paragraph_format.line_spacing = Pt(18)
            p_dt.paragraph_format.space_after  = Pt(0)
            r_dt = p_dt.add_run(f'{city_name}, le {today_str}')
            set_run_tnr(r_dt, size=12, color=DARK)

            doc.add_paragraph()

            # ─── Destinataire (aligné à droite) ──────────────────────────────
            add_tnr_para('A', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
            if dest_line:
                if ' chez ' in dest_line:
                    role_p, co_p = dest_line.split(' chez ', 1)
                    add_tnr_para('Madame/Monsieur,', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                    add_tnr_para(role_p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                    add_tnr_para(co_p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                else:
                    add_tnr_para('Madame/Monsieur,', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                    add_tnr_para(dest_line, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
            else:
                add_tnr_para('Madame/Monsieur,', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

            doc.add_paragraph()

            # ─── Objet ────────────────────────────────────────────────────────
            objet_txt = objet_line or 'Candidature'
            p_obj = doc.add_paragraph()
            p_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_obj.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_obj.paragraph_format.line_spacing = Pt(18)
            p_obj.paragraph_format.space_after  = Pt(10)
            r_obj = p_obj.add_run(f'Objet : {objet_txt}')
            set_run_tnr(r_obj, size=12, bold=True, color=DARK)
            r_obj.font.underline = True

            # ─── Corps — alinéa rentrant ──────────────────────────────────────
            for para_txt in body_paras:
                p_body = doc.add_paragraph()
                p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p_body.paragraph_format.line_spacing      = Pt(18)
                p_body.paragraph_format.first_line_indent = Cm(1.0)
                p_body.paragraph_format.space_after       = Pt(8)
                parts = re_mod.split(r'(\*\*[^*]+\*\*)', para_txt)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p_body.add_run(part[2:-2])
                        set_run_tnr(r, size=12, bold=True, color=DARK)
                    else:
                        r = p_body.add_run(part)
                        set_run_tnr(r, size=12, color=DARK)

            # ─── Formule de politesse ─────────────────────────────────────────
            p_close = doc.add_paragraph()
            p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_close.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_close.paragraph_format.line_spacing      = Pt(18)
            p_close.paragraph_format.first_line_indent = Cm(1.0)
            p_close.paragraph_format.space_after       = Pt(36)
            r_close = p_close.add_run(
                "Dans l'attente d'une réponse favorable, je vous prie d'agréer "
                "Madame/Monsieur, l'expression de ma considération distinguée."
            )
            set_run_tnr(r_close, size=12, color=DARK)

            # ─── Signature ────────────────────────────────────────────────────
            add_tnr_para(user_name, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

        elif is_email:
            # ── EMAIL PROFESSIONNEL — même mise en forme que LM ──────────────
            skip_normal_body = True
            from datetime import date as _date_cls

            # Parsing du contenu structuré : DESTINATAIRE: / OBJET: / corps
            dest_line   = ''
            objet_line  = ''
            body_paras  = []
            header_done = False
            cur_para    = []

            for raw_ln in content_text.split('\n'):
                sl = raw_ln.strip()
                if not header_done:
                    if sl.startswith('DESTINATAIRE:'):
                        dest_line = sl[13:].strip()
                    elif sl.startswith('OBJET:'):
                        objet_line = sl[6:].strip()
                        header_done = True
                else:
                    if sl:
                        cur_para.append(sl)
                    elif cur_para:
                        body_paras.append(' '.join(cur_para))
                        cur_para = []
            if cur_para:
                body_paras.append(' '.join(cur_para))
            if not body_paras:
                body_paras = [p.strip() for p in content_text.split('\n\n') if p.strip()
                              and not p.strip().startswith(('DESTINATAIRE:', 'OBJET:'))]

            city_name = user_loc.split(',')[0].strip() if user_loc else 'Abidjan'
            today_str = _date_cls.today().strftime('%d/%m/%Y')

            # ─── Expéditeur (gauche) | Date (droite) ─────────────────────────
            em_tbl = doc.add_table(rows=1, cols=2)
            em_tbl.style = 'Table Grid'
            em_tbl_el = em_tbl._tbl
            em_tblPr = em_tbl_el.find(qn('w:tblPr'))
            if em_tblPr is None:
                em_tblPr = OxmlElement('w:tblPr')
                em_tbl_el.insert(0, em_tblPr)
            em_bdr = OxmlElement('w:tblBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                em_bdr.append(el)
            em_tblPr.append(em_bdr)

            # Cellule gauche : expéditeur
            cell_snd = em_tbl.rows[0].cells[0]
            cell_snd.width = Cm(9.5)
            cell_snd.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p_snd = cell_snd.paragraphs[0]
            p_snd.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_snd.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_snd.paragraph_format.line_spacing = Pt(18)
            p_snd.paragraph_format.space_before = Pt(14)
            p_snd.paragraph_format.space_after  = Pt(0)
            r_snd = p_snd.add_run(user_name)
            set_run_tnr(r_snd, size=12, bold=True, color=DARK)
            for info in [user_loc, user_email, user_phone]:
                if info:
                    p_i = cell_snd.add_paragraph()
                    p_i.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_i.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p_i.paragraph_format.line_spacing = Pt(18)
                    p_i.paragraph_format.space_before = Pt(0)
                    p_i.paragraph_format.space_after  = Pt(0)
                    r_i = p_i.add_run(info)
                    set_run_tnr(r_i, size=12, color=DARK)

            # Cellule droite : date
            cell_dt = em_tbl.rows[0].cells[1]
            cell_dt.width = Cm(7.0)
            cell_dt.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p_dt = cell_dt.paragraphs[0]
            p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_dt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_dt.paragraph_format.line_spacing = Pt(18)
            p_dt.paragraph_format.space_after  = Pt(0)
            r_dt = p_dt.add_run(f'{city_name}, le {today_str}')
            set_run_tnr(r_dt, size=12, color=DARK)

            doc.add_paragraph()

            # ─── Destinataire (aligné à droite) ──────────────────────────────
            if dest_line:
                if ' chez ' in dest_line:
                    role_p, co_p = dest_line.split(' chez ', 1)
                    add_tnr_para('A', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                    add_tnr_para('Madame/Monsieur,', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                    add_tnr_para(role_p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                    add_tnr_para(co_p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                else:
                    add_tnr_para('A', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                    add_tnr_para(dest_line, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
            else:
                add_tnr_para('A', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
                add_tnr_para('Madame/Monsieur,', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

            doc.add_paragraph()

            # ─── Objet ────────────────────────────────────────────────────────
            objet_txt = objet_line or 'Sujet'
            p_obj = doc.add_paragraph()
            p_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_obj.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_obj.paragraph_format.line_spacing = Pt(18)
            p_obj.paragraph_format.space_after  = Pt(10)
            r_obj = p_obj.add_run(f'Objet : {objet_txt}')
            set_run_tnr(r_obj, size=12, bold=True, color=DARK)
            r_obj.font.underline = True

            # ─── Corps — alinéa rentrant (salutation centrée) ────────────────
            salut_re_docx = re_mod.compile(r'^(madame|monsieur|cher|chère?s?)', re_mod.I)
            for para_txt in body_paras:
                p_body = doc.add_paragraph()
                is_salut = bool(salut_re_docx.match(para_txt.strip()))
                p_body.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_salut else WD_ALIGN_PARAGRAPH.JUSTIFY
                p_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p_body.paragraph_format.line_spacing      = Pt(18)
                p_body.paragraph_format.first_line_indent = Cm(0) if is_salut else Cm(1.0)
                p_body.paragraph_format.space_after       = Pt(8)
                parts = re_mod.split(r'(\*\*[^*]+\*\*)', para_txt)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p_body.add_run(part[2:-2])
                        set_run_tnr(r, size=12, bold=True, color=DARK)
                    else:
                        r = p_body.add_run(part)
                        set_run_tnr(r, size=12, color=DARK)

            # ─── Signature ────────────────────────────────────────────────────
            doc.add_paragraph()
            add_tnr_para(user_name, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

        else:
            # Document générique
            add_tnr_para(title, bold=True, size=16, color=DARK,
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
            add_horizontal_line('C9A84C', 8)

        # ── CONTENU (CV et documents génériques — pas lettre) ────────────────
        if not skip_normal_body:
            # Pour le CV : supprimer les premières lignes qui dupliquent l'en-tête
            cv_header_skip = set()
            if is_cv:
                _name_norm = re_mod.sub(r'\s+', ' ', user_name).strip().lower()
                cv_header_skip.add(_name_norm)
                try:
                    _desired = (request.user.profile.desired_title or '').strip().lower()
                    if _desired:
                        cv_header_skip.add(_desired)
                except Exception:
                    pass
                if user_loc:
                    cv_header_skip.add(user_loc.strip().lower())

            _cur_section  = ''
            _exp_counter  = 0
            _EXP_KEYWORDS = ('expérience', 'experience', 'professionnel', 'emploi', 'poste')

            for raw_line in content_text.split('\n'):
                is_sub = bool(re_mod.match(r'^  +[-•*]\s', raw_line))
                line = raw_line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                # Sauter les lignes d'en-tête dupliquées dans le CV
                if is_cv and cv_header_skip:
                    _clean = re_mod.sub(r'[*#_]', '', line).strip().lower()
                    if _clean in cv_header_skip:
                        continue
                    # Sauter toute ligne "Candidature : ..." (déjà dans l'en-tête)
                    if re_mod.match(r'^candidature\s*:', _clean):
                        continue

                # Section titre **...**
                if re_mod.match(r'^\*\*[^*]+\*\*$', line):
                    text = line.strip('*').strip()
                    _cur_section = text.lower()
                    _exp_counter = 0
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after  = Pt(2)
                    r = p.add_run(text.upper())
                    set_run_tnr(r, size=12, bold=True, color=ACCENT)
                    add_horizontal_line('CCCCCC', 6)
                    continue

                # Titre Markdown ###
                if re_mod.match(r'^#{1,3}\s', line):
                    level = len(re_mod.match(r'^(#+)', line).group(1))
                    text  = re_mod.sub(r'^#+\s*', '', line)
                    _cur_section = text.lower()
                    _exp_counter = 0
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after  = Pt(2)
                    r = p.add_run(text)
                    sz = 13 if level == 1 else 12
                    set_run_tnr(r, size=sz, bold=True, color=ACCENT if level <= 2 else DARK)
                    continue

                # Puces
                if re_mod.match(r'^[-•*]\s', line):
                    text = re_mod.sub(r'^[-•*]\s+', '', line)
                    in_exp = any(k in _cur_section for k in _EXP_KEYWORDS)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p.paragraph_format.line_spacing  = Pt(18)
                    p.paragraph_format.space_after   = Pt(3)
                    if in_exp and not is_sub:
                        # Numérotation expérience pro
                        _exp_counter += 1
                        p.paragraph_format.left_indent       = Cm(0.7)
                        p.paragraph_format.first_line_indent = Cm(-0.5)
                        num_run = p.add_run(f'{_exp_counter}.  ')
                        set_run_tnr(num_run, size=12, bold=True, color=ACCENT)
                    elif is_sub:
                        p.paragraph_format.left_indent       = Cm(1.2)
                        p.paragraph_format.first_line_indent = Cm(-0.5)
                        bullet_run = p.add_run('\u2013  ')
                        set_run_tnr(bullet_run, size=12, color=RGBColor(0x88, 0x88, 0x88))
                    else:
                        p.paragraph_format.left_indent       = Cm(0.7)
                        p.paragraph_format.first_line_indent = Cm(-0.5)
                        bullet_run = p.add_run('\u25a0  ')
                        set_run_tnr(bullet_run, size=10, bold=True, color=ACCENT)
                    # Texte avec gras inline
                    parts = re_mod.split(r'(\*\*[^*]+\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            r = p.add_run(part[2:-2])
                            set_run_tnr(r, size=12, bold=True, color=DARK)
                        else:
                            r = p.add_run(part)
                            set_run_tnr(r, size=12, color=DARK)
                    continue

                # Paragraphe normal avec gras inline
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing  = Pt(18)  # 12 * 1.5
                p.paragraph_format.space_after   = Pt(4)
                parts = re_mod.split(r'(\*\*[^*]+\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        set_run_tnr(r, size=12, bold=True, color=DARK)
                    else:
                        r = p.add_run(part)
                        set_run_tnr(r, size=12, color=DARK)

            # ── PIED DE PAGE ──────────────────────────────────────────────────
            doc.add_paragraph()
            add_horizontal_line('CCCCCC', 4)
            p_f = doc.add_paragraph()
            p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_f = p_f.add_run('Généré par JobFinder AI — jobfinder.ci')
            set_run_tnr(r_f, size=9, italic=True, color=LIGHT)

        # ── EXPORT ───────────────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        resp = HttpResponse(
            buf.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        safe   = re_mod.sub(r'[^a-zA-Z0-9_-]', '_', title)[:40]
        prefix = 'EMAIL' if is_email else ('LM' if is_lm else ('CV' if is_cv else 'JobFinder'))
        resp['Content-Disposition'] = f'attachment; filename="{prefix}_{safe}.docx"'
        return resp

    except ImportError:
        return JsonResponse({'error': 'python-docx manquant. pip install python-docx'}, status=500)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def _add_formatted_run(paragraph, text, color):
    """Gère le markdown **gras** dans les runs DOCX avec Times New Roman."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        is_bold = part.startswith('**') and part.endswith('**')
        run = paragraph.add_run(part[2:-2] if is_bold else part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold   = is_bold
        run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)
