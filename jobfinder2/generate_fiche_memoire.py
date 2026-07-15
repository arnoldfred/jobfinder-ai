"""
Génère la fiche de projet de mémoire M2 MIAGE — JobFinder AI
en format DOCX et PDF sur le Bureau.
"""
import os, sys

DESKTOP = os.path.join(os.path.expanduser('~'), 'Desktop')
OUT_DOCX = os.path.join(DESKTOP, 'Fiche_Projet_Memoire_JobFinder_AI.docx')
OUT_PDF  = os.path.join(DESKTOP, 'Fiche_Projet_Memoire_JobFinder_AI.pdf')

# ═══════════════════════════════════════════════════════════════
#  CONTENU STRUCTURÉ
# ═══════════════════════════════════════════════════════════════

TITRE     = "FICHE DE PROJET DE MÉMOIRE – ARNOLD FREDDY ADOPO\nM2 MIAGE"
THEME     = ("Proposition de thème : JobFinder AI — Plateforme web intelligente de recherche "
             "d'emploi automatisée et de génération de candidatures personnalisées par "
             "intelligence artificielle générative")

SECTIONS = [
    {
        "num": "1", "title": "Contexte du projet",
        "body": [
            ("para", (
                "Le marché de l'emploi en Côte d'Ivoire fait face à plusieurs défis structurels : "
                "un chômage élevé chez les jeunes diplômés, une dispersion des offres sur de multiples "
                "plateformes, et l'absence d'outils intelligents d'accompagnement adaptés au contexte "
                "local. Les candidats consacrent un temps considérable à consulter les offres, adapter "
                "leurs CV et rédiger des lettres de motivation, souvent sans résultat concret."
            )),
            ("para", (
                "De leur côté, les recruteurs reçoivent des candidatures peu ciblées, ce qui ralentit "
                "leurs processus de sélection. C'est pour répondre à ces défis qu'a été conçue "
                "JobFinder AI, une plateforme web intelligente qui automatise la recherche d'emploi "
                "et la génération de candidatures personnalisées grâce à l'intelligence artificielle "
                "générative (IA générative) et au traitement du langage naturel (NLP)."
            )),
        ]
    },
    {
        "num": "2", "title": "Problématique",
        "body": [
            ("para", (
                "Comment concevoir une application web intelligente capable d'automatiser la collecte "
                "d'offres d'emploi, d'apprendre des comportements des utilisateurs et de générer des "
                "candidatures personnalisées et optimisées grâce aux grands modèles de langage (LLM), "
                "tout en restant adaptée au marché de l'emploi ivoirien ?"
            )),
        ]
    },
    {
        "num": "3", "title": "Objectifs du projet",
        "body": [
            ("heading3", "Objectif général :"),
            ("para", (
                "Développer une plateforme web intelligente qui facilite la recherche d'emploi en "
                "Côte d'Ivoire par l'automatisation du scraping, le matching comportemental et la "
                "génération de candidatures personnalisées par IA générative."
            )),
            ("heading3", "Objectifs spécifiques :"),
            ("list", [
                "Concevoir un moteur de scraping automatique collectant les offres d'emploi depuis "
                "plusieurs sources (LinkedIn, AEJI, etc.) toutes les 6 heures via APScheduler.",
                "Développer un module de matching intelligent basé sur les signaux comportementaux "
                "des utilisateurs (consultations, sauvegardes, candidatures, intérêts, rejets).",
                "Intégrer un système de génération automatique de CV et de lettres de motivation "
                "personnalisées via LLM (LLaMA 3.3 70B — Groq API), optimisés pour les filtres ATS.",
                "Implémenter un mécanisme anti-détection IA dans les documents générés, afin de "
                "produire des candidatures à l'aspect authentiquement humain.",
                "Proposer plusieurs styles de lettres de motivation (Classique, Impact & Résultats, "
                "Storytelling, Enthousiaste) adaptés aux préférences du candidat.",
                "Offrir une interface web ergonomique et intuitive.",
                "Évaluer la pertinence du système et la satisfaction des utilisateurs.",
            ]),
        ]
    },
    {
        "num": "4", "title": "Hypothèses de travail",
        "body": [
            ("list", [
                "Les grands modèles de langage (LLM) peuvent générer des candidatures professionnelles "
                "de qualité humaine, adaptées à chaque offre.",
                "Un système de scoring comportemental améliore la pertinence des recommandations au "
                "fil des interactions.",
                "L'optimisation ATS dans les documents générés améliore le taux de passage des "
                "filtres automatiques de recrutement.",
                "Une interface conviviale favorise l'adoption du système par les candidats ivoiriens.",
            ]),
        ]
    },
    {
        "num": "5", "title": "Fonctionnalités de l'application",
        "body": [
            ("heading3", "Côté Candidat"),
            ("numlist", [
                "Inscription / Authentification : création de compte et gestion sécurisée du profil.",
                "Création de profil professionnel : informations personnelles, expériences, formations, "
                "compétences (hard skills et soft skills séparés), photo de profil.",
                "Import automatique de CV : extraction des compétences à partir d'un CV existant (PDF ou DOCX).",
                "Scraping automatique d'offres : collecte toutes les 6 heures depuis LinkedIn, AEJI "
                "et autres sources via APScheduler.",
                "Matching IA comportemental : score de recommandation calculé selon les interactions "
                "passées (candidature : +3, sauvegarde : +2, intérêt : +2, consultation : +1, rejet : −2).",
                "Génération automatique de CV adapté : CV personnalisé avec séparation hard skills / "
                "soft skills et titre du poste ciblé en en-tête.",
                "Génération de lettres de motivation multi-styles : 4 styles distincts (Classique, "
                "Impact & Résultats, Storytelling, Enthousiaste), optimisés ATS, anti-détection IA.",
                "Génération d'emails professionnels de candidature : 4 types (candidature spontanée, "
                "réponse à offre, relance, remerciement post-entretien).",
                "Export PDF et DOCX des documents générés (CV, lettre, email).",
                "Historique des candidatures : suivi des emplois postulés et de leurs statuts.",
                "Alertes personnalisées : notifications d'offres correspondant au profil.",
                "Analyse IA du profil et estimation des chances de succès par offre.",
            ]),
            ("heading3", "Côté Système (IA et automatisation)"),
            ("numlist", [
                "Scraping / Agrégation automatique : collecte périodique (APScheduler, toutes les 6h) "
                "depuis LinkedIn, AEJI et autres plateformes d'emploi.",
                "Traitement NLP des offres : extraction de mots-clés, compétences requises, niveau "
                "d'expérience, type de contrat.",
                "Système de scoring comportemental (ML) : modèle de recommandation basé sur les "
                "interactions utilisateur-offre (modèle JobInteraction), avec pondération adaptative.",
                "Génération de documents par LLM : appels à l'API Groq (LLaMA 3.3 70B / LLaMA 3.1 8B) "
                "avec prompts structurés, retry automatique et chaîne de fallback sur 5 modèles.",
                "Optimisation ATS : les documents générés intègrent les mots-clés exacts du poste, "
                "des verbes d'action ciblés et la terminologie du secteur.",
                "Anti-détection IA : variation de la longueur des phrases, interdiction des formules "
                "IA-typiques, injection de faits personnels non reproductibles.",
                "Génération PDF professionnelle : mise en page Times New Roman 12pt, interligne 1.5, "
                "photo intégrée, en-tête structuré (ReportLab).",
                "Tableau de bord analytique : statistiques sur les candidatures, taux de matching, "
                "historique des interactions.",
            ]),
            ("heading3", "Côté Administrateur"),
            ("numlist", [
                "Gestion des utilisateurs (ajout, suppression, mise à jour).",
                "Supervision et validation des offres collectées.",
                "Statistiques générales (utilisateurs, offres, candidatures).",
                "Paramétrage du moteur de scraping et des sources.",
            ]),
        ]
    },
    {
        "num": "6", "title": "Architecture technique",
        "body": [
            ("para", (
                "L'application suit une architecture MVT (Model-View-Template) implémentée via le "
                "framework Django 5.x. Les principaux modèles de données sont : User, Profile, Job, "
                "Application, JobInteraction, SavedJob et SearchAlert. Le module IA est découplé dans "
                "un package dédié (ai_tools) comprenant deux composants clés :"
            )),
            ("list", [
                "groq_client.py : client API Groq avec retry automatique (503/429/timeout) et "
                "fallback sur une chaîne de 5 modèles LLaMA.",
                "views.py : logique complète de génération (CV adapté, lettre multi-styles, email "
                "professionnel, analyse de profil, scoring d'offre).",
            ]),
            ("para", (
                "La génération de documents PDF utilise ReportLab avec une mise en page professionnelle "
                "(Times New Roman, police 12pt, interligne 1.5, marges calibrées pour tenir sur 1 page). "
                "Le format DOCX est généré via python-docx avec la même structure."
            )),
        ]
    },
    {
        "num": "7", "title": "Méthodologie",
        "body": [
            ("para", "Approche Agile (itérative et incrémentale) :"),
            ("numlist", [
                "Analyse des besoins et rédaction du cahier des charges.",
                "Conception UML (diagrammes de cas d'utilisation, de classes, de séquences).",
                "Développement itératif du backend Django et du frontend HTML/JS/Bootstrap 5.",
                "Intégration IA : connexion API Groq, ingénierie des prompts, optimisation ATS "
                "et système anti-détection IA.",
                "Développement du module de scoring comportemental (JobInteraction ML).",
                "Tests unitaires, tests d'intégration et tests utilisateurs.",
                "Évaluation et optimisation après retours utilisateurs.",
            ]),
        ]
    },
    {
        "num": "8", "title": "Outils et technologies",
        "body": [
            ("table", [
                ["Domaine", "Technologies"],
                ["Langage", "Python 3.12"],
                ["Framework backend", "Django 5.x"],
                ["Frontend", "HTML5, CSS3, JavaScript, Bootstrap 5"],
                ["Base de données", "MySQL (production) / SQLite (développement)"],
                ["Scraping", "BeautifulSoup4, Selenium, Requests"],
                ["IA Générative (LLM)", "Groq API — LLaMA 3.3 70B / LLaMA 3.1 8B (gratuit)"],
                ["NLP / ML", "Scoring comportemental custom (JobInteraction), extraction de mots-clés"],
                ["Génération de documents", "ReportLab (PDF), python-docx (DOCX)"],
                ["Planification automatique", "APScheduler (scraping toutes les 6h)"],
                ["Environnement", "VS Code, WAMP Server, Git"],
            ]),
        ]
    },
    {
        "num": "9", "title": "Résultats attendus",
        "body": [
            ("list", [
                "Réduction significative du temps de recherche et de candidature.",
                "Meilleur taux d'adéquation entre profils candidats et offres d'emploi.",
                "Documents de candidature (CV, lettres) de qualité professionnelle, optimisés ATS "
                "et indétectables comme générés par IA.",
                "Système de recommandation qui s'améliore avec l'usage grâce au scoring comportemental.",
                "Outil intelligent localement adapté au marché ivoirien.",
            ]),
        ]
    },
    {
        "num": "10", "title": "Conclusion",
        "body": [
            ("para", (
                "JobFinder AI vise à transformer la recherche d'emploi en Côte d'Ivoire en la rendant "
                "plus rapide, intelligente et personnalisée. En combinant l'intelligence artificielle "
                "générative (LLaMA 3.3 70B via Groq API), un scoring comportemental adaptatif et la "
                "génération de documents optimisés ATS avec neutralisation de la détection IA, "
                "l'application offre à chaque candidat un avantage concret dans un marché compétitif."
            )),
            ("para", (
                "Ce projet associe innovation technologique, impact social et applicabilité réelle, "
                "en phase avec les enjeux actuels du marché du travail ivoirien et les avancées "
                "récentes de l'intelligence artificielle générative."
            )),
        ]
    },
]


# ═══════════════════════════════════════════════════════════════
#  GÉNÉRATION DOCX
# ═══════════════════════════════════════════════════════════════

def build_docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    def set_font(run, bold=False, size=12, color=None):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(text)
        set_font(run, bold=bold, size=size, color=color)
        return p

    # Titre principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(TITRE)
    set_font(run, bold=True, size=14)

    # Séparateur
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Thème
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = Pt(18)
    run_label = p.add_run("Proposition de thème : ")
    set_font(run_label, bold=True, size=12)
    run_text = p.add_run(
        "JobFinder AI — Plateforme web intelligente de recherche d'emploi automatisée "
        "et de génération de candidatures personnalisées par intelligence artificielle générative"
    )
    set_font(run_text, bold=False, size=12)

    for sec in SECTIONS:
        # Titre de section
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(f"{sec['num']}. {sec['title']}")
        set_font(run, bold=True, size=12)
        run.font.underline = True

        for item_type, item_content in sec['body']:

            if item_type == 'para':
                add_para(item_content)

            elif item_type == 'heading3':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.line_spacing = Pt(18)
                run = p.add_run(item_content)
                set_font(run, bold=True, size=12)

            elif item_type == 'list':
                for bullet in item_content:
                    p = doc.add_paragraph(style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after  = Pt(3)
                    p.paragraph_format.line_spacing = Pt(18)
                    run = p.add_run(bullet)
                    set_font(run, size=12)

            elif item_type == 'numlist':
                for i, item in enumerate(item_content, 1):
                    p = doc.add_paragraph(style='List Number')
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after  = Pt(3)
                    p.paragraph_format.line_spacing = Pt(18)
                    run = p.add_run(item)
                    set_font(run, size=12)

            elif item_type == 'table':
                rows = item_content
                tbl = doc.add_table(rows=len(rows), cols=2)
                tbl.style = 'Table Grid'
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(cell_text)
                        bold = (r_idx == 0)
                        set_font(run, bold=bold, size=11)
                        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                        cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
                doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.save(OUT_DOCX)
    print(f"DOCX généré : {OUT_DOCX}")


# ═══════════════════════════════════════════════════════════════
#  GÉNÉRATION PDF
# ═══════════════════════════════════════════════════════════════

def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, Table, TableStyle)
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

    DARK   = colors.HexColor('#1A1A2E')
    GRAY   = colors.HexColor('#555555')
    ACCENT = colors.HexColor('#C9A84C')
    LBORDER= colors.HexColor('#CCCCCC')

    doc = SimpleDocTemplate(
        OUT_PDF, pagesize=A4,
        leftMargin=3.0*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm,  bottomMargin=2.5*cm
    )

    FS = 12
    LD = 18

    s_title = ParagraphStyle('Title', fontName='Times-Bold', fontSize=14, leading=20,
                              alignment=TA_CENTER, textColor=DARK, spaceAfter=6)
    s_theme_label = ParagraphStyle('ThL', fontName='Times-Bold', fontSize=FS, leading=LD,
                                    textColor=DARK)
    s_body  = ParagraphStyle('Body', fontName='Times-Roman', fontSize=FS, leading=LD,
                              alignment=TA_JUSTIFY, textColor=DARK, spaceAfter=6)
    s_sec   = ParagraphStyle('Sec', fontName='Times-Bold', fontSize=FS, leading=LD,
                              textColor=DARK, spaceBefore=12, spaceAfter=4,
                              underlineProportion=0.04)
    s_h3    = ParagraphStyle('H3', fontName='Times-Bold', fontSize=FS, leading=LD,
                              textColor=DARK, spaceBefore=6, spaceAfter=2)
    s_bullet= ParagraphStyle('Bul', fontName='Times-Roman', fontSize=FS, leading=LD,
                              textColor=DARK, leftIndent=20, firstLineIndent=-10,
                              spaceAfter=3, alignment=TA_JUSTIFY)
    s_num   = ParagraphStyle('Num', fontName='Times-Roman', fontSize=FS, leading=LD,
                              textColor=DARK, leftIndent=22, firstLineIndent=-14,
                              spaceAfter=3, alignment=TA_JUSTIFY)

    story = []

    # Titre
    story.append(Paragraph(TITRE.replace('\n', '<br/>'), s_title))
    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8))

    # Thème
    theme_p = Paragraph(
        '<b>Proposition de thème :</b> JobFinder AI — Plateforme web intelligente de recherche '
        "d'emploi automatisée et de génération de candidatures personnalisées par intelligence "
        "artificielle générative",
        s_body
    )
    story.append(theme_p)
    story.append(Spacer(1, 0.3*cm))

    for sec in SECTIONS:
        story.append(Paragraph(f'<u><b>{sec["num"]}. {sec["title"]}</b></u>', s_sec))
        story.append(HRFlowable(width='100%', thickness=0.5, color=LBORDER, spaceAfter=4))

        for item_type, item_content in sec['body']:

            if item_type == 'para':
                story.append(Paragraph(item_content, s_body))

            elif item_type == 'heading3':
                story.append(Paragraph(f'<b>{item_content}</b>', s_h3))

            elif item_type == 'list':
                for bullet in item_content:
                    story.append(Paragraph(
                        f'<font color="#C9A84C"><b>\u25a0</b></font>  {bullet}',
                        s_bullet
                    ))

            elif item_type == 'numlist':
                for i, item in enumerate(item_content, 1):
                    story.append(Paragraph(
                        f'<font color="#C9A84C"><b>{i}.</b></font>  {item}',
                        s_num
                    ))

            elif item_type == 'table':
                rows = item_content
                tbl = Table(rows, colWidths=[6.0*cm, 10.5*cm])
                tbl.setStyle(TableStyle([
                    ('BACKGROUND',   (0, 0), (-1, 0),  colors.HexColor('#F5F0E8')),
                    ('TEXTCOLOR',    (0, 0), (-1, 0),  DARK),
                    ('FONTNAME',     (0, 0), (-1, 0),  'Times-Bold'),
                    ('FONTNAME',     (0, 1), (-1, -1), 'Times-Roman'),
                    ('FONTSIZE',     (0, 0), (-1, -1), 11),
                    ('LEADING',      (0, 0), (-1, -1), 16),
                    ('ROWBACKGROUNDS',(0,1), (-1,-1),  [colors.white, colors.HexColor('#FAFAFA')]),
                    ('GRID',         (0, 0), (-1, -1), 0.5, LBORDER),
                    ('VALIGN',       (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING',   (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING',(0, 0), (-1, -1), 5),
                    ('LEFTPADDING',  (0, 0), (-1, -1), 6),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    print(f"PDF généré : {OUT_PDF}")


# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    errors = []
    try:
        build_docx()
    except Exception as e:
        errors.append(f"DOCX erreur : {e}")
        print(f"DOCX erreur : {e}")
    try:
        build_pdf()
    except Exception as e:
        errors.append(f"PDF erreur : {e}")
        print(f"PDF erreur : {e}")

    if not errors:
        print("\nLes deux fichiers ont été générés avec succès sur le Bureau.")
    else:
        print(f"\n{len(errors)} erreur(s) rencontrée(s).")
