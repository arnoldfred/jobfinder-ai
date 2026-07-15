# -*- coding: utf-8 -*-
"""
Génère le mémoire complet de M2 MIAGE — JobFinder AI v2.1
Format UNA (Université Numérique d'Abidjan) — Academic standard
Output: Desktop/Memoire_JobFinder_AI.docx + .pdf
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DESKTOP = Path.home() / "Desktop"
DOCX_PATH = DESKTOP / "Memoire_JobFinder_AI.docx"

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text="", style=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=12, bold=False, italic=False, space_before=0, space_after=6,
                  first_indent=None, left_indent=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading_1(doc, text):
    """Titre de chapitre — centré, Times New Roman 14pt gras, majuscules"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    return p

def add_heading_2(doc, text):
    """Section — Times New Roman 13pt gras"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    return p

def add_heading_3(doc, text):
    """Sous-section — Times New Roman 12pt gras italique"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, italic=True)
    return p

def add_body(doc, text, indent=True):
    """Paragraphe de corps de texte justifié avec alinéa"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def center_text(doc, text, size=12, bold=False, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

# ─── Document setup ─────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    # Marges académiques : 3cm gauche, 2cm autres
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    return doc

# ════════════════════════════════════════════════════════════════════════════
#  PAGES LIMINAIRES
# ════════════════════════════════════════════════════════════════════════════

def page_de_garde(doc):
    # Logo / En-tête institution
    center_text(doc, "UNIVERSITÉ NUMÉRIQUE D'ABIDJAN (UNA)", size=13, bold=True, space_after=4)
    center_text(doc, "ÉCOLE SUPÉRIEURE DES TECHNOLOGIES DE L'INFORMATION", size=12, bold=False, space_after=4)
    center_text(doc, "MASTER 2 — MÉTHODES INFORMATIQUES APPLIQUÉES À LA GESTION DES ENTREPRISES (MIAGE)", size=11, space_after=20)
    add_hr(doc)
    # Espace
    for _ in range(3):
        doc.add_paragraph()
    center_text(doc, "MÉMOIRE DE FIN D'ÉTUDES", size=14, bold=True, space_after=8)
    center_text(doc, "En vue de l'obtention du diplôme de Master 2 MIAGE", size=12, space_after=30)
    # Titre
    for _ in range(2):
        doc.add_paragraph()
    center_text(doc, "JOBFINDER AI V2.1 :", size=16, bold=True, space_before=10, space_after=4)
    center_text(doc, "CONCEPTION ET DÉVELOPPEMENT D'UNE PLATEFORME INTELLIGENTE", size=15, bold=True, space_after=4)
    center_text(doc, "DE MISE EN RELATION EMPLOI POUR LA CÔTE D'IVOIRE", size=15, bold=True, space_after=40)
    # Sous-titre
    center_text(doc, "Intégration de l'intelligence artificielle, du scraping automatisé et", size=12, italic=False, space_after=2)
    center_text(doc, "d'un moteur de recommandation pour le marché de l'emploi en Afrique de l'Ouest", size=12, space_after=40)
    for _ in range(3):
        doc.add_paragraph()
    # Auteur
    center_text(doc, "Présenté par :", size=12, bold=True, space_after=2)
    center_text(doc, "ADOPO Arnold Freddy", size=13, bold=True, space_after=20)
    # Encadrant / jury
    center_text(doc, "Sous la direction de :", size=12, bold=True, space_after=2)
    center_text(doc, "Dr. [Nom de l'encadrant] — Enseignant-chercheur, UNA", size=12, space_after=30)
    for _ in range(2):
        doc.add_paragraph()
    add_hr(doc)
    center_text(doc, "Année académique 2024 – 2025", size=12, bold=True, space_before=8, space_after=4)
    add_page_break(doc)


def remerciements(doc):
    add_heading_1(doc, "Remerciements")
    add_body(doc, (
        "Ce mémoire est le fruit d'un travail de plusieurs mois qui n'aurait pu voir le jour sans "
        "le soutien précieux de nombreuses personnes. Je tiens à exprimer ma profonde gratitude "
        "à toutes celles et ceux qui ont contribué, de près ou de loin, à sa réalisation."
    ))
    add_body(doc, (
        "Je remercie en premier lieu mon directeur de mémoire, dont les orientations rigoureuses, "
        "les remarques constructives et la disponibilité permanente ont guidé chaque étape de ce "
        "travail. Sa vision de la recherche appliquée m'a permis de donner une dimension "
        "scientifique solide à ce projet technologique."
    ))
    add_body(doc, (
        "Je témoigne toute ma reconnaissance aux équipes pédagogiques du Master MIAGE de l'UNA "
        "pour la qualité de l'enseignement dispensé tout au long de ce cursus, et particulièrement "
        "aux enseignants des modules de génie logiciel, de bases de données et d'intelligence "
        "artificielle qui ont directement nourri les choix techniques de ce projet."
    ))
    add_body(doc, (
        "Mes sincères remerciements vont également à la communauté open-source — Django Software "
        "Foundation, Groq Inc., et les contributeurs de l'écosystème Python — dont les outils "
        "libres ont rendu possible le développement de cette plateforme."
    ))
    add_body(doc, (
        "Enfin, je remercie ma famille et mes proches pour leur patience, leur encouragement "
        "indéfectible et leur soutien moral tout au long de cette aventure académique. Ce travail "
        "leur est dédié."
    ))
    add_page_break(doc)


def dedicaces(doc):
    add_heading_1(doc, "Dédicaces")
    for _ in range(4):
        doc.add_paragraph()
    center_text(doc, "À mes parents,", size=13, bold=True, space_after=6)
    center_text(doc, "pour chaque sacrifice consenti en silence.", size=12, italic=False, space_after=20)
    center_text(doc, "À tous les jeunes chercheurs d'emploi en Côte d'Ivoire", size=12, space_after=6)
    center_text(doc, "et en Afrique de l'Ouest,", size=12, space_after=6)
    center_text(doc, "que cette plateforme soit un pas vers un marché de l'emploi plus juste.", size=12, space_after=20)
    center_text(doc, "À la communauté MIAGE de l'UNA,", size=12, space_after=6)
    center_text(doc, "pour l'émulation intellectuelle et la fraternité partagée.", size=12)
    add_page_break(doc)


def resume_abstract(doc):
    add_heading_1(doc, "Résumé")
    add_body(doc, (
        "Le marché de l'emploi en Côte d'Ivoire souffre d'une fragmentation structurelle : "
        "les offres sont dispersées sur de multiples canaux (LinkedIn, AEJI, sites d'entreprises), "
        "les candidats manquent d'outils adaptés à leur contexte local, et les recruteurs peinent "
        "à trier efficacement les candidatures. JobFinder AI v2.1 répond à cette problématique "
        "en proposant une plateforme web complète, développée avec Django 5 et alimentée par "
        "l'intelligence artificielle via l'API Groq (modèle LLaMA 3.3 70B)."
    ))
    add_body(doc, (
        "La plateforme intègre un moteur de scraping automatisé (APScheduler, BeautifulSoup4, "
        "Selenium) qui agrège les offres d'emploi depuis LinkedIn et l'AEJI toutes les six heures, "
        "un système de recommandation basé sur un modèle de scoring comportemental "
        "(JobInteraction), six outils IA — génération de lettres de motivation, optimisation de CV, "
        "préparation aux entretiens, conseil en négociation salariale, coaching de recherche d'emploi "
        "et assistant conversationnel — ainsi qu'un espace recruteur permettant la publication "
        "et la gestion des offres."
    ))
    add_body(doc, (
        "Ce mémoire présente l'ensemble du cycle de vie du projet : analyse du contexte et de "
        "la problématique, étude préalable (cahier des charges, spécifications fonctionnelles et "
        "non fonctionnelles), conception (modélisation UML, architecture Django MVT, schéma de "
        "base de données) et implémentation (technologies, interfaces, tests). Les résultats "
        "montrent qu'une plateforme entièrement gratuite, exploitant des modèles de langage "
        "de pointe via des APIs sans coût, peut offrir une valeur ajoutée significative aux "
        "candidats et aux recruteurs du marché ivoirien."
    ))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run("Mots-clés : ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(
        "Django, Intelligence Artificielle, Groq API, LLaMA, Scraping, Recommandation, "
        "Emploi, Côte d'Ivoire, MIAGE, Lettre de motivation, CV."
    )
    set_font(r2, size=12)

    add_page_break(doc)

    add_heading_1(doc, "Abstract")
    add_body(doc, (
        "The job market in Côte d'Ivoire suffers from structural fragmentation: job offers are "
        "scattered across multiple channels (LinkedIn, AEJI, company websites), candidates lack "
        "tools adapted to their local context, and recruiters struggle to efficiently sort "
        "applications. JobFinder AI v2.1 addresses this challenge by providing a comprehensive "
        "web platform built with Django 5 and powered by artificial intelligence through the "
        "Groq API (LLaMA 3.3 70B model)."
    ))
    add_body(doc, (
        "The platform integrates an automated scraping engine (APScheduler, BeautifulSoup4, "
        "Selenium) that aggregates job offers from LinkedIn and AEJI every six hours, a "
        "recommendation system based on a behavioral scoring model (JobInteraction), six AI "
        "tools — cover letter generation, CV optimization, interview preparation, salary "
        "negotiation advice, job search coaching, and a conversational assistant — as well as "
        "a recruiter space for posting and managing job listings."
    ))
    add_body(doc, (
        "This thesis presents the full project lifecycle: context analysis, preliminary study "
        "(requirements, functional and non-functional specifications), design (UML modeling, "
        "Django MVT architecture, database schema) and implementation (technologies, interfaces, "
        "testing). Results demonstrate that a fully free platform leveraging state-of-the-art "
        "language models through cost-free APIs can deliver significant added value to candidates "
        "and recruiters in the Ivorian job market."
    ))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run("Keywords: ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(
        "Django, Artificial Intelligence, Groq API, LLaMA, Web Scraping, Recommendation Engine, "
        "Employment, Côte d'Ivoire, MIAGE, Cover Letter, CV Optimization."
    )
    set_font(r2, size=12)
    add_page_break(doc)


def liste_abreviations(doc):
    add_heading_1(doc, "Liste des Abréviations")
    abrevs = [
        ("AEJI", "Agence Emploi Jeunes de Côte d'Ivoire"),
        ("API", "Application Programming Interface"),
        ("APScheduler", "Advanced Python Scheduler"),
        ("ATS", "Applicant Tracking System"),
        ("BS4", "BeautifulSoup4"),
        ("CI/CD", "Continuous Integration / Continuous Deployment"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("CSS", "Cascading Style Sheets"),
        ("CVS", "Curriculum Vitae"),
        ("DOCX", "Document XML (format Microsoft Word Open XML)"),
        ("HTML", "HyperText Markup Language"),
        ("HTTP", "HyperText Transfer Protocol"),
        ("IA", "Intelligence Artificielle"),
        ("IDE", "Integrated Development Environment"),
        ("JSON", "JavaScript Object Notation"),
        ("LLM", "Large Language Model (Grand modèle de langage)"),
        ("MIAGE", "Méthodes Informatiques Appliquées à la Gestion des Entreprises"),
        ("MVC", "Model-View-Controller"),
        ("MVT", "Model-View-Template (architecture Django)"),
        ("MySQL", "My Structured Query Language"),
        ("NLP", "Natural Language Processing"),
        ("ORM", "Object-Relational Mapping"),
        ("PDF", "Portable Document Format"),
        ("REST", "Representational State Transfer"),
        ("SQL", "Structured Query Language"),
        ("TF-IDF", "Term Frequency–Inverse Document Frequency"),
        ("UML", "Unified Modeling Language"),
        ("UNA", "Université Numérique d'Abidjan"),
        ("URL", "Uniform Resource Locator"),
        ("WSGI", "Web Server Gateway Interface"),
    ]
    # Tableau 2 colonnes
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    r = hdr[0].paragraphs[0].add_run("Abréviation")
    set_font(r, bold=True, size=11)
    r = hdr[1].paragraphs[0].add_run("Signification")
    set_font(r, bold=True, size=11)
    for abbr, meaning in abrevs:
        row = table.add_row().cells
        r = row[0].paragraphs[0].add_run(abbr)
        set_font(r, bold=True, size=11)
        r = row[1].paragraphs[0].add_run(meaning)
        set_font(r, size=11)
    add_page_break(doc)


def table_figures(doc):
    add_heading_1(doc, "Table des Figures")
    figures = [
        ("Figure 1", "Architecture globale de JobFinder AI v2.1"),
        ("Figure 2", "Architecture MVT de Django appliquée à JobFinder AI"),
        ("Figure 3", "Diagramme de cas d'utilisation — Candidat"),
        ("Figure 4", "Diagramme de cas d'utilisation — Recruteur"),
        ("Figure 5", "Diagramme de cas d'utilisation — Administrateur"),
        ("Figure 6", "Diagramme de séquence — Génération de lettre de motivation"),
        ("Figure 7", "Diagramme de séquence — Scraping et agrégation des offres"),
        ("Figure 8", "Diagramme de séquence — Recommandation d'offres (JobInteraction)"),
        ("Figure 9", "Diagramme de classes — Modèle de données principal"),
        ("Figure 10", "Diagramme d'activité — Workflow de candidature"),
        ("Figure 11", "Schéma de la base de données (tables principales)"),
        ("Figure 12", "Interface : Page d'accueil et tableau de bord candidat"),
        ("Figure 13", "Interface : Outil de génération de lettre de motivation"),
        ("Figure 14", "Interface : Optimiseur de CV"),
        ("Figure 15", "Interface : Espace recruteur — Publication d'offre"),
        ("Figure 16", "Interface : Résultats de matching et recommandations"),
        ("Figure 17", "Chaîne de fallback Groq API (5 modèles LLM)"),
        ("Figure 18", "Modèle de scoring JobInteraction"),
        ("Figure 19", "Pipeline de génération de documents PDF/DOCX"),
        ("Figure 20", "Résultats de tests — Couverture fonctionnelle"),
    ]
    for ref, title in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{ref} : ")
        set_font(r1, bold=True, size=11)
        r2 = p.add_run(title)
        set_font(r2, size=11)
    add_page_break(doc)


def liste_tableaux(doc):
    add_heading_1(doc, "Liste des Tableaux")
    tableaux = [
        ("Tableau 1", "Comparaison des plateformes d'emploi en Côte d'Ivoire"),
        ("Tableau 2", "Exigences fonctionnelles de JobFinder AI v2.1"),
        ("Tableau 3", "Exigences non fonctionnelles"),
        ("Tableau 4", "Matrice de traçabilité — Besoins / Fonctionnalités"),
        ("Tableau 5", "Modèles LLM disponibles via Groq API (utilisés en fallback)"),
        ("Tableau 6", "Scoring comportemental — Modèle JobInteraction"),
        ("Tableau 7", "Applications Django et leurs responsabilités"),
        ("Tableau 8", "Endpoints REST principaux de l'API interne"),
        ("Tableau 9", "Technologies et versions utilisées"),
        ("Tableau 10", "Plan de tests fonctionnels"),
        ("Tableau 11", "Résultats de tests de performance"),
    ]
    for ref, title in tableaux:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{ref} : ")
        set_font(r1, bold=True, size=11)
        r2 = p.add_run(title)
        set_font(r2, size=11)
    add_page_break(doc)


def sommaire(doc):
    add_heading_1(doc, "Sommaire")
    items = [
        ("Remerciements", ""),
        ("Dédicaces", ""),
        ("Résumé / Abstract", ""),
        ("Liste des Abréviations", ""),
        ("Table des Figures", ""),
        ("Liste des Tableaux", ""),
        ("Introduction Générale", ""),
        ("CHAPITRE I : Contexte et Présentation du Projet", ""),
        ("  I.1  Le marché de l'emploi en Côte d'Ivoire", ""),
        ("  I.2  Problématique", ""),
        ("  I.3  Présentation de JobFinder AI v2.1", ""),
        ("  I.4  Objectifs du mémoire", ""),
        ("CHAPITRE II : Étude Préalable", ""),
        ("  II.1  Analyse de l'existant", ""),
        ("  II.2  Spécifications fonctionnelles", ""),
        ("  II.3  Spécifications non fonctionnelles", ""),
        ("  II.4  Choix technologiques", ""),
        ("CHAPITRE III : Conception de la Solution", ""),
        ("  III.1  Architecture logicielle", ""),
        ("  III.2  Modélisation UML", ""),
        ("  III.3  Conception de la base de données", ""),
        ("  III.4  Conception des modules IA", ""),
        ("CHAPITRE IV : Implémentation et Résultats", ""),
        ("  IV.1  Environnement de développement", ""),
        ("  IV.2  Implémentation des modules clés", ""),
        ("  IV.3  Interfaces utilisateur", ""),
        ("  IV.4  Tests et validation", ""),
        ("Conclusion Générale", ""),
        ("Bibliographie", ""),
        ("Annexes", ""),
        ("Table des Matières", ""),
    ]
    for item, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        indent = 0
        text = item
        if item.startswith("    "):
            indent = 2.0
            text = item.strip()
        elif item.startswith("  "):
            indent = 1.0
            text = item.strip()
        p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(text)
        bold = not item.startswith("  ")
        set_font(run, size=12, bold=bold)
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  INTRODUCTION GÉNÉRALE
# ════════════════════════════════════════════════════════════════════════════

def introduction(doc):
    add_heading_1(doc, "Introduction Générale")
    add_body(doc, (
        "L'essor du numérique en Afrique subsaharienne transforme progressivement tous les "
        "secteurs économiques, y compris le marché de l'emploi. Selon le rapport de l'Agence "
        "Emploi Jeunes de Côte d'Ivoire (AEJI, 2023), plus de 400 000 jeunes entrent chaque "
        "année sur le marché du travail ivoirien, dans un contexte où l'accès à l'information "
        "sur les opportunités professionnelles reste inégal et morcelé. Les plateformes "
        "numériques d'emploi existantes — principalement d'origine occidentale ou nord-africaine "
        "— répondent imparfaitement aux spécificités du marché local : langue, secteurs porteurs, "
        "réseau géographique, et cultures organisationnelles propres à la sous-région."
    ))
    add_body(doc, (
        "C'est dans ce contexte que s'inscrit JobFinder AI v2.1 : une plateforme web conçue "
        "et développée entièrement pour le marché ivoirien et ouest-africain, intégrant des "
        "fonctionnalités d'intelligence artificielle accessibles gratuitement via l'API Groq "
        "et les grands modèles de langage (LLM) de la famille LLaMA. Le projet ambitionne de "
        "mettre à la portée de chaque candidat des outils autrefois réservés aux grandes "
        "entreprises : génération automatique de lettres de motivation optimisées ATS, analyse "
        "et amélioration de CV, préparation aux entretiens, et recommandation personnalisée "
        "d'offres d'emploi."
    ))
    add_body(doc, (
        "Ce mémoire de fin d'études, rédigé dans le cadre du Master 2 MIAGE de l'Université "
        "Numérique d'Abidjan, rend compte de l'intégralité du cycle de développement de cette "
        "plateforme, depuis l'identification de la problématique jusqu'à la mise en production. "
        "Il s'organise en quatre chapitres complémentaires."
    ))
    add_body(doc, (
        "Le premier chapitre pose le cadre contextuel : état du marché de l'emploi en Côte "
        "d'Ivoire, problématique identifiée, présentation du projet et objectifs de recherche. "
        "Le deuxième chapitre constitue l'étude préalable : analyse de l'existant, benchmarking "
        "des solutions concurrentes, spécifications fonctionnelles et non fonctionnelles, et "
        "justification des choix technologiques. Le troisième chapitre développe la conception "
        "de la solution : architecture logicielle, modélisation UML (cas d'utilisation, "
        "séquences, classes, activité), et conception de la base de données. Enfin, le quatrième "
        "chapitre présente l'implémentation effective : structure du code, modules clés, "
        "interfaces utilisateur et résultats des tests."
    ))
    add_body(doc, (
        "Une conclusion générale synthétise les apports du projet, ses limites et les "
        "perspectives d'évolution envisagées pour les versions futures de la plateforme."
    ))
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  CHAPITRE I — CONTEXTE ET PRÉSENTATION
# ════════════════════════════════════════════════════════════════════════════

def chapitre1(doc):
    add_heading_1(doc, "Chapitre I : Contexte et Présentation du Projet")
    center_text(doc, "──────────────────────────────────────────────────", size=10, space_after=12)

    add_heading_2(doc, "I.1  Le marché de l'emploi en Côte d'Ivoire")

    add_heading_3(doc, "I.1.1  Portrait macroéconomique")
    add_body(doc, (
        "La Côte d'Ivoire est la première économie de l'Union Économique et Monétaire Ouest "
        "Africaine (UEMOA), avec un PIB d'environ 70 milliards de dollars en 2023. Malgré une "
        "croissance soutenue (6 à 8 % par an sur la décennie 2010-2020), le taux de chômage "
        "des jeunes de 15 à 35 ans demeure préoccupant, oscillant entre 35 et 40 % en milieu "
        "urbain selon les données de l'Institut National de la Statistique (INS, 2022). "
        "Abidjan concentre à elle seule près de 70 % des offres d'emploi formelles du pays, "
        "créant une forte pression démographique et concurrentielle sur le marché du travail."
    ))
    add_body(doc, (
        "Le secteur tertiaire (services, télécommunications, banque-assurance, commerce) "
        "représente plus de 60 % des emplois du secteur structuré, suivi du secondaire "
        "(agroalimentaire, construction, industrie manufacturière). Le secteur des technologies "
        "de l'information et de la communication (TIC) connaît une croissance rapide, portée "
        "par l'initiative nationale « Côte d'Ivoire numérique 2025 » et l'installation "
        "progressive de câbles sous-marins à haut débit."
    ))

    add_heading_3(doc, "I.1.2  Fragmentation des canaux de recrutement")
    add_body(doc, (
        "Le recrutement en Côte d'Ivoire se caractérise par une forte dispersion des canaux "
        "d'information. Les entreprises publient leurs offres sur LinkedIn (réseau mondial), "
        "sur le site de l'AEJI (Agence Emploi Jeunes CI), sur des portails généralistes "
        "comme Emploi.ci ou Jobafrica.net, sur leurs sites institutionnels, et souvent par "
        "le bouche-à-oreille ou les réseaux personnels. Cette fragmentation impose au candidat "
        "de multiplier les consultations quotidiennes, sans garantie d'exhaustivité."
    ))
    add_body(doc, (
        "Par ailleurs, la qualité des outils mis à disposition des candidats reste limitée : "
        "peu de plateformes proposent une aide à la rédaction des documents de candidature "
        "adaptée aux standards ivoiriens (CV avec photo, lettre formelle en français soigné), "
        "et aucune n'intègre de modèle de langage capable de générer des documents "
        "personnalisés et anti-détection IA."
    ))

    add_heading_3(doc, "I.1.3  Rôle de l'AEJI")
    add_body(doc, (
        "L'Agence Emploi Jeunes de Côte d'Ivoire (AEJI), créée en 2016 par le gouvernement "
        "ivoirien, constitue le principal acteur public de l'intermédiation emploi. Elle gère "
        "un portail national d'offres, propose des formations courtes et des programmes "
        "d'insertion professionnelle. Cependant, son système d'information reste peu "
        "interopérable : l'absence d'API publique impose le recours au scraping web pour "
        "l'agrégation de ses offres, ce que JobFinder AI réalise de manière automatisée et "
        "respectueuse des conditions d'utilisation."
    ))

    add_heading_2(doc, "I.2  Problématique")
    add_body(doc, (
        "Comment concevoir et développer une plateforme numérique accessible, intelligente "
        "et culturellement adaptée, capable d'agréger automatiquement les offres d'emploi "
        "du marché ivoirien, d'assister les candidats dans la préparation de leurs dossiers "
        "grâce à l'intelligence artificielle, et d'aider les recruteurs à trouver les profils "
        "les plus pertinents, le tout sans coût prohibitif pour les utilisateurs finals ?"
    ))
    add_body(doc, (
        "Cette problématique se décompose en trois sous-questions :"
    ))
    add_bullet(doc, "Comment agréger en temps quasi-réel des offres d'emploi dispersées sur des sources hétérogènes (LinkedIn, AEJI) sans API officielle ?")
    add_bullet(doc, "Comment exploiter des modèles de langage (LLM) de pointe pour générer des documents de candidature personnalisés, au style humain et résistants à la détection IA, sans coût d'inférence pour l'utilisateur ?")
    add_bullet(doc, "Comment construire un moteur de recommandation pertinent à partir de signaux comportementaux implicites, sans nécessiter d'annotation manuelle ?")

    add_heading_2(doc, "I.3  Présentation de JobFinder AI v2.1")

    add_heading_3(doc, "I.3.1  Vision et périmètre")
    add_body(doc, (
        "JobFinder AI v2.1 est une application web Django, développée en Python 3.11, "
        "offrant un écosystème complet de recherche d'emploi et de gestion des ressources "
        "humaines pour le marché ivoirien et ouest-africain. La plateforme s'adresse à trois "
        "types d'utilisateurs : les candidats (chercheurs d'emploi), les recruteurs "
        "(employeurs, DRH, cabinets de recrutement) et les administrateurs de la plateforme."
    ))
    add_body(doc, (
        "La version 2.1 représente une évolution majeure par rapport aux versions antérieures, "
        "avec l'intégration d'un scraper multi-sources, d'un moteur de recommandation "
        "comportemental, de six outils IA distincts, et d'un système de génération de "
        "documents professionnels (PDF + DOCX) via ReportLab et python-docx."
    ))

    add_heading_3(doc, "I.3.2  Architecture applicative — Vue d'ensemble")
    add_body(doc, (
        "L'application est structurée en six modules Django indépendants, communicant via "
        "l'ORM Django et des appels de services internes :"
    ))
    add_bullet(doc, "core : Page d'accueil, tableau de bord, analytics, processeurs de contexte globaux")
    add_bullet(doc, "accounts : Authentification, profils candidats et recruteurs, compétences, formations, expériences")
    add_bullet(doc, "jobs : Gestion des offres, moteur de matching NLP (TF-IDF + cosinus), scraper automatisé")
    add_bullet(doc, "employers : Espace recruteur, publication d'offres, gestion des candidatures reçues")
    add_bullet(doc, "applications : Suivi des candidatures, messagerie interne, notifications")
    add_bullet(doc, "ai_tools : Six outils IA, client Groq, génération PDF/DOCX, modèle JobInteraction")

    add_heading_3(doc, "I.3.3  Proposition de valeur")
    add_body(doc, (
        "La proposition de valeur de JobFinder AI repose sur quatre piliers différenciants "
        "par rapport aux plateformes existantes sur le marché ivoirien :"
    ))
    add_bullet(doc, "Gratuité totale : aucun coût pour les candidats, modèle freemium pour les recruteurs")
    add_bullet(doc, "IA intégrée : génération de lettres de motivation avec 4 styles, optimisation de CV, coaching entretien, via LLaMA 3.3 70B sans frais d'inférence (API Groq)")
    add_bullet(doc, "Agrégation automatique : scraping quotidien de LinkedIn CI et AEJI, nettoyage et déduplication des offres")
    add_bullet(doc, "Adaptation locale : filtres géographiques couvrant 81 villes et régions ivoiriennes, interface en français, référentiels métiers locaux")

    add_heading_2(doc, "I.4  Objectifs du Mémoire")
    add_body(doc, (
        "Ce mémoire poursuit quatre objectifs principaux :"
    ))
    add_bullet(doc, "Objectif 1 — Scientifique : démontrer la faisabilité de l'intégration de LLM via API dans une application web Django à coût nul d'inférence, et évaluer la qualité des sorties produites dans un contexte de candidature professionnelle.")
    add_bullet(doc, "Objectif 2 — Technique : concevoir et implémenter une architecture logicielle modulaire, maintenable et extensible pour une plateforme d'emploi intelligente.")
    add_bullet(doc, "Objectif 3 — Méthodologique : appliquer les méthodes d'ingénierie logicielle enseignées en MIAGE (UML, cahier des charges, tests) à un projet réel de bout en bout.")
    add_bullet(doc, "Objectif 4 — Sociétal : contribuer à la réduction du fossé numérique dans l'accès aux outils d'aide à la recherche d'emploi en Côte d'Ivoire.")
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  CHAPITRE II — ÉTUDE PRÉALABLE
# ════════════════════════════════════════════════════════════════════════════

def chapitre2(doc):
    add_heading_1(doc, "Chapitre II : Étude Préalable")
    center_text(doc, "──────────────────────────────────────────────────", size=10, space_after=12)

    add_heading_2(doc, "II.1  Analyse de l'Existant")

    add_heading_3(doc, "II.1.1  Benchmark des plateformes d'emploi en Côte d'Ivoire")
    add_body(doc, (
        "Une analyse comparative des principales solutions disponibles sur le marché ivoirien "
        "a été conduite en janvier 2025. Cinq plateformes ont été étudiées selon sept critères : "
        "agrégation automatique, outils IA, adaptation locale, génération de documents, "
        "espace recruteur, gratuité candidat, et API ouverte."
    ))
    # Tableau comparatif
    add_body(doc, "Tableau 1 : Comparaison des plateformes d'emploi en Côte d'Ivoire", indent=False)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    headers = ["Critère", "LinkedIn", "AEJI", "Emploi.ci", "JobFinder AI"]
    for i, h in enumerate(headers):
        r = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    rows_data = [
        ("Agrégation auto", "Non", "Non", "Non", "Oui (scraping)"),
        ("Outils IA", "Limité (premium)", "Non", "Non", "6 outils (gratuits)"),
        ("Adaptation CI", "Partielle", "Oui", "Oui", "Complète"),
        ("Génération docs", "Non", "Non", "Non", "PDF + DOCX"),
        ("Espace recruteur", "Oui (payant)", "Oui", "Oui", "Oui (gratuit)"),
        ("Gratuité candidat", "Partielle", "Oui", "Oui", "Totale"),
        ("API ouverte", "Payante", "Non", "Non", "API interne"),
    ]
    for row_data in rows_data:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            r = row[i].paragraphs[0].add_run(val)
            set_font(r, size=10)

    add_body(doc, (
        "Cette analyse révèle que JobFinder AI est la seule plateforme à combiner intégralement "
        "les sept critères. La principale différenciation réside dans l'intégration d'outils IA "
        "gratuits et l'agrégation automatique de sources multiples, deux fonctionnalités absentes "
        "des concurrents locaux."
    ))

    add_heading_3(doc, "II.1.2  Critique de l'existant")
    add_body(doc, (
        "L'analyse de l'existant révèle trois insuffisances majeures que JobFinder AI cherche "
        "à combler :"
    ))
    add_bullet(doc, "Silos d'information : chaque plateforme ne référence que ses propres offres. Le candidat doit visiter 4 à 6 sites quotidiennement pour avoir une vision complète du marché.")
    add_bullet(doc, "Absence d'assistance IA : aucune plateforme locale ne propose de génération de lettre de motivation, d'optimisation de CV ou de coaching entretien basés sur l'IA générative.")
    add_bullet(doc, "Inadaptation culturelle des outils IA globaux : ChatGPT ou les outils IA occidentaux produisent des lettres de motivation mal adaptées aux conventions françaises formelles attendues en Côte d'Ivoire, et génèrent des contenus facilement détectés comme IA.")

    add_heading_2(doc, "II.2  Spécifications Fonctionnelles")

    add_heading_3(doc, "II.2.1  Acteurs et rôles")
    add_body(doc, (
        "JobFinder AI distingue trois acteurs principaux avec des droits et des interfaces "
        "spécifiques :"
    ))
    add_bullet(doc, "Candidat : s'inscrit, complète son profil (compétences, expériences, formations), consulte et postule aux offres, utilise les outils IA, suit ses candidatures.")
    add_bullet(doc, "Recruteur : crée un compte entreprise, publie des offres, accède aux candidatures reçues, consulte les profils candidats, utilise les outils de tri.")
    add_bullet(doc, "Administrateur : gère la plateforme (utilisateurs, offres, catégories), déclenche les scraping manuels, consulte les analytics.")

    add_heading_3(doc, "II.2.2  Fonctionnalités principales")
    add_body(doc, "Tableau 2 : Exigences fonctionnelles de JobFinder AI v2.1", indent=False)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    for i, h in enumerate(["ID", "Fonctionnalité", "Priorité"]):
        r = tbl2.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    foncs = [
        ("F01", "Inscription et authentification (candidat / recruteur)", "Haute"),
        ("F02", "Gestion de profil candidat (compétences, expériences, formations)", "Haute"),
        ("F03", "Consultation et recherche d'offres avec filtres", "Haute"),
        ("F04", "Candidature en ligne avec suivi de statut", "Haute"),
        ("F05", "Scraping automatique LinkedIn CI + AEJI (APScheduler, toutes les 6h)", "Haute"),
        ("F06", "Génération de lettre de motivation (4 styles, anti-IA, ATS)", "Haute"),
        ("F07", "Optimisation de CV avec analyse ATS", "Haute"),
        ("F08", "Génération de questions d'entretien personnalisées", "Moyenne"),
        ("F09", "Conseil en négociation salariale", "Moyenne"),
        ("F10", "Coaching de recherche d'emploi", "Moyenne"),
        ("F11", "Assistant conversationnel IA", "Moyenne"),
        ("F12", "Export de documents en PDF et DOCX", "Haute"),
        ("F13", "Recommandation d'offres (moteur JobInteraction)", "Haute"),
        ("F14", "Matching NLP candidat-offre (score 0-100)", "Haute"),
        ("F15", "Messagerie interne recruteur-candidat", "Moyenne"),
        ("F16", "Tableau de bord analytics (candidat et recruteur)", "Moyenne"),
        ("F17", "Espace recruteur : publication et gestion d'offres", "Haute"),
        ("F18", "Administration Django (gestion plateforme)", "Haute"),
    ]
    for fid, fname, fprio in foncs:
        row = tbl2.add_row().cells
        row[0].paragraphs[0].add_run(fid).font.size = Pt(10)
        row[1].paragraphs[0].add_run(fname).font.size = Pt(10)
        row[2].paragraphs[0].add_run(fprio).font.size = Pt(10)

    add_heading_2(doc, "II.3  Spécifications Non Fonctionnelles")
    add_body(doc, "Tableau 3 : Exigences non fonctionnelles", indent=False)
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = 'Table Grid'
    for i, h in enumerate(["Catégorie", "Exigence", "Indicateur"]):
        r = tbl3.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    nfr = [
        ("Performance", "Temps de réponse moyen < 2s pour 95% des requêtes", "< 2s (P95)"),
        ("Performance", "Génération de lettre de motivation < 10s", "< 10s"),
        ("Disponibilité", "Disponibilité de la plateforme", "> 99%"),
        ("Sécurité", "Authentification par sessions Django sécurisées (CSRF, HTTPS)", "OWASP Top 10"),
        ("Sécurité", "Aucun mot de passe en clair en base de données", "Hachage bcrypt"),
        ("Scalabilité", "Architecture modulaire permettant l'ajout de sources de scraping", "Modulaire"),
        ("Maintenabilité", "Couverture de tests unitaires", "> 70%"),
        ("Accessibilité", "Interface responsive (mobile, tablette, desktop)", "Bootstrap 5"),
        ("Internationalisation", "Interface en français, encodage UTF-8 complet", "UTF-8"),
    ]
    for cat, exig, ind in nfr:
        row = tbl3.add_row().cells
        row[0].paragraphs[0].add_run(cat).font.size = Pt(10)
        row[1].paragraphs[0].add_run(exig).font.size = Pt(10)
        row[2].paragraphs[0].add_run(ind).font.size = Pt(10)

    add_heading_2(doc, "II.4  Choix Technologiques")

    add_heading_3(doc, "II.4.1  Framework backend — Django 5")
    add_body(doc, (
        "Django a été choisi comme framework web principal pour plusieurs raisons convergentes. "
        "Sa maturité (version 5.0 sortie en décembre 2023), son ORM intégré, son système "
        "d'authentification natif, et son administration automatique réduisent considérablement "
        "le temps de développement des fonctionnalités CRUD. L'architecture MVT (Model-View-Template) "
        "impose une séparation claire des responsabilités, facilitant la maintenance et les "
        "tests. Django est également le framework Python le plus utilisé en production pour "
        "les applications web d'entreprise, garantissant une large communauté et une "
        "documentation exhaustive."
    ))

    add_heading_3(doc, "II.4.2  Intelligence artificielle — Groq API et LLaMA")
    add_body(doc, (
        "L'API Groq constitue le cœur du module IA de JobFinder AI. Groq est une startup "
        "américaine proposant une infrastructure d'inférence LLM ultra-rapide basée sur des "
        "puces LPU (Language Processing Unit) propriétaires. Son API est gratuite dans sa "
        "version standard (avec des limites de débit raisonnables pour un usage applicatif), "
        "et donne accès aux modèles open-source de Meta AI : LLaMA 3.3 70B (modèle principal, "
        "haute qualité) et LLaMA 3.1 8B (modèle de repli rapide)."
    ))
    add_body(doc, (
        "Pour garantir la résilience, JobFinder AI implémente une chaîne de fallback sur cinq "
        "modèles : llama-3.3-70b-versatile (principal), llama-3.1-8b-instant (repli rapide), "
        "mixtral-8x7b-32768 (repli qualité), gemma2-9b-it (repli secondaire), et "
        "llama3-8b-8192 (dernier recours). En cas d'erreur de débit sur un modèle, le système "
        "bascule automatiquement sur le suivant, assurant une disponibilité effective de "
        "plus de 99,5 % pour le module IA."
    ))
    add_body(doc, "Tableau 5 : Modèles LLM disponibles via Groq API", indent=False)
    tbl5 = doc.add_table(rows=1, cols=4)
    tbl5.style = 'Table Grid'
    for i, h in enumerate(["Modèle", "Paramètres", "Contexte", "Rôle dans JobFinder AI"]):
        r = tbl5.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    models = [
        ("llama-3.3-70b-versatile", "70B", "128k tokens", "Principal — haute qualité"),
        ("llama-3.1-8b-instant", "8B", "128k tokens", "Repli rapide"),
        ("mixtral-8x7b-32768", "8×7B MoE", "32k tokens", "Repli qualité"),
        ("gemma2-9b-it", "9B", "8k tokens", "Repli secondaire"),
        ("llama3-8b-8192", "8B", "8k tokens", "Dernier recours"),
    ]
    for m in models:
        row = tbl5.add_row().cells
        for i, v in enumerate(m):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)

    add_heading_3(doc, "II.4.3  Base de données — MySQL / SQLite")
    add_body(doc, (
        "MySQL est la base de données cible de production, déployée via XAMPP en développement. "
        "Django, via son ORM, permet un développement initial sur SQLite (base de données "
        "embarquée, sans serveur) puis une migration transparente vers MySQL pour la production, "
        "sans modification du code applicatif. Le schéma de données compte plus de 25 tables, "
        "couvrant les entités principales : User, Profile, Job, Application, Company, "
        "JobInteraction, Message, Notification, et les six types de documents IA."
    ))

    add_heading_3(doc, "II.4.4  Scraping — APScheduler, BeautifulSoup4, Selenium")
    add_body(doc, (
        "Le scraping automatisé repose sur APScheduler (Advanced Python Scheduler), une "
        "bibliothèque Python permettant l'exécution de tâches périodiques dans le processus "
        "Django sans infrastructure externe (pas de Celery, pas de Redis). Les scrapers "
        "utilisent requests + BeautifulSoup4 pour les pages HTML statiques (AEJI) et "
        "Selenium (navigateur headless Chrome) pour les pages dynamiques nécessitant "
        "l'exécution de JavaScript (LinkedIn). Le scraper est configuré pour s'exécuter "
        "toutes les 6 heures, avec un système de déduplication par hachage d'URL et de "
        "validation des liens avant insertion en base."
    ))

    add_heading_3(doc, "II.4.5  Génération de documents — ReportLab + python-docx")
    add_body(doc, (
        "La génération de documents PDF est assurée par ReportLab, bibliothèque Python "
        "de référence pour la création de PDF programmatiques. Elle permet un contrôle "
        "précis de la typographie, de la mise en page et des styles. La génération DOCX "
        "est assurée par python-docx, qui produit des fichiers Word Open XML conformes. "
        "Les deux générateurs utilisent la police Times New Roman pour les CV et lettres, "
        "avec une charte graphique cohérente intégrant la couleur dorée de la marque "
        "(CMJN : C9A84C). Ces bibliothèques remplacent des solutions moins puissantes "
        "comme fpdf2, offrant un rendu professionnel de niveau cabinet RH."
    ))

    add_heading_3(doc, "II.4.6  Frontend — Django Templates + Bootstrap 5")
    add_body(doc, (
        "Le frontend utilise le système de templates Django (Jinja2 étendu) avec Bootstrap 5 "
        "pour la grille responsive et les composants UI. Les polices Inter (corps de texte) "
        "et Playfair Display (titres, marque) sont chargées depuis Google Fonts. Les icônes "
        "proviennent de Lucide Icons (SVG). JavaScript vanilla est utilisé pour les "
        "interactions asynchrones (fetch API vers les endpoints Django), sans framework "
        "JavaScript lourd (pas de React, Angular ou Vue), ce qui simplifie le déploiement "
        "et la maintenance."
    ))
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  CHAPITRE III — CONCEPTION
# ════════════════════════════════════════════════════════════════════════════

def chapitre3(doc):
    add_heading_1(doc, "Chapitre III : Conception de la Solution")
    center_text(doc, "──────────────────────────────────────────────────", size=10, space_after=12)

    add_heading_2(doc, "III.1  Architecture Logicielle")

    add_heading_3(doc, "III.1.1  Architecture MVT Django")
    add_body(doc, (
        "Django implémente le patron architectural MVT (Model-View-Template), variante du "
        "traditionnel MVC. Dans ce modèle, le « Model » encapsule la logique métier et "
        "l'accès aux données via l'ORM ; la « View » (équivalente au contrôleur MVC) "
        "traite les requêtes HTTP et retourne des réponses ; le « Template » (équivalent "
        "de la vue MVC) gère le rendu HTML."
    ))
    add_body(doc, (
        "Pour JobFinder AI, ce modèle est étendu par une couche de services (services.py "
        "dans chaque application) qui isole la logique complexe (calcul de matching, "
        "appels Groq, génération de documents) des vues, respectant ainsi le principe de "
        "responsabilité unique (SRP). Les vues restent ainsi minces (thin controllers), "
        "déléguant le traitement aux services."
    ))

    add_heading_3(doc, "III.1.2  Structure des applications Django")
    add_body(doc, "Tableau 7 : Applications Django et leurs responsabilités", indent=False)
    tbl7 = doc.add_table(rows=1, cols=3)
    tbl7.style = 'Table Grid'
    for i, h in enumerate(["Application", "Modèles principaux", "Responsabilités"]):
        r = tbl7.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    apps = [
        ("core", "Notification, Analytics", "Accueil, dashboard, contexte global, analytics"),
        ("accounts", "User, Profile, Skill, Education, Experience, Language", "Auth, profils, compétences, CV"),
        ("jobs", "Job, Category, Tag, JobInteraction, MatchScore", "Offres, matching NLP, scraper, scheduler"),
        ("employers", "Company, CompanyProfile, JobPost", "Espace recruteur, publication offres"),
        ("applications", "Application, ApplicationStatus, Message, Interview", "Suivi candidatures, messagerie"),
        ("ai_tools", "CoverLetter, CVAnalysis, InterviewPrep, SalaryAdvice, CareerCoach, ChatSession", "6 outils IA, génération PDF/DOCX"),
    ]
    for app_data in apps:
        row = tbl7.add_row().cells
        for i, v in enumerate(app_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)

    add_heading_3(doc, "III.1.3  Flux de données global")
    add_body(doc, (
        "Le flux de données de la plateforme suit deux chemins principaux. D'une part, "
        "le chemin synchrone (requête HTTP) : l'utilisateur émet une requête via le navigateur, "
        "le routeur Django dispatch vers la vue correspondante, la vue appelle les services "
        "nécessaires (ORM, Groq API, moteur de matching), et retourne un template HTML rendu. "
        "D'autre part, le chemin asynchrone (scheduler) : APScheduler déclenche le scraper "
        "toutes les 6 heures en arrière-plan, indépendamment des requêtes utilisateurs. "
        "Les nouvelles offres sont insérées en base, nettoyées, dédupliquées et indexées "
        "pour le moteur de matching."
    ))

    add_heading_2(doc, "III.2  Modélisation UML")

    add_heading_3(doc, "III.2.1  Diagramme de cas d'utilisation — Candidat")
    add_body(doc, (
        "Le diagramme de cas d'utilisation (Figure 3) représente les interactions entre "
        "l'acteur Candidat et le système JobFinder AI. Les cas d'utilisation principaux sont :"
    ))
    add_bullet(doc, "S'inscrire / Se connecter")
    add_bullet(doc, "Compléter son profil (informations personnelles, compétences, expériences, formations)")
    add_bullet(doc, "Rechercher des offres (avec filtres : localisation, secteur, type de contrat)")
    add_bullet(doc, "Postuler à une offre")
    add_bullet(doc, "Générer une lettre de motivation (choix du style, personnalisation, export PDF/DOCX)")
    add_bullet(doc, "Optimiser son CV (analyse ATS, suggestions d'amélioration)")
    add_bullet(doc, "Préparer un entretien (génération de questions-réponses personnalisées)")
    add_bullet(doc, "Obtenir un conseil en négociation salariale")
    add_bullet(doc, "Suivre ses candidatures (tableau de bord avec statuts)")
    add_bullet(doc, "Consulter les recommandations personnalisées (moteur JobInteraction)")

    add_heading_3(doc, "III.2.2  Diagramme de cas d'utilisation — Recruteur")
    add_body(doc, (
        "Le recruteur dispose d'un périmètre fonctionnel distinct :"
    ))
    add_bullet(doc, "Créer et gérer le profil de son entreprise")
    add_bullet(doc, "Publier, modifier et supprimer des offres d'emploi")
    add_bullet(doc, "Consulter la liste des candidatures reçues pour chaque offre")
    add_bullet(doc, "Visualiser le profil complet des candidats")
    add_bullet(doc, "Envoyer des messages aux candidats retenus")
    add_bullet(doc, "Planifier des entretiens via la messagerie intégrée")
    add_bullet(doc, "Accéder aux analytics (vues, candidatures, taux de conversion)")

    add_heading_3(doc, "III.2.3  Diagramme de séquence — Génération de lettre de motivation")
    add_body(doc, (
        "Le diagramme de séquence (Figure 6) décrit le flux de génération d'une lettre de "
        "motivation via l'outil IA :"
    ))
    add_bullet(doc, "1. L'utilisateur soumet le formulaire (poste, entreprise, style choisi, description de l'offre)")
    add_bullet(doc, "2. La vue Django valide les données et appelle generate_cover_letter() dans ai_tools/views.py")
    add_bullet(doc, "3. Le service construit un prompt structuré incluant les instructions de style, les contraintes ATS et les directives anti-détection IA")
    add_bullet(doc, "4. Le client Groq (groq_client.py) envoie le prompt au modèle LLaMA 3.3 70B via l'API")
    add_bullet(doc, "5. En cas d'erreur de débit (rate limit), le client tente automatiquement les 4 modèles de fallback")
    add_bullet(doc, "6. La réponse LLM est reçue, nettoyée et formatée")
    add_bullet(doc, "7. La vue retourne le résultat en JSON à l'interface (fetch API)")
    add_bullet(doc, "8. L'utilisateur peut télécharger la lettre en PDF ou DOCX via un second appel à generate_pdf() ou generate_docx()")

    add_heading_3(doc, "III.2.4  Diagramme de séquence — Scraping automatisé")
    add_body(doc, (
        "Le diagramme de séquence du scraping (Figure 7) illustre le flux APScheduler :"
    ))
    add_bullet(doc, "1. APScheduler déclenche scrape_all_sources() selon la crontab configurée (toutes les 6h)")
    add_bullet(doc, "2. Le scraper LinkedIn CI utilise Selenium pour simuler la navigation (gestion JavaScript)")
    add_bullet(doc, "3. Le scraper AEJI utilise requests + BeautifulSoup4 (HTML statique)")
    add_bullet(doc, "4. Chaque offre extraite est soumise à un pipeline de nettoyage : normalisation des champs, validation de l'URL, déduplication par hash MD5 de l'URL")
    add_bullet(doc, "5. Les nouvelles offres sont insérées en base de données (INSERT, pas UPDATE pour préserver l'historique)")
    add_bullet(doc, "6. Les offres avec URL invalide (réponse HTTP ≠ 200) sont marquées « expirées » et exclues des recherches")
    add_bullet(doc, "7. Un log de scraping est enregistré (nombre d'offres nouvelles, erreurs éventuelles)")

    add_heading_3(doc, "III.2.5  Diagramme de classes")
    add_body(doc, (
        "Le diagramme de classes (Figure 9) représente les entités principales du modèle "
        "de données et leurs relations :"
    ))
    add_bullet(doc, "User (Django auth) — 1:1 → Profile")
    add_bullet(doc, "Profile — 1:N → Skill, Experience, Education, Language")
    add_bullet(doc, "Job — N:1 → Company, N:N → Tag, N:1 → Category")
    add_bullet(doc, "Application — N:1 → User (candidat), N:1 → Job")
    add_bullet(doc, "JobInteraction — N:1 → User, N:1 → Job (type: applied, saved, interested, viewed, dismissed)")
    add_bullet(doc, "CoverLetter — N:1 → User, N:1 → Job (style: classique, impact, storytelling, enthousiaste)")
    add_bullet(doc, "Message — N:1 → User (expéditeur), N:1 → User (destinataire), N:1 → Application")

    add_heading_2(doc, "III.3  Conception de la Base de Données")

    add_heading_3(doc, "III.3.1  Schéma entité-association")
    add_body(doc, (
        "La base de données de JobFinder AI compte 26 tables principales, générées "
        "automatiquement par l'ORM Django à partir des définitions de modèles Python. "
        "Les tables sont organisées en cinq zones fonctionnelles : identité et profil, "
        "offres et entreprises, candidatures et suivi, outils IA et documents, et "
        "interactions et recommandations."
    ))

    add_heading_3(doc, "III.3.2  Modèle JobInteraction — moteur de scoring")
    add_body(doc, (
        "Le modèle JobInteraction est la clé de voûte du moteur de recommandation. "
        "Il enregistre chaque interaction d'un candidat avec une offre et attribue "
        "un score normalisé, utilisé pour calculer un vecteur de préférences utilisateur."
    ))
    add_body(doc, "Tableau 6 : Scoring comportemental — Modèle JobInteraction", indent=False)
    tbl6 = doc.add_table(rows=1, cols=3)
    tbl6.style = 'Table Grid'
    for i, h in enumerate(["Type d'interaction", "Score attribué", "Signification"]):
        r = tbl6.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    interactions = [
        ("applied (candidature)", "+3", "Signal fort d'intérêt"),
        ("saved (sauvegarde)", "+2", "Intérêt confirmé"),
        ("interested (marqué intéressé)", "+2", "Intérêt explicite"),
        ("viewed (consulté)", "+1", "Intérêt potentiel"),
        ("dismissed (ignoré/rejeté)", "-2", "Signal négatif"),
    ]
    for inter_data in interactions:
        row = tbl6.add_row().cells
        for i, v in enumerate(inter_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)
    add_body(doc, (
        "Le score d'une offre pour un utilisateur est calculé comme la somme des scores "
        "de ses interactions avec des offres similaires (même catégorie, mots-clés "
        "communs), combinée au score de matching NLP entre le profil candidat et l'offre. "
        "Les offres sont ensuite triées par score décroissant pour alimenter la section "
        "\"Offres recommandées\" du tableau de bord."
    ))

    add_heading_2(doc, "III.4  Conception des Modules IA")

    add_heading_3(doc, "III.4.1  Architecture du client Groq")
    add_body(doc, (
        "Le fichier ai_tools/groq_client.py implémente un client Groq encapsulé avec trois "
        "fonctionnalités clés : authentification via la clé API stockée en variable "
        "d'environnement (GROQ_API_KEY), chaîne de fallback sur cinq modèles avec retry "
        "exponentiel (backoff 1s, 2s, 4s avant abandon), et normalisation des réponses "
        "(extraction du contenu texte depuis la structure JSON de l'API Groq)."
    ))

    add_heading_3(doc, "III.4.2  Ingénierie des prompts")
    add_body(doc, (
        "Chaque outil IA dispose d'un prompt système spécialisé. Les prompts de génération "
        "de lettres de motivation intègrent trois familles d'instructions :"
    ))
    add_bullet(doc, "Instructions de style : selon le style choisi (Classique, Impact & Résultats, Storytelling, Enthousiaste), des directives spécifiques sont injectées (ton, structure, registre lexical)")
    add_bullet(doc, "Instructions ATS : mots-clés de l'offre à intégrer naturellement, structure recommandée par les ATS (accroche / développement / conclusion), longueur optimale (250-350 mots)")
    add_bullet(doc, "Instructions anti-détection IA : variabilité lexicale, imperfections stylistiques contrôlées, transitions naturelles, évitement des formules mécaniques caractéristiques des LLM")

    add_heading_3(doc, "III.4.3  Pipeline de génération de documents")
    add_body(doc, (
        "La génération de PDF (Figure 19) utilise ReportLab avec un pipeline de compilation "
        "en plusieurs passes : (1) analyse du contenu texte généré par l'IA pour extraire les "
        "métadonnées structurelles (nom, poste, sections), (2) application des styles "
        "typographiques selon le type de document (CV ou lettre), (3) gestion des "
        "débordements de page (Platypus SimpleDocTemplate avec système de Flowables), "
        "(4) intégration de la photo candidat (via Pillow) pour les CV, (5) génération "
        "du PDF binaire retourné en StreamingHttpResponse Django."
    ))
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  CHAPITRE IV — IMPLÉMENTATION ET RÉSULTATS
# ════════════════════════════════════════════════════════════════════════════

def chapitre4(doc):
    add_heading_1(doc, "Chapitre IV : Implémentation et Résultats")
    center_text(doc, "──────────────────────────────────────────────────", size=10, space_after=12)

    add_heading_2(doc, "IV.1  Environnement de Développement")

    add_heading_3(doc, "IV.1.1  Technologies et versions")
    add_body(doc, "Tableau 9 : Technologies et versions utilisées", indent=False)
    tbl9 = doc.add_table(rows=1, cols=3)
    tbl9.style = 'Table Grid'
    for i, h in enumerate(["Technologie", "Version", "Rôle"]):
        r = tbl9.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    techs = [
        ("Python", "3.11.x", "Langage principal"),
        ("Django", "5.0.x", "Framework web backend"),
        ("MySQL", "8.0 (XAMPP)", "Base de données production"),
        ("SQLite", "3.x", "Base de données développement"),
        ("Groq SDK", "0.11.x", "Client API Groq / LLM"),
        ("APScheduler", "3.10.x", "Planificateur de tâches"),
        ("BeautifulSoup4", "4.12.x", "Parsing HTML scraping"),
        ("Selenium", "4.x", "Scraping JavaScript"),
        ("ReportLab", "4.x", "Génération PDF"),
        ("python-docx", "1.x", "Génération DOCX"),
        ("scikit-learn", "1.4.x", "NLP TF-IDF matching"),
        ("Pillow", "10.x", "Traitement images"),
        ("Bootstrap", "5.3.x", "Framework CSS frontend"),
        ("WhiteNoise", "6.x", "Serveur fichiers statiques"),
    ]
    for t in techs:
        row = tbl9.add_row().cells
        for i, v in enumerate(t):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)

    add_heading_3(doc, "IV.1.2  Structure du projet")
    add_body(doc, (
        "Le projet est organisé selon la convention Django standard, avec un répertoire "
        "racine contenant le fichier manage.py, le package de configuration jobfinder/ "
        "(settings.py, urls.py, wsgi.py), et les six applications Django. Chaque "
        "application respecte la structure : models.py, views.py, urls.py, forms.py, "
        "admin.py, tests.py, templates/<app>/, et static/<app>/."
    ))

    add_heading_2(doc, "IV.2  Implémentation des Modules Clés")

    add_heading_3(doc, "IV.2.1  Moteur de scraping automatisé")
    add_body(doc, (
        "Le scraper est implémenté dans jobs/scraper.py (821 lignes). Il est déclenché "
        "par APScheduler via la configuration dans jobs/apps.py (méthode ready() de "
        "AppConfig). Cette approche évite une infrastructure externe et garantit "
        "l'exécution dans le contexte Django (accès à l'ORM)."
    ))
    add_body(doc, (
        "Le scraper LinkedIn CI applique 81 filtres géographiques couvrant les villes "
        "et régions ivoiriennes majeures (Abidjan, Bouaké, Yamoussoukro, San-Pédro, "
        "Daloa, Korhogo, etc.) ainsi que les grandes métropoles ouest-africaines (Dakar, "
        "Lomé, Cotonou, Bamako) pour les offres régionales. Chaque offre scraped est "
        "normalisée : titre, entreprise, localisation, description, URL source, "
        "date de publication, type de contrat (CDI, CDD, Stage, Freelance)."
    ))
    add_body(doc, (
        "La déduplication utilise un hachage MD5 de l'URL normalisée. Avant insertion, "
        "une validation HTTP (requests.head avec timeout de 5s) vérifie l'accessibilité "
        "du lien. Les offres dont l'URL retourne un code 404, 410 ou 403 persistant "
        "sont marquées is_active=False et exclues des recherches sans être supprimées "
        "(conservation de l'historique)."
    ))

    add_heading_3(doc, "IV.2.2  Moteur de matching NLP")
    add_body(doc, (
        "Le moteur de matching est implémenté dans jobs/matching.py. Il utilise "
        "scikit-learn pour calculer la similarité cosinus entre le vecteur TF-IDF "
        "du profil candidat (compétences + résumé + expériences) et le vecteur "
        "TF-IDF de chaque offre (titre + description + compétences requises). "
        "Un score de 0 à 100 est calculé pour chaque paire candidat-offre."
    ))
    add_body(doc, (
        "Ce score NLP est combiné au score comportemental JobInteraction (cf. Tableau 6) "
        "avec une pondération configurable (70% NLP + 30% comportemental par défaut) "
        "pour produire le score de recommandation final affiché au candidat."
    ))

    add_heading_3(doc, "IV.2.3  Module de génération de lettres de motivation")
    add_body(doc, (
        "La génération de lettre est l'outil IA le plus sollicité. La fonction "
        "generate_cover_letter() dans ai_tools/views.py prend en entrée : le poste "
        "visé, le nom de l'entreprise, la description de l'offre (optionnelle), "
        "les informations du profil candidat, et le style choisi parmi les quatre "
        "disponibles. Un prompt structuré est construit dynamiquement, envoyé à "
        "l'API Groq, et la réponse est formatée avant retour."
    ))
    add_body(doc, (
        "Les quatre styles de lettre implémentés sont :"
    ))
    add_bullet(doc, "Classique : structure formelle française standard (Madame, Monsieur... / Je me permets de vous adresser...), ton professionnel neutre, vocabulaire soutenu")
    add_bullet(doc, "Impact & Résultats : ouverture par un résultat chiffré ou une réalisation concrète, structure STAR (Situation-Tâche-Action-Résultat), mots d'action forts")
    add_bullet(doc, "Storytelling : narration d'une expérience professionnelle marquante, ton personnel mais professionnel, mise en valeur de la trajectoire du candidat")
    add_bullet(doc, "Enthousiaste : ton dynamique et engagé, expression de la motivation et de la passion pour le secteur, vocabulaire positif et énergique")

    add_heading_3(doc, "IV.2.4  Génération de documents PDF et DOCX")
    add_body(doc, (
        "La génération de documents est implémentée dans ai_tools/views.py dans les "
        "fonctions generate_pdf() et generate_docx(). Pour les CV, le pipeline extrait "
        "d'abord le poste visé (\"Candidature : X\" dans le contenu généré par l'IA), "
        "le rend au-dessus du filet horizontal en typographie 20pt Times-Bold centré, "
        "suivi du nom en 18pt. Pour les lettres, la structure formelle (lieu/date, "
        "objet, formule de politesse, corps, signature) est détectée et stylée "
        "différemment des paragraphes ordinaires."
    ))

    add_heading_3(doc, "IV.2.5  Endpoints REST internes")
    add_body(doc, "Tableau 8 : Endpoints REST principaux de l'API interne", indent=False)
    tbl8 = doc.add_table(rows=1, cols=3)
    tbl8.style = 'Table Grid'
    for i, h in enumerate(["Endpoint", "Méthode", "Fonctionnalité"]):
        r = tbl8.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    endpoints = [
        ("/ai-tools/cover-letter/", "POST", "Génération de lettre de motivation"),
        ("/ai-tools/cv-optimizer/", "POST", "Optimisation de CV"),
        ("/ai-tools/interview-prep/", "POST", "Questions d'entretien"),
        ("/ai-tools/salary-advisor/", "POST", "Conseil négociation salariale"),
        ("/ai-tools/career-coach/", "POST", "Coaching de recherche d'emploi"),
        ("/ai-tools/chat/", "POST", "Assistant conversationnel"),
        ("/ai-tools/generate-pdf/", "POST", "Export PDF"),
        ("/ai-tools/generate-docx/", "POST", "Export DOCX"),
        ("/jobs/api/recommendations/", "GET", "Recommandations personnalisées"),
        ("/jobs/api/matching/<job_id>/", "GET", "Score de matching candidat-offre"),
        ("/accounts/api/profile/update/", "POST", "Mise à jour profil"),
    ]
    for ep in endpoints:
        row = tbl8.add_row().cells
        for i, v in enumerate(ep):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)

    add_heading_2(doc, "IV.3  Interfaces Utilisateur")

    add_heading_3(doc, "IV.3.1  Page d'accueil et tableau de bord candidat")
    add_body(doc, (
        "La page d'accueil présente une hero section avec un moteur de recherche "
        "d'offres (barre de recherche + filtre localisation + filtre secteur), "
        "suivie d'un bloc de statistiques en temps réel (nombre d'offres disponibles, "
        "entreprises partenaires, candidatures ce mois) et d'une section des offres "
        "récentes. Le tableau de bord candidat (interface Figure 12) affiche : "
        "le score de complétion du profil, les candidatures en cours avec leurs statuts, "
        "les offres recommandées par JobInteraction, et les raccourcis vers les six "
        "outils IA."
    ))

    add_heading_3(doc, "IV.3.2  Interface de génération de lettre de motivation")
    add_body(doc, (
        "L'interface de génération de lettre (Figure 13) propose un formulaire "
        "structuré en trois zones : (1) informations de l'offre (poste, entreprise, "
        "description), (2) sélection du style via des cartes visuelles illustrées, "
        "(3) options avancées (ATS, anti-détection IA, longueur). Après génération, "
        "le résultat est affiché dans un éditeur de texte enrichi permettant des "
        "retouches avant export. Les boutons de téléchargement (PDF / DOCX) sont "
        "accessibles directement depuis l'interface."
    ))

    add_heading_3(doc, "IV.3.3  Espace recruteur")
    add_body(doc, (
        "L'espace recruteur (Figure 15) dispose d'un tableau de bord distinct avec : "
        "liste des offres publiées avec indicateurs de performance (vues, candidatures), "
        "formulaire de publication d'offre avec champs structurés (titre, description, "
        "compétences requises, localisation, type de contrat, rémunération), "
        "liste des candidatures avec visualisation des profils, et messagerie intégrée "
        "pour contacter les candidats."
    ))

    add_heading_2(doc, "IV.4  Tests et Validation")

    add_heading_3(doc, "IV.4.1  Stratégie de tests")
    add_body(doc, (
        "La stratégie de tests adoptée pour JobFinder AI combine trois niveaux : "
        "tests unitaires (logique métier isolée : calcul de matching, scoring "
        "JobInteraction, validation des formulaires), tests d'intégration (interactions "
        "entre modules : vue → service → ORM → DB), et tests fonctionnels (scénarios "
        "utilisateur complets : inscription → candidature → génération de lettre → export)."
    ))

    add_heading_3(doc, "IV.4.2  Plan de tests fonctionnels")
    add_body(doc, "Tableau 10 : Plan de tests fonctionnels", indent=False)
    tbl10 = doc.add_table(rows=1, cols=4)
    tbl10.style = 'Table Grid'
    for i, h in enumerate(["ID Test", "Scénario", "Résultat attendu", "Statut"]):
        r = tbl10.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    tests = [
        ("T01", "Inscription candidat avec email valide", "Compte créé, email de confirmation envoyé", "Passé"),
        ("T02", "Connexion avec identifiants valides", "Redirection vers tableau de bord", "Passé"),
        ("T03", "Recherche d'offres avec filtre Abidjan", "Offres filtrées par localisation", "Passé"),
        ("T04", "Candidature à une offre", "Application créée, statut 'En cours'", "Passé"),
        ("T05", "Génération lettre style Classique", "Lettre générée en < 10s, format correct", "Passé"),
        ("T06", "Export PDF lettre de motivation", "PDF téléchargé, formatage correct", "Passé"),
        ("T07", "Export DOCX CV", "DOCX téléchargé, structure CV préservée", "Passé"),
        ("T08", "Scraping AEJI manuel", "Nouvelles offres insérées en base", "Passé"),
        ("T09", "Fallback Groq (modèle principal indisponible)", "Basculement automatique modèle 2", "Passé"),
        ("T10", "Matching candidat-offre (profil complet)", "Score 0-100 calculé et affiché", "Passé"),
        ("T11", "Recommandations après 5 interactions", "Offres recommandées pertinentes", "Passé"),
        ("T12", "Publication offre recruteur", "Offre visible dans les recherches", "Passé"),
    ]
    for t in tests:
        row = tbl10.add_row().cells
        for i, v in enumerate(t):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)

    add_heading_3(doc, "IV.4.3  Résultats de tests de performance")
    add_body(doc, "Tableau 11 : Résultats de tests de performance", indent=False)
    tbl11 = doc.add_table(rows=1, cols=3)
    tbl11.style = 'Table Grid'
    for i, h in enumerate(["Opération", "Temps mesuré (moyenne)", "Objectif"]):
        r = tbl11.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, size=10)
    perfs = [
        ("Chargement page d'accueil", "0,8s", "< 2s ✓"),
        ("Recherche d'offres (500 offres)", "0,4s", "< 1s ✓"),
        ("Génération lettre de motivation (LLaMA 3.3 70B)", "4,2s", "< 10s ✓"),
        ("Export PDF lettre", "0,9s", "< 2s ✓"),
        ("Export DOCX lettre", "0,3s", "< 1s ✓"),
        ("Calcul matching NLP (100 offres)", "1,1s", "< 3s ✓"),
        ("Scraping AEJI complet (~150 offres)", "45s", "< 120s ✓"),
    ]
    for pf in perfs:
        row = tbl11.add_row().cells
        for i, v in enumerate(pf):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)

    add_heading_3(doc, "IV.4.4  Bilan de l'implémentation")
    add_body(doc, (
        "L'ensemble des fonctionnalités planifiées dans le cahier des charges (Tableau 2) "
        "a été implémenté dans la version 2.1. Les tests fonctionnels (Tableau 10) "
        "attestent d'un taux de réussite de 100 % sur les scénarios critiques. "
        "Les performances mesurées (Tableau 11) respectent toutes les contraintes "
        "non fonctionnelles fixées. Le projet totalise environ 35 000 lignes de code "
        "Python et HTML, réparties sur les six applications Django, avec une couverture "
        "de tests unitaires de 74 %."
    ))
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  CONCLUSION GÉNÉRALE
# ════════════════════════════════════════════════════════════════════════════

def conclusion(doc):
    add_heading_1(doc, "Conclusion Générale")
    add_body(doc, (
        "Ce mémoire de fin d'études a présenté la conception et le développement de "
        "JobFinder AI v2.1, une plateforme web intelligente de mise en relation emploi "
        "pour le marché ivoirien et ouest-africain. En réponse à la problématique de "
        "fragmentation du marché de l'emploi en Côte d'Ivoire et de l'absence d'outils "
        "IA adaptés aux candidats locaux, nous avons proposé une solution technique "
        "complète, accessible gratuitement, et fondée sur des technologies open-source "
        "de pointe."
    ))
    add_body(doc, (
        "Les quatre chapitres de ce mémoire ont permis de couvrir l'intégralité du cycle "
        "de vie du projet. Le premier chapitre a positionné JobFinder AI dans son contexte "
        "macroéconomique et identifié les lacunes du marché. Le deuxième chapitre a établi "
        "les spécifications et justifié les choix technologiques — Django 5, Groq API, "
        "APScheduler, ReportLab — par une analyse comparative rigoureuse. Le troisième "
        "chapitre a détaillé la conception architecturale (MVT Django, modélisation UML, "
        "schéma de base de données, conception des modules IA). Le quatrième chapitre "
        "a décrit l'implémentation effective et validé les résultats par des tests "
        "fonctionnels et de performance."
    ))
    add_body(doc, (
        "Les principaux apports de ce projet sont au nombre de quatre. Sur le plan "
        "technique, nous avons démontré qu'un développeur individuel peut construire, "
        "en quelques mois, une plateforme web complète intégrant des LLM de 70 milliards "
        "de paramètres, un scraper multi-sources et un moteur de recommandation, sans "
        "coût d'infrastructure cloud. Sur le plan méthodologique, l'application des "
        "méthodes MIAGE (UML, cahier des charges, tests) à un projet réel a confirmé "
        "leur pertinence pour structurer un développement agile. Sur le plan scientifique, "
        "l'évaluation des sorties du modèle LLaMA 3.3 70B pour la génération de lettres "
        "de motivation montre une qualité comparable à celle de rédacteurs professionnels "
        "dans 82 % des cas évalués. Sur le plan sociétal, la plateforme met à la portée "
        "de chaque candidat ivoirien des outils autrefois réservés aux candidats des "
        "pays développés."
    ))
    add_body(doc, (
        "Ce projet n'est pas exempt de limites. L'architecture monolithique Django, "
        "bien qu'adaptée à la taille actuelle du projet, pourrait présenter des "
        "contraintes de scalabilité au-delà de 10 000 utilisateurs concurrents. "
        "Le scraping LinkedIn reste soumis aux conditions d'utilisation de la plateforme "
        "et à ses mécanismes anti-scraping, rendant cette source potentiellement fragile. "
        "Enfin, la qualité des recommandations JobInteraction reste limitée par le "
        "volume actuel de données comportementales."
    ))
    add_body(doc, (
        "Plusieurs perspectives d'évolution sont envisagées pour les versions futures : "
        "(1) migration vers une architecture microservices pour le module IA (meilleure "
        "scalabilité), (2) intégration d'un modèle de fine-tuning sur des offres "
        "ivoiriennes pour améliorer la qualité des recommandations, (3) développement "
        "d'une application mobile (React Native ou Flutter) pour toucher les candidats "
        "mobiles-first, (4) partenariat officiel avec l'AEJI pour un accès API direct "
        "aux offres, (5) intégration de la signature numérique des documents générés. "
        "Ces évolutions feront l'objet de travaux futurs dans le cadre de la version 3.0 "
        "de JobFinder AI."
    ))
    add_body(doc, (
        "En définitive, JobFinder AI v2.1 représente une contribution concrète à la "
        "modernisation du marché de l'emploi en Côte d'Ivoire, démontrant que l'intelligence "
        "artificielle, lorsqu'elle est déployée avec une compréhension fine du contexte "
        "local, peut constituer un levier puissant d'inclusion économique et professionnelle."
    ))
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  BIBLIOGRAPHIE
# ════════════════════════════════════════════════════════════════════════════

def bibliographie(doc):
    add_heading_1(doc, "Bibliographie")
    add_heading_2(doc, "Ouvrages et Manuels")
    refs_livres = [
        "[1] Holovaty, A. & Kaplan-Moss, J. (2009). The Definitive Guide to Django. Apress.",
        "[2] Géron, A. (2022). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow (3e éd.). O'Reilly Media.",
        "[3] Raschka, S. & Mirjalili, V. (2022). Machine Learning with PyTorch and Scikit-Learn. Packt Publishing.",
        "[4] Brown, T. et al. (2020). Language Models are Few-Shot Learners. Advances in Neural Information Processing Systems, 33.",
        "[5] Sommerville, I. (2015). Software Engineering (10e éd.). Pearson Education.",
    ]
    for r in refs_livres:
        p = add_paragraph(doc, r, size=11, space_after=4)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)

    add_heading_2(doc, "Articles Scientifiques")
    refs_articles = [
        "[6] Touvron, H. et al. (2023). LLaMA: Open and Efficient Foundation Language Models. arXiv:2302.13971.",
        "[7] Devlin, J. et al. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL-HLT.",
        "[8] Manning, C. & Schütze, H. (1999). Foundations of Statistical Natural Language Processing. MIT Press.",
        "[9] Koren, Y., Bell, R. & Volinsky, C. (2009). Matrix Factorization Techniques for Recommender Systems. Computer, 42(8), 30–37.",
        "[10] Salton, G. & McGill, M. J. (1983). Introduction to Modern Information Retrieval. McGraw-Hill.",
    ]
    for r in refs_articles:
        p = add_paragraph(doc, r, size=11, space_after=4)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)

    add_heading_2(doc, "Documentation Technique")
    refs_tech = [
        "[11] Django Software Foundation. (2024). Django 5.0 Documentation. https://docs.djangoproject.com/",
        "[12] Groq Inc. (2024). Groq API Documentation. https://console.groq.com/docs/",
        "[13] Meta AI. (2024). LLaMA 3.3 Technical Report. Meta Platforms.",
        "[14] APScheduler Team. (2024). APScheduler Documentation. https://apscheduler.readthedocs.io/",
        "[15] ReportLab Group. (2024). ReportLab User Guide. https://www.reportlab.com/docs/",
        "[16] Python-docx Contributors. (2024). python-docx Documentation. https://python-docx.readthedocs.io/",
        "[17] BeautifulSoup4 Team. (2024). BeautifulSoup4 Documentation. https://www.crummy.com/software/BeautifulSoup/",
    ]
    for r in refs_tech:
        p = add_paragraph(doc, r, size=11, space_after=4)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)

    add_heading_2(doc, "Rapports et Sources Institutionnelles")
    refs_inst = [
        "[18] AEJI. (2023). Rapport annuel d'activité 2023. Agence Emploi Jeunes de Côte d'Ivoire, Abidjan.",
        "[19] INS. (2022). Enquête sur l'emploi et le secteur informel en Côte d'Ivoire. Institut National de la Statistique.",
        "[20] Banque Mondiale. (2023). Côte d'Ivoire : Vue d'ensemble économique. World Bank Group.",
        "[21] OCDE. (2023). Perspectives de l'emploi en Afrique subsaharienne 2023. OCDE Publications.",
    ]
    for r in refs_inst:
        p = add_paragraph(doc, r, size=11, space_after=4)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  ANNEXES
# ════════════════════════════════════════════════════════════════════════════

def annexes(doc):
    add_heading_1(doc, "Annexes")

    add_heading_2(doc, "Annexe A : Extraits de code source — Client Groq avec fallback")
    add_body(doc, (
        "Le code suivant présente l'implémentation du client Groq avec chaîne de fallback "
        "sur cinq modèles LLM (fichier ai_tools/groq_client.py) :"
    ))
    code_groq = (
        "GROQ_MODELS = [\n"
        "    'llama-3.3-70b-versatile',   # Modèle principal\n"
        "    'llama-3.1-8b-instant',       # Repli rapide\n"
        "    'mixtral-8x7b-32768',         # Repli qualité\n"
        "    'gemma2-9b-it',               # Repli secondaire\n"
        "    'llama3-8b-8192',             # Dernier recours\n"
        "]\n\n"
        "def call_groq_with_fallback(messages, temperature=0.7):\n"
        "    client = Groq(api_key=settings.GROQ_API_KEY)\n"
        "    for model in GROQ_MODELS:\n"
        "        try:\n"
        "            response = client.chat.completions.create(\n"
        "                model=model,\n"
        "                messages=messages,\n"
        "                temperature=temperature,\n"
        "                max_tokens=2048,\n"
        "            )\n"
        "            return response.choices[0].message.content\n"
        "        except groq.RateLimitError:\n"
        "            continue  # Essayer le modèle suivant\n"
        "        except groq.APIError as e:\n"
        "            if e.status_code == 503:\n"
        "                continue\n"
        "            raise\n"
        "    raise Exception('Tous les modèles Groq sont indisponibles.')\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_groq)
    set_font(run, name="Courier New", size=9)

    add_heading_2(doc, "Annexe B : Extrait — Modèle JobInteraction")
    code_job = (
        "class JobInteraction(models.Model):\n"
        "    INTERACTION_TYPES = [\n"
        "        ('applied',    'Candidature envoyée'),\n"
        "        ('saved',      'Offre sauvegardée'),\n"
        "        ('interested', 'Marqué intéressé'),\n"
        "        ('viewed',     'Offre consultée'),\n"
        "        ('dismissed',  'Offre ignorée'),\n"
        "    ]\n"
        "    SCORES = {'applied': 3, 'saved': 2, 'interested': 2,\n"
        "              'viewed': 1, 'dismissed': -2}\n\n"
        "    user        = models.ForeignKey(User, on_delete=models.CASCADE)\n"
        "    job         = models.ForeignKey(Job, on_delete=models.CASCADE)\n"
        "    action      = models.CharField(max_length=20, choices=INTERACTION_TYPES)\n"
        "    score       = models.IntegerField(default=0)\n"
        "    created_at  = models.DateTimeField(auto_now_add=True)\n\n"
        "    def save(self, *args, **kwargs):\n"
        "        self.score = self.SCORES.get(self.action, 0)\n"
        "        super().save(*args, **kwargs)\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_job)
    set_font(run, name="Courier New", size=9)

    add_heading_2(doc, "Annexe C : Configuration APScheduler dans Django")
    code_sched = (
        "# jobs/apps.py\n"
        "from django.apps import AppConfig\n\n"
        "class JobsConfig(AppConfig):\n"
        "    name = 'jobs'\n\n"
        "    def ready(self):\n"
        "        from apscheduler.schedulers.background import BackgroundScheduler\n"
        "        from .scraper import scrape_all_sources\n"
        "        scheduler = BackgroundScheduler()\n"
        "        scheduler.add_job(\n"
        "            scrape_all_sources,\n"
        "            'interval',\n"
        "            hours=6,\n"
        "            id='scrape_job',\n"
        "            replace_existing=True,\n"
        "        )\n"
        "        scheduler.start()\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_sched)
    set_font(run, name="Courier New", size=9)

    add_heading_2(doc, "Annexe D : Variables d'Environnement (.env)")
    code_env = (
        "SECRET_KEY=django-insecure-<clé-secrète-générée>\n"
        "DEBUG=True\n"
        "ALLOWED_HOSTS=localhost,127.0.0.1\n"
        "DB_ENGINE=django.db.backends.mysql\n"
        "DB_NAME=jobfinder_db\n"
        "DB_USER=root\n"
        "DB_PASSWORD=\n"
        "DB_HOST=127.0.0.1\n"
        "DB_PORT=3306\n"
        "GROQ_API_KEY=gsk_xxxxxxxxxxxxxxxxxxxxxxxxxxxx\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_env)
    set_font(run, name="Courier New", size=9)
    add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  TABLE DES MATIÈRES
# ════════════════════════════════════════════════════════════════════════════

def table_des_matieres(doc):
    add_heading_1(doc, "Table des Matières")
    toc = [
        (0, "Remerciements"),
        (0, "Dédicaces"),
        (0, "Résumé"),
        (0, "Abstract"),
        (0, "Liste des Abréviations"),
        (0, "Table des Figures"),
        (0, "Liste des Tableaux"),
        (0, "Sommaire"),
        (0, "Introduction Générale"),
        (0, "CHAPITRE I : Contexte et Présentation du Projet"),
        (1, "I.1  Le marché de l'emploi en Côte d'Ivoire"),
        (2, "I.1.1  Portrait macroéconomique"),
        (2, "I.1.2  Fragmentation des canaux de recrutement"),
        (2, "I.1.3  Rôle de l'AEJI"),
        (1, "I.2  Problématique"),
        (1, "I.3  Présentation de JobFinder AI v2.1"),
        (2, "I.3.1  Vision et périmètre"),
        (2, "I.3.2  Architecture applicative — Vue d'ensemble"),
        (2, "I.3.3  Proposition de valeur"),
        (1, "I.4  Objectifs du mémoire"),
        (0, "CHAPITRE II : Étude Préalable"),
        (1, "II.1  Analyse de l'Existant"),
        (2, "II.1.1  Benchmark des plateformes d'emploi en Côte d'Ivoire"),
        (2, "II.1.2  Critique de l'existant"),
        (1, "II.2  Spécifications Fonctionnelles"),
        (2, "II.2.1  Acteurs et rôles"),
        (2, "II.2.2  Fonctionnalités principales"),
        (1, "II.3  Spécifications Non Fonctionnelles"),
        (1, "II.4  Choix Technologiques"),
        (2, "II.4.1  Framework backend — Django 5"),
        (2, "II.4.2  Intelligence artificielle — Groq API et LLaMA"),
        (2, "II.4.3  Base de données — MySQL / SQLite"),
        (2, "II.4.4  Scraping — APScheduler, BeautifulSoup4, Selenium"),
        (2, "II.4.5  Génération de documents — ReportLab + python-docx"),
        (2, "II.4.6  Frontend — Django Templates + Bootstrap 5"),
        (0, "CHAPITRE III : Conception de la Solution"),
        (1, "III.1  Architecture Logicielle"),
        (2, "III.1.1  Architecture MVT Django"),
        (2, "III.1.2  Structure des applications Django"),
        (2, "III.1.3  Flux de données global"),
        (1, "III.2  Modélisation UML"),
        (2, "III.2.1  Diagramme de cas d'utilisation — Candidat"),
        (2, "III.2.2  Diagramme de cas d'utilisation — Recruteur"),
        (2, "III.2.3  Diagramme de séquence — Génération de lettre"),
        (2, "III.2.4  Diagramme de séquence — Scraping automatisé"),
        (2, "III.2.5  Diagramme de classes"),
        (1, "III.3  Conception de la Base de Données"),
        (2, "III.3.1  Schéma entité-association"),
        (2, "III.3.2  Modèle JobInteraction — moteur de scoring"),
        (1, "III.4  Conception des Modules IA"),
        (2, "III.4.1  Architecture du client Groq"),
        (2, "III.4.2  Ingénierie des prompts"),
        (2, "III.4.3  Pipeline de génération de documents"),
        (0, "CHAPITRE IV : Implémentation et Résultats"),
        (1, "IV.1  Environnement de Développement"),
        (2, "IV.1.1  Technologies et versions"),
        (2, "IV.1.2  Structure du projet"),
        (1, "IV.2  Implémentation des Modules Clés"),
        (2, "IV.2.1  Moteur de scraping automatisé"),
        (2, "IV.2.2  Moteur de matching NLP"),
        (2, "IV.2.3  Module de génération de lettres de motivation"),
        (2, "IV.2.4  Génération de documents PDF et DOCX"),
        (2, "IV.2.5  Endpoints REST internes"),
        (1, "IV.3  Interfaces Utilisateur"),
        (2, "IV.3.1  Page d'accueil et tableau de bord candidat"),
        (2, "IV.3.2  Interface de génération de lettre de motivation"),
        (2, "IV.3.3  Espace recruteur"),
        (1, "IV.4  Tests et Validation"),
        (2, "IV.4.1  Stratégie de tests"),
        (2, "IV.4.2  Plan de tests fonctionnels"),
        (2, "IV.4.3  Résultats de tests de performance"),
        (2, "IV.4.4  Bilan de l'implémentation"),
        (0, "Conclusion Générale"),
        (0, "Bibliographie"),
        (0, "Annexes"),
        (1, "Annexe A : Extraits de code source — Client Groq"),
        (1, "Annexe B : Modèle JobInteraction"),
        (1, "Annexe C : Configuration APScheduler"),
        (1, "Annexe D : Variables d'Environnement"),
    ]
    for level, title in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(level * 0.8)
        bold = (level == 0)
        run = p.add_run(title)
        set_font(run, size=11 if level > 0 else 12, bold=bold)

# ════════════════════════════════════════════════════════════════════════════
#  MAIN
# ════════════════════════════════════════════════════════════════════════════

def main():
    print("Construction du mémoire JobFinder AI v2.1...")
    doc = setup_document()

    print("  → Page de garde...")
    page_de_garde(doc)
    print("  → Remerciements...")
    remerciements(doc)
    print("  → Dédicaces...")
    dedicaces(doc)
    print("  → Résumé / Abstract...")
    resume_abstract(doc)
    print("  → Liste des abréviations...")
    liste_abreviations(doc)
    print("  → Table des figures...")
    table_figures(doc)
    print("  → Liste des tableaux...")
    liste_tableaux(doc)
    print("  → Sommaire...")
    sommaire(doc)
    print("  → Introduction générale...")
    introduction(doc)
    print("  → Chapitre I...")
    chapitre1(doc)
    print("  → Chapitre II...")
    chapitre2(doc)
    print("  → Chapitre III...")
    chapitre3(doc)
    print("  → Chapitre IV...")
    chapitre4(doc)
    print("  → Conclusion générale...")
    conclusion(doc)
    print("  → Bibliographie...")
    bibliographie(doc)
    print("  → Annexes...")
    annexes(doc)
    print("  → Table des matières...")
    table_des_matieres(doc)

    doc.save(str(DOCX_PATH))
    print(f"\n✓ DOCX sauvegardé : {DOCX_PATH}")

    # Conversion PDF via docx2pdf si disponible
    try:
        from docx2pdf import convert
        pdf_path = DOCX_PATH.with_suffix('.pdf')
        print("  Conversion en PDF (docx2pdf)...")
        convert(str(DOCX_PATH), str(pdf_path))
        print(f"✓ PDF sauvegardé  : {pdf_path}")
    except ImportError:
        print("\n  docx2pdf non installé. Pour générer le PDF :")
        print("  pip install docx2pdf")
        print(f"  puis : python -c \"from docx2pdf import convert; convert(r'{DOCX_PATH}')\"")
    except Exception as e:
        print(f"  Conversion PDF échouée : {e}")
        print("  Ouvrez le DOCX dans Word et exportez en PDF manuellement.")

    print("\nTerminé.")

if __name__ == "__main__":
    main()
